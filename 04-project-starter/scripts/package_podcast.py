#!/usr/bin/env python3
"""
NotebookLM Podcast Packager for UHN Accessibility First Course Series
Prepares and validates podcast source documents for optimal NotebookLM generation.

Usage:
    python package_podcast.py --guide 1                     # Validate + package Guide 1
    python package_podcast.py --guide 1 --enhance           # Add metadata headers for NotebookLM
    python package_podcast.py --guide 1 --word              # Export as .docx (NotebookLM prefers Word)
    python package_podcast.py --guide 1 --stats             # Show document statistics
    python package_podcast.py --all --stats                 # Stats for all available guides
"""

import argparse
import re
import sys
from pathlib import Path

BASE_DIR = Path(__file__).resolve().parent.parent.parent
BUILD_OUTPUT = BASE_DIR / "05-build-output"

GUIDE_FOLDERS = {
    1: "01-Foundations-of-Disability-Inclusion-and-Accessible-Design",
    2: "02-Perceptions-Attitudes-and-Barriers",
    3: "03-Vision-Disabilities",
    4: "04-Sensory-Hearing-and-Communication-Disabilities",
    5: "05-Physical-Disabilities-and-Mobility",
    6: "06-Mental-Health-Disabilities",
    7: "07-Intellectual-Developmental-and-Learning-Disabilities",
    8: "08-Non-Visible-Disabilities",
    9: "09-Aging-Disability-and-Intersectionality",
    10: "10-Engaging-with-Confidence-and-Respect",
    11: "11-Service-Animals-Guide-Dogs-and-Non-Service-Animals",
    12: "12-Support-Persons",
    13: "13-Assistive-Devices",
    14: "14-Communication-and-Information-Accessibility",
    15: "15-Neurodiversity-and-Sensory-Regulation",
    16: "16-Trauma-Informed-Accessibility",
    17: "17-Accessibility-in-Crisis-Situations-and-De-escalation",
    18: "18-Indigenous-Peoples-and-Accessibility",
}

# NotebookLM optimal ranges
OPTIMAL_WORD_MIN = 1500
OPTIMAL_WORD_MAX = 5000
OPTIMAL_PODCAST_MIN_MIN = 8
OPTIMAL_PODCAST_MAX_MIN = 20


def analyze_document(filepath: Path) -> dict:
    """Analyze a podcast source document for NotebookLM readiness."""
    with open(filepath, "r", encoding="utf-8") as f:
        content = f.read()

    words = content.split()
    word_count = len(words)
    sentences = re.split(r'[.!?]+', content)
    sentence_count = len([s for s in sentences if s.strip()])
    paragraphs = [p for p in content.split("\n\n") if p.strip()]
    paragraph_count = len(paragraphs)

    # Check for elements NotebookLM likes
    has_questions = bool(re.search(r'\?', content))
    question_count = len(re.findall(r'\?', content))
    has_statistics = bool(re.search(r'\d+\s*(?:percent|%)', content, re.IGNORECASE))
    has_stories = bool(re.search(r'(?:story|scenario|imagine|picture this|consider)', content, re.IGNORECASE))
    has_headers = bool(re.search(r'^#+\s', content, re.MULTILINE))
    header_count = len(re.findall(r'^#+\s', content, re.MULTILINE))

    # Conversational markers NotebookLM responds well to
    conversational_markers = [
        r"did you know",
        r"here's (?:the thing|what's interesting|what matters)",
        r"let's (?:think about|consider|look at)",
        r"you might (?:be surprised|not realize|wonder)",
        r"the truth is",
        r"what if",
        r"imagine",
    ]
    conv_count = sum(
        len(re.findall(pattern, content, re.IGNORECASE))
        for pattern in conversational_markers
    )

    # Estimated podcast duration (rough: 150 words/min of source → ~3:1 compression)
    est_podcast_minutes = word_count / 150 * 0.8  # NotebookLM typically covers ~80% of content

    # Quality score
    issues = []
    score = 100

    if word_count < OPTIMAL_WORD_MIN:
        issues.append(f"Too short ({word_count} words, need {OPTIMAL_WORD_MIN}+)")
        score -= 20
    elif word_count > OPTIMAL_WORD_MAX:
        issues.append(f"Too long ({word_count} words, max {OPTIMAL_WORD_MAX})")
        score -= 10

    if question_count < 3:
        issues.append(f"Few questions ({question_count}) — add more for discussion triggers")
        score -= 10

    if not has_statistics:
        issues.append("No statistics found — add data points for engaging discussion")
        score -= 10

    if not has_stories:
        issues.append("No story/scenario markers — add narrative elements")
        score -= 15

    if conv_count < 2:
        issues.append(f"Low conversational tone ({conv_count} markers) — add 'Did you know...', 'Here's what's interesting...'")
        score -= 10

    if header_count < 3:
        issues.append(f"Few section headers ({header_count}) — add structure for topic transitions")
        score -= 5

    # Check for bullet-heavy content (NotebookLM prefers prose)
    bullet_lines = len(re.findall(r'^\s*[-*•]\s', content, re.MULTILINE))
    prose_ratio = 1 - (bullet_lines / max(len(content.split('\n')), 1))
    if prose_ratio < 0.6:
        issues.append(f"Too many bullet points ({bullet_lines} lines) — convert to flowing prose")
        score -= 15

    return {
        "word_count": word_count,
        "sentence_count": sentence_count,
        "paragraph_count": paragraph_count,
        "header_count": header_count,
        "question_count": question_count,
        "has_statistics": has_statistics,
        "has_stories": has_stories,
        "conversational_markers": conv_count,
        "bullet_lines": bullet_lines,
        "prose_ratio": prose_ratio,
        "est_podcast_minutes": est_podcast_minutes,
        "score": max(score, 0),
        "issues": issues,
    }


def enhance_document(filepath: Path, guide_num: int) -> str:
    """Add NotebookLM optimization headers and metadata to the document."""
    with open(filepath, "r", encoding="utf-8") as f:
        content = f.read()

    guide_name = GUIDE_FOLDERS.get(guide_num, "").replace("-", " ").lstrip("0123456789-").strip()

    header = f"""# {guide_name} — Accessibility First Series

## About This Document

This document is a companion resource for Guide {guide_num} of the Accessibility First Course Series at University Health Network (UHN) in Toronto, Ontario, Canada. The series helps healthcare employees understand and apply accessibility principles in their daily work, aligned with the Accessibility for Ontarians with Disabilities Act (AODA) and the Ontario Human Rights Code (OHRC).

This guide is designed to be read as a narrative — exploring key concepts, real-world scenarios, and practical takeaways that healthcare workers can apply immediately.

---

"""

    # Check if document already has this header
    if "About This Document" not in content:
        # Remove any existing title if it's just a simple header
        content = re.sub(r'^#\s+.*?\n', '', content, count=1)
        content = header + content.strip()

    # Add discussion prompts section if not present
    if "Discussion" not in content and "discussion" not in content.lower():
        content += """

---

## Questions to Consider

After listening to or reading this guide, reflect on these questions:

- What is one thing you learned that changes how you think about accessibility in your role?
- Can you think of a time when a process or system at your workplace created an unintended barrier?
- How would you apply the Accessibility Decision Path in your next shift?
- What is one concrete action you could take this week to make your workspace more accessible?

These questions are also part of the My Action Planning (MAP) activity in the full course.
"""

    return content


def export_to_word(content: str, output_path: Path) -> None:
    """Export content to a simple .docx file for NotebookLM upload."""
    try:
        from docx import Document
        from docx.shared import Pt, Inches
    except ImportError:
        print("Error: python-docx not installed. Run: pip install python-docx")
        print("Alternatively, copy the .md file content into Google Docs or Word manually.")
        sys.exit(1)

    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Parse markdown-ish content into Word
    for line in content.split('\n'):
        line = line.rstrip()

        if line.startswith('# '):
            p = doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            p = doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            p = doc.add_heading(line[4:], level=3)
        elif line.startswith('---'):
            doc.add_paragraph('_' * 50)
        elif line.startswith('- ') or line.startswith('* '):
            doc.add_paragraph(line[2:], style='List Bullet')
        elif line.strip():
            # Handle bold markdown
            clean_line = re.sub(r'\*\*(.+?)\*\*', r'\1', line)
            clean_line = re.sub(r'\*(.+?)\*', r'\1', clean_line)
            doc.add_paragraph(clean_line)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))


def main():
    parser = argparse.ArgumentParser(description="Package podcast source docs for NotebookLM")
    parser.add_argument("--guide", type=int, help="Guide number (1-18)")
    parser.add_argument("--all", action="store_true", help="Process all available guides")
    parser.add_argument("--stats", action="store_true", help="Show document statistics and quality score")
    parser.add_argument("--enhance", action="store_true", help="Add NotebookLM optimization headers")
    parser.add_argument("--word", action="store_true", help="Export as .docx for NotebookLM upload")
    args = parser.parse_args()

    if not args.guide and not args.all:
        print("Error: Specify --guide N or --all")
        sys.exit(1)

    guides_to_process = range(1, 19) if args.all else [args.guide]

    for guide_num in guides_to_process:
        if guide_num not in GUIDE_FOLDERS:
            continue

        guide_folder = BUILD_OUTPUT / GUIDE_FOLDERS[guide_num]
        podcast_dir = guide_folder / "08-notebooklm-podcast"
        source_file = podcast_dir / f"PODCAST-SOURCE-GUIDE-{guide_num:02d}.md"

        if not source_file.exists():
            if not args.all:
                print(f"Error: Podcast source not found: {source_file}")
            continue

        print(f"\n{'=' * 50}")
        print(f"Guide {guide_num:02d}: {GUIDE_FOLDERS[guide_num]}")
        print(f"{'=' * 50}")

        # Stats
        if args.stats:
            stats = analyze_document(source_file)
            print(f"\n  Word count:      {stats['word_count']}")
            print(f"  Sentences:       {stats['sentence_count']}")
            print(f"  Paragraphs:      {stats['paragraph_count']}")
            print(f"  Section headers: {stats['header_count']}")
            print(f"  Questions:       {stats['question_count']}")
            print(f"  Has statistics:  {'Yes' if stats['has_statistics'] else 'No'}")
            print(f"  Has stories:     {'Yes' if stats['has_stories'] else 'No'}")
            print(f"  Conv. markers:   {stats['conversational_markers']}")
            print(f"  Prose ratio:     {stats['prose_ratio']:.0%}")
            print(f"  Est. podcast:    ~{stats['est_podcast_minutes']:.0f} min")
            print(f"\n  Quality Score:   {stats['score']}/100")

            if stats['issues']:
                print(f"\n  Issues:")
                for issue in stats['issues']:
                    print(f"    ⚠ {issue}")
            else:
                print(f"\n  ✓ Document is well-optimized for NotebookLM")

        # Enhance
        if args.enhance:
            enhanced = enhance_document(source_file, guide_num)
            enhanced_path = podcast_dir / f"PODCAST-SOURCE-GUIDE-{guide_num:02d}-enhanced.md"
            enhanced_path.parent.mkdir(parents=True, exist_ok=True)
            with open(enhanced_path, "w", encoding="utf-8") as f:
                f.write(enhanced)
            print(f"\n  Enhanced version saved: {enhanced_path.name}")

        # Word export
        if args.word:
            if args.enhance:
                content = enhance_document(source_file, guide_num)
            else:
                with open(source_file, "r", encoding="utf-8") as f:
                    content = f.read()

            word_path = podcast_dir / f"PODCAST-SOURCE-GUIDE-{guide_num:02d}.docx"
            export_to_word(content, word_path)
            print(f"  Word export saved: {word_path.name}")

    print("\nDone!")


if __name__ == "__main__":
    main()
