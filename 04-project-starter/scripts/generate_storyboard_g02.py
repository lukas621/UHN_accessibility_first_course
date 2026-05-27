#!/usr/bin/env python3
"""
Master Storyboard DOCX Generator — UHN Accessibility First Series
Guide 02: Perceptions, Attitudes, and Barriers

Generates a fully-formatted Word document matching the v0.5.1 storyboard
reference format: Part 1 (review) + Part 2 (designer toolkit), 3-column tables.

Usage:
    python generate_storyboard_g02.py

Output:
    05-build-output/02-Perceptions-Attitudes-and-Barriers/
        02-production/master-storyboard/MASTER-STORYBOARD-GUIDE-02.docx

Requirements:
    pip install python-docx
"""

import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor, Emu
    from docx.enum.section import WD_ORIENT
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("Error: 'python-docx' not installed. Run: pip install python-docx")
    sys.exit(1)

# ---------------------------------------------------------------------------
# Paths
# ---------------------------------------------------------------------------
BASE_DIR = Path(__file__).resolve().parent.parent.parent
LOGO_PATH = BASE_DIR / "02-branding-and-style" / "logos" / "uhn_logo_2.png"
OUTPUT_DIR = (
    BASE_DIR
    / "05-build-output"
    / "02-Perceptions-Attitudes-and-Barriers"
    / "02-production"
    / "master-storyboard"
)
OUTPUT_PATH = OUTPUT_DIR / "MASTER-STORYBOARD-GUIDE-02.docx"

# ---------------------------------------------------------------------------
# Colour constants (hex strings — no leading #)
# ---------------------------------------------------------------------------
NAVY        = "192858"
COBALT      = "255CAA"
BODY        = "2B2B2B"
LIGHT_FILL  = "F5F5F5"
HEADER_FILL = "EDF0F7"
BORDER_CLR  = "CCCCCC"
WHITE       = "FFFFFF"
RED         = "C0233B"

# ---------------------------------------------------------------------------
# Helper: apply shading to a table cell
# ---------------------------------------------------------------------------
def shade_cell(cell, fill_hex: str):
    """Set the background fill of a table cell."""
    tc_pr = cell._element.get_or_add_tcPr()
    for existing in tc_pr.findall(qn("w:shd")):
        tc_pr.remove(existing)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  fill_hex.upper())
    tc_pr.append(shd)


def set_cell_margins(cell, top_pt=6, bottom_pt=6, left_pt=8, right_pt=8):
    """Set cell internal margins in points (converted to twips)."""
    tc_pr = cell._element.get_or_add_tcPr()
    mar = OxmlElement("w:tcMar")
    for side, val in [("top", top_pt), ("bottom", bottom_pt),
                      ("left", left_pt), ("right", right_pt)]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"), str(int(val * 20)))
        el.set(qn("w:type"), "dxa")
        mar.append(el)
    for existing in tc_pr.findall(qn("w:tcMar")):
        tc_pr.remove(existing)
    tc_pr.append(mar)


def set_table_borders(table, color_hex: str = BORDER_CLR, size_pt: int = 4):
    """Apply thin uniform borders to every cell in a table."""
    border_size = str(size_pt * 8)
    for row in table.rows:
        for cell in row.cells:
            tc_pr = cell._element.get_or_add_tcPr()
            tc_borders = OxmlElement("w:tcBorders")
            for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
                border = OxmlElement(f"w:{side}")
                border.set(qn("w:val"),   "single")
                border.set(qn("w:sz"),    border_size)
                border.set(qn("w:space"), "0")
                border.set(qn("w:color"), color_hex.upper())
                tc_borders.append(border)
            for existing in tc_pr.findall(qn("w:tcBorders")):
                tc_pr.remove(existing)
            tc_pr.append(tc_borders)


def set_col_widths(table, widths_inches: list):
    """Set column widths for a table (list of floats in inches)."""
    for i, row in enumerate(table.rows):
        for j, cell in enumerate(row.cells):
            if j < len(widths_inches):
                tc = cell._element
                tc_pr = tc.get_or_add_tcPr()
                tc_w = OxmlElement("w:tcW")
                tc_w.set(qn("w:w"),    str(int(widths_inches[j] * 1440)))
                tc_w.set(qn("w:type"), "dxa")
                for existing in tc_pr.findall(qn("w:tcW")):
                    tc_pr.remove(existing)
                tc_pr.append(tc_w)


def rgb(hex_str: str):
    h = hex_str.lstrip("#")
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


def add_run(para, text: str, bold=False, italic=False, size_pt=10,
            color_hex=BODY, font_name="Arial"):
    run = para.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.size       = Pt(size_pt)
    run.font.color.rgb  = rgb(color_hex)
    run.font.name       = font_name
    return run


def cell_para(cell, text: str, bold=False, italic=False, size_pt=10,
              color_hex=BODY, font_name="Arial", align=WD_ALIGN_PARAGRAPH.LEFT):
    """Clear a cell and add a single paragraph with one run."""
    for p in cell.paragraphs[1:]:
        p._element.getparent().remove(p._element)
    para = cell.paragraphs[0]
    para.clear()
    para.alignment = align
    add_run(para, text, bold=bold, italic=italic, size_pt=size_pt,
            color_hex=color_hex, font_name=font_name)
    return para


def cell_para_multiline(cell, lines: list, size_pt=10, color_hex=BODY,
                        font_name="Arial"):
    """Write multiple paragraphs into a cell. Each item in lines is either
    a plain str or a dict with keys: text, bold, italic, size_pt, color_hex."""
    for p in cell.paragraphs[1:]:
        p._element.getparent().remove(p._element)
    first = True
    for line in lines:
        if first:
            para = cell.paragraphs[0]
            para.clear()
            first = False
        else:
            para = cell.add_paragraph()
        if isinstance(line, dict):
            para.alignment = line.get("align", WD_ALIGN_PARAGRAPH.LEFT)
            add_run(para, line.get("text", ""),
                    bold=line.get("bold", False),
                    italic=line.get("italic", False),
                    size_pt=line.get("size_pt", size_pt),
                    color_hex=line.get("color_hex", color_hex),
                    font_name=line.get("font_name", font_name))
        else:
            add_run(para, str(line), size_pt=size_pt, color_hex=color_hex,
                    font_name=font_name)


# ---------------------------------------------------------------------------
# Document setup
# ---------------------------------------------------------------------------
def setup_document() -> Document:
    doc = Document()
    # Landscape 11 x 8.5 inches
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width  = Inches(11)
    section.page_height = Inches(8.5)
    section.left_margin   = Inches(0.7)
    section.right_margin  = Inches(0.7)
    section.top_margin    = Inches(0.55)
    section.bottom_margin = Inches(0.5)

    # Default paragraph style
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(10)
    style.font.color.rgb = rgb(BODY)

    return doc


# ---------------------------------------------------------------------------
# Reusable block builders
# ---------------------------------------------------------------------------
def add_section_heading(doc, number: str, title: str):
    """Add a numbered section heading (e.g. '1.1 Course Information')."""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(12)
    para.paragraph_format.space_after  = Pt(4)
    add_run(para, f"{number}  {title}", bold=True, size_pt=13,
            color_hex=NAVY, font_name="Arial Black")


def add_part_banner(doc, text: str):
    """Full-width navy banner for PART 1 / PART 2."""
    doc.add_paragraph()  # spacer
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = tbl.rows[0].cells[0]
    shade_cell(cell, NAVY)
    set_cell_margins(cell, top_pt=8, bottom_pt=8, left_pt=12, right_pt=12)
    cell_para(cell, text, bold=True, size_pt=14,
              color_hex=WHITE, font_name="Arial Black",
              align=WD_ALIGN_PARAGRAPH.CENTER)
    set_table_borders(tbl, color_hex=NAVY)
    doc.add_paragraph()  # spacer after


def add_two_col_table(doc, rows_data: list,
                      col_widths=(2.5, 6.8),
                      header_row: list = None):
    """
    Build a 2-column info table.
    rows_data: list of (label, value) tuples.
    header_row: optional [left_header, right_header] for tables with a header row.
    """
    num_rows = len(rows_data) + (1 if header_row else 0)
    tbl = doc.add_table(rows=num_rows, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    row_idx = 0
    if header_row:
        hrow = tbl.rows[row_idx].cells
        shade_cell(hrow[0], NAVY)
        shade_cell(hrow[1], NAVY)
        cell_para(hrow[0], header_row[0], bold=True, size_pt=10,
                  color_hex=WHITE, font_name="Arial Black")
        cell_para(hrow[1], header_row[1], bold=True, size_pt=10,
                  color_hex=WHITE, font_name="Arial Black")
        set_cell_margins(hrow[0])
        set_cell_margins(hrow[1])
        row_idx += 1

    for label, value in rows_data:
        cells = tbl.rows[row_idx].cells
        shade_cell(cells[0], HEADER_FILL)
        shade_cell(cells[1], WHITE)
        cell_para(cells[0], label, bold=True, size_pt=10, color_hex=NAVY)
        if "\n" in str(value):
            lines = []
            for line in str(value).split("\n"):
                lines.append({"text": line, "size_pt": 10, "color_hex": BODY})
            cell_para_multiline(cells[1], lines)
        else:
            cell_para(cells[1], str(value), size_pt=10, color_hex=BODY)
        set_cell_margins(cells[0])
        set_cell_margins(cells[1])
        row_idx += 1

    set_col_widths(tbl, list(col_widths))
    set_table_borders(tbl)
    return tbl


def add_four_col_table(doc, rows_data: list, col_widths=(1.8, 3.5, 1.2, 1.8),
                       headers=None):
    """Build a 4-column table (used for summative assessments)."""
    num_rows = len(rows_data) + (1 if headers else 0)
    tbl = doc.add_table(rows=num_rows, cols=4)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    row_idx = 0
    if headers:
        hcells = tbl.rows[0].cells
        for i, h in enumerate(headers):
            shade_cell(hcells[i], NAVY)
            cell_para(hcells[i], h, bold=True, size_pt=10,
                      color_hex=WHITE, font_name="Arial Black")
            set_cell_margins(hcells[i])
        row_idx += 1

    for r in rows_data:
        cells = tbl.rows[row_idx].cells
        for i, val in enumerate(r):
            shade_cell(cells[i], LIGHT_FILL if i == 0 else WHITE)
            cell_para(cells[i], str(val), bold=(i == 0), size_pt=10,
                      color_hex=NAVY if i == 0 else BODY)
            set_cell_margins(cells[i])
        row_idx += 1

    set_col_widths(tbl, list(col_widths))
    set_table_borders(tbl)
    return tbl


def add_five_col_table(doc, rows_data: list,
                       col_widths=(0.45, 0.65, 2.8, 2.1, 3.3),
                       headers=None):
    """Build a 5-column table (Master Screen Schedule)."""
    num_rows = len(rows_data) + (1 if headers else 0)
    tbl = doc.add_table(rows=num_rows, cols=5)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    row_idx = 0
    if headers:
        hcells = tbl.rows[0].cells
        for i, h in enumerate(headers):
            shade_cell(hcells[i], NAVY)
            cell_para(hcells[i], h, bold=True, size_pt=10,
                      color_hex=WHITE, font_name="Arial Black")
            set_cell_margins(hcells[i])
        row_idx += 1

    for r in rows_data:
        cells = tbl.rows[row_idx].cells
        for i, val in enumerate(r):
            is_first = (i == 0)
            shade_cell(cells[i], HEADER_FILL if is_first else WHITE)
            cell_para(cells[i], str(val), bold=is_first, size_pt=10,
                      color_hex=NAVY if is_first else BODY)
            set_cell_margins(cells[i])
        row_idx += 1

    set_col_widths(tbl, list(col_widths))
    set_table_borders(tbl)
    return tbl


# ---------------------------------------------------------------------------
# Merge cells helper (for banner rows in Part 2 tables)
# ---------------------------------------------------------------------------
def merge_row(table, row_idx: int, start_col: int, end_col: int):
    """Merge cells from start_col to end_col (inclusive) in a given row."""
    row = table.rows[row_idx]
    merged = row.cells[start_col]
    for c in range(start_col + 1, end_col + 1):
        merged = merged.merge(row.cells[c])
    return merged


# ---------------------------------------------------------------------------
# Part 2: Three-column screen table
# ---------------------------------------------------------------------------
def add_screen_table(doc, screen_num: str, screen_title: str,
                     step: str, activities: str, design_guide: str):
    """
    Build the 3-column designer-toolkit table for one screen.
    Layout:
      Row 0: merged header — "SCREEN {num}" | "Topic: {title}"
      Row 1: column headers — Step | Activities | Design Guide
      Row 2: content
    """
    tbl = doc.add_table(rows=3, cols=3)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    # --- Row 0: screen banner ---
    r0 = tbl.rows[0].cells
    left_cell = r0[0].merge(r0[1])
    shade_cell(left_cell, NAVY)
    set_cell_margins(left_cell, top_pt=6, bottom_pt=6, left_pt=8, right_pt=8)
    cell_para(left_cell, f"SCREEN {screen_num}  (Online · Self-Paced)",
              bold=True, size_pt=10, color_hex=WHITE, font_name="Arial Black")

    shade_cell(r0[2], NAVY)
    set_cell_margins(r0[2], top_pt=6, bottom_pt=6, left_pt=8, right_pt=8)
    cell_para(r0[2], f"Topic: {screen_title}",
              bold=True, size_pt=10, color_hex=WHITE, font_name="Arial Black")

    # --- Row 1: column headers ---
    r1 = tbl.rows[1].cells
    headers = ["Step", "Activities", "Design Guide"]
    for i, h in enumerate(headers):
        shade_cell(r1[i], "EDF0F7")
        cell_para(r1[i], h, bold=True, size_pt=10, color_hex=NAVY,
                  font_name="Arial Black")
        set_cell_margins(r1[i])

    # --- Row 2: content ---
    r2 = tbl.rows[2].cells

    # Step column
    shade_cell(r2[0], LIGHT_FILL)
    cell_para(r2[0], step, bold=True, size_pt=10, color_hex=NAVY)
    set_cell_margins(r2[0])

    # Activities column
    shade_cell(r2[1], WHITE)
    _fill_cell_lines(r2[1], activities, size_pt=10, color_hex=BODY)
    set_cell_margins(r2[1])

    # Design Guide column
    shade_cell(r2[2], LIGHT_FILL)
    _fill_cell_lines(r2[2], design_guide, size_pt=9, color_hex=BODY)
    set_cell_margins(r2[2])

    set_col_widths(tbl, [1.5, 4.5, 3.3])
    set_table_borders(tbl)
    doc.add_paragraph()  # spacer between screen tables


def _fill_cell_lines(cell, text: str, size_pt=10, color_hex=BODY):
    """Write text into a cell, splitting on newlines into separate paragraphs."""
    lines = str(text).split("\n")
    first = True
    for line in lines:
        if first:
            para = cell.paragraphs[0]
            para.clear()
            first = False
        else:
            para = cell.add_paragraph()
        add_run(para, line.strip(), size_pt=size_pt, color_hex=color_hex)


# ---------------------------------------------------------------------------
# COVER SECTION
# ---------------------------------------------------------------------------
def build_cover(doc):
    if LOGO_PATH.exists():
        doc.add_picture(str(LOGO_PATH), width=Inches(2.0))
    else:
        p = doc.add_paragraph("[UHN LOGO]")
        p.runs[0].font.size = Pt(10)
        p.runs[0].font.color.rgb = rgb(COBALT)

    doc.add_paragraph()  # spacer

    # Series name
    series_para = doc.add_paragraph()
    add_run(series_para, "Accessibility First eLearning Series",
            bold=True, size_pt=12, color_hex=COBALT)

    # Title
    title_para = doc.add_paragraph()
    add_run(title_para,
            "Guide 02 — Perceptions, Attitudes, and Barriers",
            bold=True, size_pt=22, color_hex=NAVY, font_name="Arial Black")

    # Subtitle
    sub_para = doc.add_paragraph()
    add_run(sub_para,
            "Understanding How Attitudes and Assumptions Create Barriers in Healthcare",
            italic=True, size_pt=12, color_hex=NAVY)

    doc.add_paragraph()  # spacer


# ---------------------------------------------------------------------------
# METADATA TABLE
# ---------------------------------------------------------------------------
def build_metadata_table(doc):
    rows = [
        ("Audience",         "All UHN clinical & administrative staff"),
        ("Modality",         "Online · Self-paced · HTML5 SCORM 1.2 to UHN MyLearning (SumTotal)"),
        ("Duration",         "~15-20 minutes (23 slides)"),
        ("Standard",         "WCAG 2.1 AA · plain language · trauma-informed"),
        ("Author",           "Yi Jin (Instructional Designer)"),
        ("Content owner",    "Jacqueline Silvera — IDEAA, People & Culture, UHN"),
        ("Document version", "v1.0 Draft Storyboard — 2026-05-25"),
    ]
    add_two_col_table(doc, rows, col_widths=(2.0, 7.3))
    doc.add_paragraph()


# ---------------------------------------------------------------------------
# HOW TO READ
# ---------------------------------------------------------------------------
def build_how_to_read(doc):
    p = doc.add_paragraph()
    add_run(p,
            "How to read this storyboard  ",
            bold=True, size_pt=10, color_hex=NAVY, font_name="Arial Black")
    p.paragraph_format.space_after = Pt(2)

    body = doc.add_paragraph()
    add_run(body,
            "This document has two parts. "
            "Part 1 (Storyboard for Review) is for SME and stakeholder review — "
            "it contains the course information, learning outcomes, assessments, "
            "and a complete screen schedule. "
            "Part 2 (Designer Toolkit) is for the instructional designer and developer — "
            "it contains per-screen 3-column tables with Step, Activities, and Design Guide details "
            "that mirror the deployed SCORM course exactly. "
            "Review Part 1 for content accuracy and sign-off. "
            "Part 2 is a production reference only.",
            italic=True, size_pt=9, color_hex=BODY)
    body.paragraph_format.space_after = Pt(8)


# ---------------------------------------------------------------------------
# PART 1 — STORYBOARD FOR REVIEW
# ---------------------------------------------------------------------------

def build_1_1(doc):
    add_section_heading(doc, "1.1", "Course Information")
    rows = [
        ("Program Name",        "UHN Accessibility First eLearning Series"),
        ("Course Title",        "Perceptions, Attitudes, and Barriers"),
        ("Course Code",         "AF-02"),
        ("Course Hours",        "~0.30 hours / 15-20 minutes (online, self-paced)"),
        ("Pre-requisite(s)",    "Guide 01 — Foundations of Disability, Inclusion and Accessible Design"),
        ("Course Description",
         "This guide examines how personal perceptions and attitudes — including implicit and explicit bias — "
         "create barriers in healthcare settings. Learners explore the lived experience of attitudinal barriers "
         "across six categories, apply a reflective practice model (Notice-Name-Examine-Act), and practise "
         "dignity-centred responses through four branching scenarios. Framed around UHN's Accessibility First "
         "principles and the OHRC duty to accommodate. Builds on the foundation established in Guide 01."),
        ("Course Rationale",
         "Technical and environmental barriers often receive the most attention in accessibility training, "
         "but research consistently shows that attitudinal barriers are among the most persistent and hardest "
         "to address. This guide equips staff to recognize bias in themselves and their workplace culture, "
         "apply practical self-reflection habits, and move from compliance-focused thinking to genuine, "
         "proactive inclusion — with a direct connection to the OHRC and UHN IDEAA framework."),
        ("Learning Materials",
         "All readings, audio, and podcast embedded inside the HTML5 SCORM module. "
         "Source: Accessibility First Guide Series 1-18 (Draft, April 2026), Guide 2."),
        ("Accrediting Standards",
         "AODA — Integrated Accessibility Standards Regulation; "
         "Ontario Human Rights Code (Part I, Disability) — Duty to Accommodate; "
         "Mental Health Commission of Canada Anti-Stigma Resources; "
         "Statistics Canada Canadian Survey on Disability (2022); "
         "UHN IDEAA Strategic Framework; WCAG 2.1 AA."),
        ("Version / Last Update",
         "v1.0 Draft Storyboard — 2026-05-25 (Yi Jin, Instructional Designer)"),
    ]
    add_two_col_table(doc, rows)
    doc.add_paragraph()


def build_1_2(doc):
    add_section_heading(doc, "1.2", "Course Development Information")
    rows = [
        ("Delivery Modality",   "Online · Self-Paced · HTML5 SCORM 1.2 to UHN MyLearning (SumTotal LMS)"),
        ("Course Developer(s)", "Yi Jin (Instructional Designer)"),
        ("SME(s)",              "Jacqueline Silvera (IDEAA, People & Culture, UHN); CAMH language review — TBC; Indigenous content review — TBC"),
        ("ID(s)",               "Yi Jin"),
        ("Existing Course",     "None — new build."),
    ]
    add_two_col_table(doc, rows)
    doc.add_paragraph()


def build_1_3(doc):
    add_section_heading(doc, "1.3", "Course Learning Outcomes (CLOs)")
    rows = [
        ("CLO 1",
         "Define attitudinal barriers and explain how implicit and explicit bias affect "
         "accessibility and equity in healthcare."),
        ("CLO 2",
         "Recognize how assumptions about ability, culture, mental health, and identity "
         "influence patient care and employee interactions at UHN."),
        ("CLO 3",
         "Apply reflective practice strategies to identify and mitigate personal bias "
         "in everyday workplace interactions."),
        ("CLO 4",
         "Demonstrate proactive, dignity-centred responses that move accessibility practice "
         "beyond compliance."),
    ]
    add_two_col_table(doc, rows, header_row=["CLO", "Description"])
    doc.add_paragraph()


def build_1_4(doc):
    add_section_heading(doc, "1.4", "Summative Assessments")
    rows = [
        ("Knowledge Check 1",
         "MCQ — define implicit bias (Screen 2.14). 2 attempts.",
         "~11%", "Screen 2.14"),
        ("Knowledge Check 2",
         "Scenario MCQ — spread effect in practice (Screen 2.15). 2 attempts.",
         "~11%", "Screen 2.15"),
        ("Knowledge Check 3",
         "Scenario MCQ — OHRC paternalism/accommodation (Screen 2.16). 2 attempts.",
         "~11%", "Screen 2.16"),
        ("Branching Scenarios",
         "4 scenarios: autonomy, workplace bias, cultural communication, mental health handoff. "
         "Select-then-submit with 3-choice branching and debrief feedback.",
         "~47%", "Screens 2.10–2.13"),
        ("My Action Plan (MAP)",
         "Stop / Start / Continue commitments. Stop and Start fields required to proceed.",
         "~20%", "Screen 2.19"),
    ]
    add_four_col_table(
        doc, rows,
        col_widths=(2.0, 4.0, 0.8, 2.5),
        headers=["Title of Assessment", "Brief Description", "Weight", "Position"],
    )
    doc.add_paragraph()


def build_1_5(doc):
    add_section_heading(doc, "1.5", "Master Screen Schedule")
    schedule = [
        ("1",  "2.1",  "Welcome and Course Purpose",                 "Title / Cover",              "Hero image, Begin button"),
        ("2",  "2.2",  "Learning Objectives",                        "Objectives",                 "4 CLO cards (2x2 grid)"),
        ("3",  "2.3",  "Why This Matters: The Stat",                 "Content + Statistic",        "3x unmet healthcare needs, navy stat block"),
        ("4",  "2.4",  "Impact: Attitudinal Barriers in Healthcare", "Content + Image",            "Red accent bar, key takeaway"),
        ("5",  "2.5",  "Impact: Stigma and Mental Health",           "Content + Image",            "Cobalt accent bar, 11-year stat"),
        ("6",  "2.6",  "Impact: Intersectionality",                  "Content + Image",            "Navy accent bar, Indigenous context box"),
        ("7",  "2.7",  "Core Concepts: Implicit vs. Explicit Bias",  "Tabbed panels (2 tabs)",     "Cobalt (implicit) / Red (explicit)"),
        ("8",  "2.8",  "Attitudinal Barriers Framework",             "Accordion (6 items)",        "Click-to-reveal, icons per category"),
        ("9",  "2.9",  "Reflective Practice Model",                  "Step-through (4 steps)",     "Notice, Name, Examine, Act — circular diagram"),
        ("10", "2.10", "Scenario 1: Assumptions About Ability",      "Branching (3-choice)",       "Mr. Santos / wheelchair scenario, graded"),
        ("11", "2.11", "Scenario 2: Implicit Bias in the Workplace", "Branching (3-choice)",       "Priya / leadership scenario, OHRC reference"),
        ("12", "2.12", "Scenario 3: Cultural Misunderstanding",      "Branching (3-choice)",       "Mrs. Nguyen / bedside assessment scenario"),
        ("13", "2.13", "Scenario 4: Mental Health in Handoff",       "Branching (3-choice)",       "Mr. Thompson / shift handoff scenario"),
        ("14", "2.14", "Knowledge Check 1",                          "MCQ (4 options)",            "Implicit bias definition, CLO 1, 2 attempts"),
        ("15", "2.15", "Knowledge Check 2",                          "Scenario MCQ (4 options)",   "Spread effect, CLO 2, 2 attempts"),
        ("16", "2.16", "Knowledge Check 3",                          "Scenario MCQ (4 options)",   "OHRC paternalism, CLO 4, 2 attempts"),
        ("17", "2.17", "Inclusive Practice Tips",                    "Card carousel (5 cards)",    "Ask, Language, Address patient, Handoff, NNEA"),
        ("18", "2.18", "Reflection Prompt",                          "Text entry (ungraded)",      "Private reflection — not submitted"),
        ("19", "2.19", "MAP Action Planning",                        "3-field form + PDF download","Stop/Start/Continue, MAP-Template-Guide-02.pdf"),
        ("20", "2.20", "Key Takeaways",                              "Summary (4 points)",         "4 takeaway cards, navy/cobalt"),
        ("21", "2.21", "Listen and Reflect (Podcast)",               "Podcast player + transcript","Episode 02 · The Invisible Barrier · ~15-18 min"),
        ("22", "2.22", "Series Progress Map",                        "Progress map (3 stages)",    "2 of 4 Foundations complete; Stage 2+3 locked"),
        ("23", "2.23", "Resources and Course Completion",            "Completion + badge",         "Badge, resource links, Guide 03 next"),
    ]
    add_five_col_table(
        doc, schedule,
        col_widths=[0.45, 0.65, 2.8, 2.1, 3.3],
        headers=["Slide", "Screen", "Title", "Type", "Notes"],
    )
    doc.add_paragraph()


def build_1_6(doc):
    add_section_heading(doc, "1.6", "Version Control & Sign-off")

    version_rows = [
        ("Version", "Date",       "Author",   "Change Summary"),
        ("v0.1",    "2026-05-25", "Yi Jin",   "Initial draft storyboard"),
        ("v0.5",    "TBD",        "Yi Jin",   "SME review draft"),
        ("v0.5.1",  "TBD",        "Yi Jin / J. Silvera", "Post-SME revisions"),
        ("v1.0",    "TBD",        "Yi Jin",   "Deployed storyboard — final"),
    ]
    tbl = doc.add_table(rows=len(version_rows), cols=4)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, row in enumerate(version_rows):
        cells = tbl.rows[i].cells
        for j, val in enumerate(row):
            is_header = (i == 0)
            shade_cell(cells[j], NAVY if is_header else (HEADER_FILL if j == 0 else WHITE))
            cell_para(cells[j], val, bold=is_header, size_pt=10,
                      color_hex=WHITE if is_header else (NAVY if j == 0 else BODY),
                      font_name="Arial Black" if is_header else "Arial")
            set_cell_margins(cells[j])
    set_col_widths(tbl, [0.8, 1.2, 2.5, 4.8])
    set_table_borders(tbl)

    doc.add_paragraph()

    # Sign-off box
    signoff_tbl = doc.add_table(rows=1, cols=1)
    cell = signoff_tbl.rows[0].cells[0]
    shade_cell(cell, LIGHT_FILL)
    set_cell_margins(cell, top_pt=10, bottom_pt=10, left_pt=12, right_pt=12)

    lines = [
        {"text": "SIGN-OFF", "bold": True, "size_pt": 11, "color_hex": NAVY,
         "font_name": "Arial Black"},
        {"text": ""},
        {"text": "SME Review — Jacqueline Silvera (IDEAA, People & Culture, UHN)"},
        {"text": "Reviewed by: ________________________    Date: ____________"},
        {"text": ""},
        {"text": "CAMH Language Review — [Reviewer TBC]"},
        {"text": "Reviewed by: ________________________    Date: ____________"},
        {"text": ""},
        {"text": "Indigenous Content Review — [Reviewer TBC]"},
        {"text": "Reviewed by: ________________________    Date: ____________"},
        {"text": ""},
        {"text": "ID Review — Yi Jin (Instructional Designer)"},
        {"text": "Approved by: ________________________    Date: ____________"},
        {"text": ""},
        {"text": "Note: Signature on this document constitutes approval of all content, "
                 "scenarios, assessment questions, and learning outcomes as accurate, "
                 "appropriate for the intended audience, and aligned with UHN policy.",
         "italic": True, "size_pt": 9, "color_hex": BODY},
    ]
    cell_para_multiline(cell, lines)
    set_table_borders(signoff_tbl, color_hex=COBALT)
    doc.add_paragraph()


# ---------------------------------------------------------------------------
# PART 2 — DESIGNER TOOLKIT: Per-screen data
# ---------------------------------------------------------------------------

SCREEN_DATA = [
    {
        "num":   "2.1",
        "title": "Welcome and Course Purpose",
        "step":  "Hook",
        "activities": (
            "TITLE: Perceptions, Attitudes, and Barriers\n"
            "SUBTITLE: Guide 02 of 18 · Foundations Stage\n"
            "BODY: Over the next 15 to 20 minutes, you will examine how our own perceptions and attitudes shape the care and environments we create.\n"
            "METADATA: Duration 15-20 min · Pre-requisite: Guide 01 · Version v1.0 · 2026\n"
            "BUTTON: Begin\n"
            "VO: voiceover_2.1.mp3\n"
            "NARRATION (~148 words): Welcome back to the Accessibility First series. This is Guide 2: Perceptions, Attitudes, and Barriers. In Guide 1, you explored what disability means, the models that shape how we think about it, and a practical framework for everyday decisions. This guide goes deeper. It asks a more personal question: how do our own perceptions and attitudes shape the care and environments we create? We all carry assumptions. Some we are aware of. Many we are not. And in a healthcare setting, those assumptions — even well-intentioned ones — can create barriers as real as any locked door or missing ramp. Over the next 15 to 20 minutes, you will examine those assumptions, explore where they come from, and build practical strategies to move from compliance toward genuine inclusion. This guide is relevant for everyone at UHN. Let us get started."
        ),
        "design_guide": (
            "Format: Cover slide · static layout with autoplay narration.\n"
            "Interaction: Begin button navigates to Screen 2.2.\n"
            "Accessibility: 1.4.3 contrast >= 4.5:1 · 2.1.1 keyboard reachable · CC ON by default · Begin button Tab + Enter.\n"
            "Visual: Warm, inclusive hero photo — diverse UHN staff and patients in a bright hospital corridor. Person using wheelchair, staff member with hijab, older patient with hearing aid.\n"
            "Image: g02-hero-welcome-01.png\n"
            "Alt: Diverse UHN healthcare workers and patients in conversation in a bright hospital corridor\n"
            "Captions: voiceover_2.1.srt"
        ),
    },
    {
        "num":   "2.2",
        "title": "Learning Objectives",
        "step":  "Outcomes",
        "activities": (
            "HEADER: What You Will Learn\n"
            "CLO 1: Understand attitudinal barriers — define attitudinal barriers and explain how implicit and explicit bias affect accessibility and equity in healthcare.\n"
            "CLO 2: Recognize assumptions in practice — identify how assumptions about ability, culture, mental health, and identity influence patient care and employee interactions at UHN.\n"
            "CLO 3: Apply reflective practice — use self-reflection strategies to recognize and mitigate personal bias in everyday interactions.\n"
            "CLO 4: Act with dignity beyond compliance — demonstrate proactive, respectful responses that move accessibility practice beyond meeting minimum standards.\n"
            "VO: voiceover_2.2.mp3\n"
            "NARRATION (~122 words): By the end of this guide, you will be able to do four things. First, you will be able to define attitudinal barriers and explain the difference between implicit and explicit bias — and why both matter in healthcare. Second, you will be able to recognize how assumptions about ability, culture, mental health, and identity show up in patient care and workplace interactions. Third, you will be able to apply reflective practice strategies to identify your own assumptions and adjust your behaviour. And fourth, you will be able to demonstrate responses that go beyond compliance — responses grounded in dignity, curiosity, and genuine inclusion. These four objectives build on Guide 1. Each one connects directly to a scenario you will encounter later in this guide."
        ),
        "design_guide": (
            "Format: Static objective cards (4 cards, 2x2 grid, navy top border, large number + title + description).\n"
            "Interaction: None — orientation slide. Continue always available.\n"
            "Accessibility: Ordered list markup (ol/li). Each icon has alt text. WCAG 2.1 AA contrast.\n"
            "Image: g02-icons-objectives-01.png (flat vector icon set — mirror, magnifying glass, light bulb, handshake)\n"
            "Captions: voiceover_2.2.srt"
        ),
    },
    {
        "num":   "2.3",
        "title": "Why This Matters: The Stat",
        "step":  "Context",
        "activities": (
            "HEADER: The Invisible Barrier\n"
            "STAT: People with disabilities are 3x more likely to report unmet healthcare needs — Statistics Canada, 2022\n"
            "PULL QUOTE: Often, the barrier is not a ramp or a doorway. It is an attitude.\n"
            "BODY: Attitudinal barriers are invisible to those who hold them — and deeply felt by those who experience them.\n"
            "VO: voiceover_2.3.mp3\n"
            "NARRATION (~90 words): Here is a number worth sitting with. According to Statistics Canada, people with disabilities are nearly three times more likely than people without disabilities to report unmet healthcare needs. Not because services do not exist. But because the experience of seeking care — the assumptions made, the language used, the way people are treated — creates barriers that make people give up before they even get the help they need. The most persistent barriers in healthcare are not physical. They are attitudinal. And the good news is that attitudes can change. That is what this guide is about.\n"
            "REFS: Statistics Canada (2023). Canadian Survey on Disability, 2022.\n"
            "SME NOTE: HIGH — verify stat accuracy with SME"
        ),
        "design_guide": (
            "Format: Large statistic prominent, centred. Animated counter. Clean white background with navy stat block. Supporting text below stat block.\n"
            "Accessibility: Stat coded as text, not image. Animation respects prefers-reduced-motion.\n"
            "Image: g02-stats-attitudinal-01.png\n"
            "Alt: Statistic with supporting text: people with disabilities are 3x more likely to report unmet healthcare needs\n"
            "Captions: voiceover_2.3.srt"
        ),
    },
    {
        "num":   "2.4",
        "title": "Impact: Attitudinal Barriers in Healthcare",
        "step":  "Impact",
        "activities": (
            "HEADER: When Attitudes Become Barriers\n"
            "BODY: Attitudinal barriers occur when assumptions, stereotypes, or lack of awareness lead to exclusion — even without intent.\n"
            "BODY 2: In healthcare, these barriers affect diagnosis, treatment decisions, and whether patients feel safe returning.\n"
            "VO: voiceover_2.4.mp3\n"
            "NARRATION (~85 words): Attitudinal barriers are among the most common — and the hardest to see — in any workplace. In healthcare, they can affect who receives equitable treatment, who gets taken seriously, and who feels safe enough to come back. An employee who speaks over a patient with a cognitive disability. A clinician who assumes a patient with a mental health history cannot make sound decisions. A colleague who speaks to a patient's support person instead of the patient directly. None of these require bad intent. They require awareness. And awareness is something we can build.\n"
            "SME NOTE: Medium"
        ),
        "design_guide": (
            "Format: Two-column layout. Left: photo placeholder (540px height). Right: red accent bar + title + body text + key takeaway box.\n"
            "Accessibility: Red accent bar for visual emphasis. WCAG 2.1 AA contrast.\n"
            "Image: g02-impact-attitudinal-01.png\n"
            "Alt: Healthcare provider and patient in a tense or disconnected interaction suggesting a communication gap\n"
            "Captions: voiceover_2.4.srt"
        ),
    },
    {
        "num":   "2.5",
        "title": "Impact: Stigma and Mental Health",
        "step":  "Impact",
        "activities": (
            "HEADER: The Weight of Stigma\n"
            "STAT: Mental health stigma in healthcare settings delays help-seeking by an average of 11 years.\n"
            "BODY: For employees with disabilities, fear of disclosure can limit career growth and access to supports.\n"
            "VO: voiceover_2.5.mp3\n"
            "NARRATION (~82 words): Stigma is one of the most powerful attitudinal barriers we know of. The Mental Health Commission of Canada reports that fear of stigma causes many people to delay seeking help for mental health conditions for over a decade. At UHN, this affects both patients and employees. Patients may not disclose conditions that are relevant to their care. Employees may not ask for the accommodations they need. When we examine our own assumptions about mental health — our language, our reactions, our expectations — we begin to break down stigma from the inside out.\n"
            "SME NOTE: HIGH — CAMH language review required"
        ),
        "design_guide": (
            "Format: Two-column layout. Left: photo placeholder (540px height). Right: cobalt accent bar + title + body text + key takeaway box.\n"
            "Accessibility: Cobalt accent bar for visual emphasis. WCAG 2.1 AA contrast.\n"
            "Image: g02-impact-stigma-01.png\n"
            "Alt: Person sitting alone in a waiting area, looking anxious or withdrawn\n"
            "Captions: voiceover_2.5.srt"
        ),
    },
    {
        "num":   "2.6",
        "title": "Impact: Intersectionality and Compounding Barriers",
        "step":  "Impact",
        "activities": (
            "HEADER: Compounding Barriers\n"
            "BODY: A person's experience of disability is shaped by race, gender, language, income, age, and culture — not disability alone.\n"
            "INDIGENOUS CONTEXT: Indigenous Peoples face compounded barriers rooted in colonial history, systemic racism, and cultural disconnection from care.\n"
            "VO: voiceover_2.6.mp3\n"
            "NARRATION (~86 words): Accessibility cannot be understood through a single lens. A First Nations elder navigating the healthcare system brings a different set of experiences than a young Black professional with an invisible disability, or a new immigrant with a cognitive disability who speaks neither English nor French. Intersectionality means that disability intersects with race, gender, language, income, and culture to shape a person's experience in ways that are unique to them. At UHN, this means our responses must be flexible, humble, and informed by the person in front of us — not a category or a checklist.\n"
            "SME NOTE: HIGH — Indigenous content requires dedicated SME sign-off"
        ),
        "design_guide": (
            "Format: Two-column layout. Left: photo placeholder (540px height). Right: navy accent bar + title + body text + Indigenous context box (earth tones).\n"
            "Accessibility: Navy accent bar + Indigenous context box in earth tones. WCAG 2.1 AA contrast.\n"
            "Image: g02-impact-intersectionality-01.png\n"
            "Alt: Diverse group of people representing multiple identities in a healthcare setting\n"
            "Captions: voiceover_2.6.srt"
        ),
    },
    {
        "num":   "2.7",
        "title": "Core Concepts: Implicit vs. Explicit Bias",
        "step":  "Core Concept",
        "activities": (
            "HEADER: Two Types of Bias\n"
            "TAB 1 — IMPLICIT BIAS: An unconscious attitude or stereotype that affects judgement and behaviour without awareness.\n"
            "  Example: A nurse automatically speaks louder when addressing a patient using a wheelchair, assuming hearing loss.\n"
            "  Impact: Patients feel patronized; trust erodes.\n"
            "TAB 2 — EXPLICIT BIAS: A conscious attitude, belief, or stereotype that is deliberately held and acted upon.\n"
            "  Example: A provider believes that patients with mental health histories are 'difficult' and delays their intake.\n"
            "  Impact: Direct discrimination; creates unsafe care environment.\n"
            "KEY INSIGHT: Both types cause harm. The difference is awareness — and awareness gives us a place to start.\n"
            "VO: voiceover_2.7.mp3\n"
            "NARRATION (~198 words): Let us be clear about something: bias is not a character flaw. It is a feature of how the human brain processes information. We all carry biases — they are shaped by our upbringing, our experiences, media, and culture. The question is not whether you have bias. The question is whether you can recognize it and choose a different response. There are two types of bias worth knowing. Implicit bias is unconscious. You may not even be aware it is shaping your behaviour. A nurse who automatically slows their speech when talking to a patient using a mobility device — not because of hearing loss, but because of an automatic assumption about disability. That nurse likely means well. The harm is real regardless. Explicit bias is conscious. It is a deliberate belief or stereotype that the person knows they hold. This is rarer in formal healthcare settings, but it exists — and it creates direct discrimination. Both types matter. Both cause harm. The difference is that implicit bias is harder to see and easier to deny. That is why reflective practice — which we will explore in the next section — is so important.\n"
            "SME NOTE: Medium"
        ),
        "design_guide": (
            "Format: Two-panel tabbed comparison layout. Implicit tab: cobalt accent. Explicit tab: red accent.\n"
            "Interaction: Click tab to reveal content. One panel visible at a time.\n"
            "Accessibility: Tab panels keyboard-navigable. Each tab announced by screen reader. Active tab visually distinct. role=tablist, aria-selected.\n"
            "Image: g02-infographic-bias-types-01.png (flat vector)\n"
            "Alt: Two-panel comparison: implicit bias (unconscious) vs explicit bias (conscious) in healthcare\n"
            "Captions: voiceover_2.7.srt"
        ),
    },
    {
        "num":   "2.8",
        "title": "Attitudinal Barriers Framework",
        "step":  "Framework",
        "activities": (
            "HEADER: Six Types of Attitudinal Barriers\n"
            "ACCORDION 1 — PATERNALISM: Assuming you know better than the person what they need.\n"
            "ACCORDION 2 — STEREOTYPING: Applying group-based assumptions to individuals.\n"
            "ACCORDION 3 — PITY OR CHARITY MODEL: Viewing people with disabilities as objects of sympathy rather than rights-holders.\n"
            "ACCORDION 4 — SPREAD EFFECT: Assuming one disability affects all aspects of a person's functioning.\n"
            "ACCORDION 5 — OTHERING: Treating people with disabilities as fundamentally different rather than as peers and patients.\n"
            "ACCORDION 6 — INVISIBILIZATION: Failing to acknowledge non-visible disabilities or dismissing disclosures.\n"
            "VO: voiceover_2.8.mp3\n"
            "NARRATION (~193 words): Attitudinal barriers are not monolithic — they come in different forms, and recognizing the specific form helps you respond more effectively. The first is paternalism: deciding for someone what is best for them without asking. It often comes from a desire to help but removes the person's autonomy. The second is stereotyping: applying a generalization about a group to an individual. Not every person who is Deaf communicates the same way. Not every person using a wheelchair has the same functional needs. The third is the pity or charity model: seeing disability as tragedy rather than recognizing people as rights-holders with agency. The fourth is the spread effect: assuming that because someone has one disability, they must have others too. A person with a visual impairment does not necessarily have cognitive or communication differences. The fifth is othering: treating people as fundamentally different, as objects of curiosity or concern rather than as people. And the sixth is invisibilization: dismissing or doubting disclosures of non-visible disabilities. 'You don't look sick' is a form of attitudinal barrier. Each of these can be unlearned. And each one requires us to pause and ask.\n"
            "SME NOTE: Medium"
        ),
        "design_guide": (
            "Format: Accordion interaction — six items, click each to expand. Clean list with icon per category.\n"
            "Interaction: Click-to-reveal accordion. Each item readable without expansion.\n"
            "Accessibility: Accordion keyboard-navigable. Each item readable without expansion. Focus indicator visible.\n"
            "Image: g02-infographic-barriers-accordion-01.png (flat vector icons)\n"
            "Alt: Six accordion panels listing attitudinal barrier types in healthcare\n"
            "Captions: voiceover_2.8.srt"
        ),
    },
    {
        "num":   "2.9",
        "title": "Reflective Practice Model",
        "step":  "Framework",
        "activities": (
            "HEADER: The Notice-Name-Examine-Act Cycle\n"
            "STEP 1 — NOTICE: Pause and become aware of your reaction, assumption, or discomfort.\n"
            "STEP 2 — NAME: Label what you observed without judgment. ('I assumed...')\n"
            "STEP 3 — EXAMINE: Consider where this assumption comes from and what impact it might have.\n"
            "STEP 4 — ACT: Choose a different response — one grounded in curiosity, dignity, and the person's actual expressed needs.\n"
            "VO: voiceover_2.9.mp3\n"
            "NARRATION (~186 words): Reflective practice is a disciplined habit of pausing to examine your own assumptions and reactions. It sounds simple. It is not always easy — especially in a busy clinical environment. But it is one of the most powerful tools we have for improving accessibility and equity. Here is a four-step cycle you can use in any interaction. Step one: Notice. Catch yourself in the moment. Did you react differently to this patient than to the last? Did something in their chart, their appearance, or their behaviour trigger an assumption? Step two: Name it. Not as an accusation against yourself, but as an observation. 'I assumed this patient could not make their own decision.' Naming it makes it visible. Step three: Examine. Where does this assumption come from? What impact could it have on this patient's care? And step four: Act. Choose a different response — one based on what the patient in front of you has actually said or expressed, not on what you expected. This cycle does not require perfection. It requires honesty and the willingness to grow.\n"
            "SME NOTE: Low"
        ),
        "design_guide": (
            "Format: Circular diagram (4 stages). Each stage: icon + label + brief description. Chartreuse accent on active step.\n"
            "Interaction: Step-through (click each step to reveal detail). Keyboard-navigable.\n"
            "Accessibility: Step interaction keyboard-navigable. Focus indicator on each step.\n"
            "Image: g02-infographic-reflection-cycle-01.png (flat vector)\n"
            "Alt: Four-step reflective practice cycle: Notice, Name, Examine, Act\n"
            "Captions: voiceover_2.9.srt"
        ),
    },
    {
        "num":   "2.10",
        "title": "Scenario 1: Assumptions About Ability",
        "step":  "Scenario",
        "activities": (
            "SETUP: You are working at the reception desk. Mr. Santos, a patient using a manual wheelchair, has been waiting for 10 minutes. There are no staff at the desk and no sign indicating when someone will return. You walk past on your way to another task. Mr. Santos has his intake form on the armrest of his chair. He has not yet started it.\n"
            "QUESTION: What do you do?\n"
            "CHOICE A (Best): You introduce yourself, ask if Mr. Santos would like any help with the form, and let him know someone will be at the desk shortly. You wait until a colleague arrives or handle the intake yourself.\n"
            "OUTCOME A: Mr. Santos completes his own form with assistance where he chooses to use it. His autonomy is respected. He feels seen, not managed.\n"
            "CHOICE B (Acceptable): You apologize for the wait, take his form, and start filling it in for him, asking him the questions as you go.\n"
            "OUTCOME B: The task gets done, but Mr. Santos' autonomy is reduced. He did not ask for the form to be taken. This approach, though well-intentioned, mirrors the original attitudinal barrier.\n"
            "CHOICE C (Poor): You pick up the form and set it on the desk, telling him someone will be with him soon.\n"
            "OUTCOME C: Mr. Santos feels dismissed and patronized. His time and independence are both disregarded. The interaction reinforces that the system was not designed with him in mind.\n"
            "DEBRIEF: Autonomy means offering help, not assuming it is needed or wanted. Always ask. 'Can I help you with anything?' is the opening. Let the person in front of you direct what happens next.\n"
            "PRINCIPLE: People First & Dignity · Independence & Autonomy\n"
            "VO: voiceover_2.10.mp3 · SME NOTE: HIGH"
        ),
        "design_guide": (
            "Format: Branching choice · 3 options · graded. Select-then-submit with feedback overlay.\n"
            "Interaction: All choices keyboard-selectable. Consequence screens have back navigation. No auto-advance.\n"
            "Image: g02-scenario-ability-01.png\n"
            "Alt: Patient using wheelchair waiting at an unattended hospital reception desk with an intake form on their armrest\n"
            "Captions: voiceover_2.10.srt\n"
            "Maps to: CLO 2, CLO 4"
        ),
    },
    {
        "num":   "2.11",
        "title": "Scenario 2: Implicit Bias in the Workplace",
        "step":  "Scenario",
        "activities": (
            "SETUP: You are a team leader responsible for assigning a lead role on a new quality improvement project. Two employees have applied: both have strong performance reviews. One of them, Priya, disclosed last year that she has a mood disorder and took a medical leave. Her performance since returning has been consistently excellent.\n"
            "QUESTION: What do you do?\n"
            "CHOICE A (Best): You review both candidates' recent performance and assign based on merit. You do not factor in Priya's medical history. If the role requires additional support, you discuss that directly with whoever is selected.\n"
            "OUTCOME A: Priya is assessed equitably. If selected, she leads successfully. The team learns that a disability history does not predict future performance.\n"
            "CHOICE B (Acceptable): You assign Priya to a supporting role instead of the lead, intending to protect her from added stress without telling her why.\n"
            "OUTCOME B: The decision is made 'for' her, without her input. Priya may sense the reason and feel her disclosure harmed her career. This is a form of paternalism and may constitute discrimination under the OHRC.\n"
            "CHOICE C (Poor): You assign the other candidate because you worry the workload might be too much for someone who has taken medical leave.\n"
            "OUTCOME C: This is a direct example of implicit bias becoming a discriminatory action. The OHRC requires that accommodation decisions not be made on the basis of assumed limitation without the employee's input.\n"
            "DEBRIEF: The duty to accommodate does not mean making decisions for people. It means removing barriers and having conversations. Assumptions about what someone can handle — without asking — are both a bias and a legal risk.\n"
            "PRINCIPLE: Equity, Rights & Intersectionality · People First & Dignity\n"
            "VO: voiceover_2.11.mp3 · SME NOTE: HIGH — OHRC/CAMH language review"
        ),
        "design_guide": (
            "Format: Branching choice · 3 options · graded. Same interaction pattern as Screen 2.10.\n"
            "Image: g02-scenario-bias-workplace-01.png\n"
            "Alt: Two healthcare professionals in a meeting room, one reviewing documents for a project assignment\n"
            "Captions: voiceover_2.11.srt\n"
            "Maps to: CLO 1, CLO 3, CLO 4"
        ),
    },
    {
        "num":   "2.12",
        "title": "Scenario 3: Cultural Misunderstanding in Care",
        "step":  "Scenario",
        "activities": (
            "SETUP: You are a nurse conducting a pre-procedure assessment for Mrs. Nguyen, an elderly Vietnamese patient. She avoids making direct eye contact and gives short, indirect answers to your questions. Her daughter is present and responds more fully when you ask questions.\n"
            "QUESTION: What do you do?\n"
            "CHOICE A (Best): You continue to address your questions to Mrs. Nguyen directly, using plain language and pausing frequently. You acknowledge the daughter's presence and invite her input after Mrs. Nguyen has had the chance to respond.\n"
            "OUTCOME A: Mrs. Nguyen feels respected. Her daughter's role as a support person is affirmed without displacing the patient's centrality. The assessment is accurate.\n"
            "CHOICE B (Acceptable): You begin directing most questions to the daughter, reasoning that communication is more efficient.\n"
            "OUTCOME B: The assessment moves faster but Mrs. Nguyen's voice is sidelined. This is a common pattern that can lead to missed information and erodes patient dignity. It may also reflect assumptions about age and cognition.\n"
            "CHOICE C (Poor): You document that the patient is 'uncooperative' or 'unable to provide history' and proceed without exploring communication alternatives.\n"
            "OUTCOME C: This is a clinical error with potential safety consequences. Mrs. Nguyen's behaviour is a cultural communication style, not a clinical finding. The documentation is inaccurate and potentially harmful.\n"
            "DEBRIEF: Cultural communication differences are not compliance problems. Indirect eye contact, deference to family, and non-linear communication patterns are meaningful and valid. The question to ask is: 'How can I create the conditions for this person to be heard?'\n"
            "PRINCIPLE: Inclusion & Integration · People First & Dignity\n"
            "VO: voiceover_2.12.mp3 · SME NOTE: HIGH — cultural review required"
        ),
        "design_guide": (
            "Format: Branching choice · 3 options · graded. Same interaction pattern as Screen 2.10.\n"
            "Image: g02-scenario-cultural-01.png\n"
            "Alt: Nurse conducting bedside assessment with elderly patient and family member present\n"
            "Captions: voiceover_2.12.srt\n"
            "Maps to: CLO 2, CLO 3, CLO 4"
        ),
    },
    {
        "num":   "2.13",
        "title": "Scenario 4: Language and Mental Health in Handoff",
        "step":  "Scenario",
        "activities": (
            "SETUP: During a shift handoff, an outgoing nurse mentions that the patient in room 4B, Mr. Thompson, 'can be difficult to work with' — and points to a note in the chart about a history of bipolar disorder. The incoming nurse, Amir, has not yet met Mr. Thompson.\n"
            "QUESTION: What do you do as Amir?\n"
            "CHOICE A (Best): You thank the outgoing nurse for the handoff and note the diagnosis, but you go in to meet Mr. Thompson without assumptions. You introduce yourself, ask how he is doing, and let the interaction inform your approach.\n"
            "OUTCOME A: Mr. Thompson is treated as an individual. Amir forms his own assessment. If challenges arise, Amir is better equipped to respond because he is responding to what is actually happening — not to a label.\n"
            "CHOICE B (Acceptable): You flag to the charge nurse that you are concerned about the characterization in the handoff note and ask for guidance before entering the room.\n"
            "OUTCOME B: This is a reasonable response, though it delays care. It shows awareness of bias and a desire to respond appropriately. The charge nurse can provide context and model good practice.\n"
            "CHOICE C (Poor): You go into the room already guarded, with less patience and a shorter interaction than you would give other patients.\n"
            "OUTCOME C: Mr. Thompson senses the difference. Any 'difficult' behaviour he exhibits may now be a response to being treated differently. The bias has become self-fulfilling.\n"
            "DEBRIEF: Diagnostic labels do not predict behaviour. Handoff language like 'difficult' transfers bias, not clinical information. Each patient deserves a fresh interaction — informed by clinical facts, not by another person's attitudinal lens.\n"
            "PRINCIPLE: People First & Dignity · Proactive Barrier Prevention\n"
            "VO: voiceover_2.13.mp3 · SME NOTE: HIGH — CAMH language review"
        ),
        "design_guide": (
            "Format: Branching choice · 3 options · graded. Same interaction pattern as Screen 2.10.\n"
            "Image: g02-scenario-mentalhealth-01.png\n"
            "Alt: Two nurses conducting a shift handoff at a nursing station, reviewing a patient chart\n"
            "Captions: voiceover_2.13.srt\n"
            "Maps to: CLO 1, CLO 2, CLO 3"
        ),
    },
    {
        "num":   "2.14",
        "title": "Knowledge Check 1",
        "step":  "Assessment",
        "activities": (
            "HEADER: Knowledge Check 1\n"
            "QUESTION: Which of the following BEST describes implicit bias?\n"
            "A: A deliberate belief that a group of people is inferior\n"
            "B: An unconscious attitude that shapes behaviour without awareness [CORRECT]\n"
            "C: A policy that intentionally excludes people with disabilities\n"
            "D: A formal complaint about discriminatory treatment\n"
            "FEEDBACK CORRECT (B): That is right. Implicit bias is unconscious — we may not even be aware it is influencing our actions. That is what makes reflective practice so important.\n"
            "FEEDBACK INCORRECT (A): This describes explicit bias — a deliberate, conscious attitude. Implicit bias is different: it operates below the level of awareness.\n"
            "FEEDBACK INCORRECT (C): This describes systemic or structural discrimination in policy. Implicit bias is an individual-level unconscious attitude.\n"
            "FEEDBACK INCORRECT (D): A formal complaint is a response to discrimination, not a type of bias. Review Screen 2.7.\n"
            "2 attempts allowed. After 2 attempts, correct answer revealed. Submit button locks after final attempt.\n"
            "VO: voiceover_2.14.mp3 (brief intro only)"
        ),
        "design_guide": (
            "Format: MCQ · 4 options · 2 attempts · submit-then-lock. Clean quiz layout. Navy/white colour scheme.\n"
            "Interaction: Select option -> Submit Answer button enables -> feedback panel appears. After 2 wrong attempts, correct answer revealed.\n"
            "Captions: voiceover_2.14.srt\n"
            "Maps to: CLO 1"
        ),
    },
    {
        "num":   "2.15",
        "title": "Knowledge Check 2",
        "step":  "Assessment",
        "activities": (
            "HEADER: Knowledge Check 2\n"
            "QUESTION: A nurse automatically speaks more slowly and loudly to a patient using a power wheelchair, even though there is no documented hearing or cognitive concern in the chart. This is an example of:\n"
            "A: Explicit bias — the nurse is deliberately patronizing the patient\n"
            "B: Spread effect — assuming one disability affects other areas of functioning [CORRECT]\n"
            "C: Invisibilization — dismissing the patient's non-visible disability\n"
            "D: Paternalism — deciding what the patient needs without asking\n"
            "FEEDBACK CORRECT (B): Exactly. The spread effect means assuming that one disability automatically affects other areas of functioning. Using a wheelchair does not affect hearing or cognitive ability.\n"
            "FEEDBACK INCORRECT (A): The nurse is likely unaware of the assumption — this is the hallmark of implicit, not explicit, bias. The specific mechanism here is the spread effect.\n"
            "FEEDBACK INCORRECT (C): Invisibilization involves dismissing disclosed non-visible disabilities. This scenario involves assuming visible disability implies other limitations.\n"
            "FEEDBACK INCORRECT (D): Paternalism involves making decisions for someone. This scenario is about an assumption about ability, not a care decision.\n"
            "2 attempts allowed.\n"
            "VO: voiceover_2.15.mp3"
        ),
        "design_guide": (
            "Format: Scenario-based MCQ · 4 options · 2 attempts · submit-then-lock.\n"
            "Captions: voiceover_2.15.srt\n"
            "Maps to: CLO 2"
        ),
    },
    {
        "num":   "2.16",
        "title": "Knowledge Check 3",
        "step":  "Assessment",
        "activities": (
            "HEADER: Knowledge Check 3\n"
            "QUESTION: Under the Ontario Human Rights Code, a manager who withholds a leadership opportunity from an employee because of a disclosed disability — without the employee's input — is most likely:\n"
            "A: Acting appropriately by protecting the employee from stress\n"
            "B: Fulfilling the duty to accommodate\n"
            "C: Committing a form of paternalism that may constitute discrimination [CORRECT]\n"
            "D: Complying with AODA requirements\n"
            "FEEDBACK CORRECT (C): That is right. Making decisions 'for' someone based on a disability assumption — without their input — is paternalistic and may constitute discrimination under the OHRC. Accommodation requires a conversation.\n"
            "FEEDBACK INCORRECT (A): Protecting someone from stress by limiting their opportunities, without their input, is not protection — it is a form of discrimination.\n"
            "FEEDBACK INCORRECT (B): The duty to accommodate requires consultation with the employee. Unilaterally making decisions on their behalf does not fulfill this duty.\n"
            "FEEDBACK INCORRECT (D): The AODA is primarily about removing barriers in the provision of services and employment practices, not about shielding employees from opportunities.\n"
            "2 attempts allowed.\n"
            "VO: voiceover_2.16.mp3 · SME NOTE: HIGH — legal accuracy review required"
        ),
        "design_guide": (
            "Format: Scenario-based MCQ · 4 options · 2 attempts · submit-then-lock.\n"
            "Captions: voiceover_2.16.srt\n"
            "Maps to: CLO 4"
        ),
    },
    {
        "num":   "2.17",
        "title": "Inclusive Practice Tips",
        "step":  "Practice",
        "activities": (
            "HEADER: Five Inclusive Practice Tips\n"
            "TIP 1: Ask, do not assume. Before offering help, ask: 'Is there anything I can do to support you?'\n"
            "TIP 2: Watch your language. Replace 'suffers from' with 'has.' Avoid terms like 'wheelchair-bound' or 'mentally ill.'\n"
            "TIP 3: Address the patient, not their companion. Direct your questions and eye contact to the patient in front of you.\n"
            "TIP 4: Examine handoff language. When you hear 'difficult patient,' ask: 'Difficult how? What has been tried?'\n"
            "TIP 5: Take the Notice-Name-Examine-Act cycle with you. One pause can change an interaction entirely.\n"
            "VO: voiceover_2.17.mp3\n"
            "NARRATION (~220 words): Before we move to the reflection activity, here are five inclusive practice tips you can use starting today. The first is simple but powerful: ask, do not assume. Before you assist, before you adapt, before you redirect — ask the person in front of you what they need. 'Is there anything I can do to support you?' is one of the most respectful questions in healthcare. The second is about language. Language shapes perception. When we say someone 'suffers from' a disability, we are framing their life as tragedy. When we say someone 'is wheelchair-bound,' we erase the agency they exercise every day. Words like 'has a visual impairment' or 'uses a wheelchair' are more accurate and more respectful. The third tip is to always address the patient directly — not their family member, not their care worker. Even when communication is slow or difficult, the patient is the person receiving care. The fourth is about handoff language. 'Difficult patient' is a red flag phrase. It transfers bias rather than clinical information. When you hear it, ask a clarifying question. And fifth: carry the Notice-Name-Examine-Act cycle with you. You do not need a reflective journal. You just need a habit of pausing.\n"
            "SME NOTE: Low"
        ),
        "design_guide": (
            "Format: Card carousel — click through 5 cards. Each card: tip text + small icon. Cobalt/chartreuse accent scheme.\n"
            "Interaction: Prev/next navigation (arrow keys). All card content accessible without hover.\n"
            "Accessibility: Carousel keyboard-navigable (arrow keys). All card content accessible without hover.\n"
            "Image: g02-tips-cards-01.png (flat vector icon set for tips)\n"
            "Captions: voiceover_2.17.srt"
        ),
    },
    {
        "num":   "2.18",
        "title": "Reflection Prompt",
        "step":  "Reflection",
        "activities": (
            "HEADER: Take a Moment\n"
            "PROMPT: Think about a recent interaction at work — with a patient, a colleague, or a family member. Was there a moment when you noticed yourself making an assumption? What did you do? What would you do differently now?\n"
            "PRIVATE NOTICE: This reflection is private — it is not submitted or reviewed.\n"
            "VO: voiceover_2.18.mp3\n"
            "NARRATION (~138 words): Take a moment to slow down. We have covered a lot in this guide — bias, attitudinal barriers, intersectionality, reflective practice. Now it is time to make it personal. In the text field below, take two minutes to respond to this prompt: Think about a recent interaction at work — with a patient, a colleague, or a family member. Was there a moment when you noticed yourself making an assumption? What did you do? What would you do differently now? There are no wrong answers here. This reflection is private. It will not be submitted, assessed, or reviewed. It is for you — a chance to begin the Notice-Name-Examine-Act cycle in your own experience. If nothing comes to mind right away, that is okay too. Sometimes the awareness comes later, in the next shift, in the next interaction. Take the habit with you.\n"
            "SME NOTE: Low"
        ),
        "design_guide": (
            "Format: Reflection · text entry (contenteditable div) · ungraded.\n"
            "Interaction: Text entry field is keyboard-accessible. Submit button saves to localStorage only. Text is never sent to server.\n"
            "Visual: Calm, contemplative mood. Soft-focus photo — quiet hospital garden or thoughtful person near a window.\n"
            "Image: g02-reflection-calm-01.png\n"
            "Alt: Quiet hospital garden or window view — space for reflection\n"
            "Captions: voiceover_2.18.srt"
        ),
    },
    {
        "num":   "2.19",
        "title": "MAP Action Planning",
        "step":  "Action",
        "activities": (
            "HEADER: My Action Planning (MAP)\n"
            "LEDE: Three prompts — Stop, Start, Continue. Completing at least the Stop and Start fields is required to proceed.\n"
            "FIELD 1 — STOP: One thing I will stop doing\n"
            "  Placeholder: 'I will stop making assumptions about what patients can do before asking.'\n"
            "FIELD 2 — START: One thing I will start doing\n"
            "  Placeholder: 'I will use the Notice-Name-Examine-Act cycle when I feel a reaction forming.'\n"
            "FIELD 3 — CONTINUE: One thing I will continue doing (optional)\n"
            "  Placeholder: 'I will continue asking patients directly how they prefer to be addressed.'\n"
            "DOWNLOAD: MAP-Template-Guide-02.pdf\n"
            "VO: voiceover_2.19.mp3\n"
            "NARRATION (~168 words): You are almost done. Before we reach the final summary, it is time to make a personal commitment. This is your My Action Planning activity — the MAP. The MAP appears in every guide in this series, and it is always about translating what you have learned into one concrete behaviour change. In the fields below, complete three prompts. Stop: one thing you will stop doing — an assumption, a habit, a phrase. Start: one thing you will start doing — a question you will ask, a pause you will take, a practice you will adopt. And Continue: one thing you are already doing well that you want to keep doing. Take two to three minutes. Be specific and honest. The goal is not a grand declaration. It is a small, concrete action that you can take in your next shift. Your MAP is private. It will not be reviewed or assessed. It is yours.\n"
            "CLO: 4 · SME NOTE: Low"
        ),
        "design_guide": (
            "Format: Form layout · 3 contenteditable inputs · downloadable PDF. Lilac accent (MAP activity signature colour).\n"
            "Interaction: Completing at least Stop and Start fields required to proceed. Save stores to localStorage. Download generates branded PDF.\n"
            "Accessibility: All fields keyboard-accessible. Download button clearly labelled. PDF is accessible (tagged PDF).\n"
            "Download: MAP-Template-Guide-02.pdf\n"
            "Captions: voiceover_2.19.srt\n"
            "Maps to: Completion requirement · CLO 4"
        ),
    },
    {
        "num":   "2.20",
        "title": "Key Takeaways",
        "step":  "Summary",
        "activities": (
            "HEADER: Key Takeaways\n"
            "TAKEAWAY 1: Attitudinal barriers are often invisible to those who hold them — but deeply felt by those who experience them.\n"
            "TAKEAWAY 2: Implicit bias is not a character flaw — it is a cognitive habit. The difference is whether we choose to notice and adjust.\n"
            "TAKEAWAY 3: Disability intersects with race, gender, language, age, and culture. A person's experience cannot be reduced to a single category.\n"
            "TAKEAWAY 4: Inclusive practice is not about perfection. It is about honesty, curiosity, and a willingness to ask before assuming.\n"
            "VO: voiceover_2.20.mp3\n"
            "NARRATION (~178 words): Let us recap what you have learned in this guide. First: attitudinal barriers are often invisible to those who hold them but are deeply felt by those who experience them. The most powerful barrier prevention begins with awareness — and awareness begins with a decision to look. Second: implicit bias is not a character flaw. It is how the brain works. The question is whether you can recognize it, name it, and choose differently. That is what the Notice-Name-Examine-Act cycle is for. Third: intersectionality matters. A person's experience of disability is shaped by race, gender, language, age, and culture. There is no one-size-fits-all response — only curiosity and the willingness to ask. And fourth: inclusive practice is not about getting it right every time. It is about the commitment to try, to notice when you did not, and to do better next time. These are the habits that build an accessible culture — one interaction at a time.\n"
            "SME NOTE: Low"
        ),
        "design_guide": (
            "Format: 4-card summary grid (2x2). Color-coded (Navy, Cobalt, Chartreuse, Red).\n"
            "Accessibility: Summary points are semantic list. WCAG 2.1 AA contrast.\n"
            "Captions: voiceover_2.20.srt"
        ),
    },
    {
        "num":   "2.21",
        "title": "Listen and Reflect (Podcast)",
        "step":  "Podcast",
        "activities": (
            "HEADER: Listen and Reflect\n"
            "EPISODE: Episode 02 of 18 · Perceptions, Attitudes, and Barriers\n"
            "TITLE: The Invisible Barrier: Bias, Dignity, and the Patient Experience\n"
            "DURATION: ~15-18 min (full podcast)\n"
            "DESCRIPTION: A two-host discussion exploring how bias manifests in healthcare, what dignity-centred care actually looks like, and how staff can build reflective habits. Full transcript and captions available.\n"
            "KEY LISTENING POINTS:\n"
            "  0:00 — Opening: 'Have you ever walked into a room and already had a story about the person?'\n"
            "  2:00 — Implicit vs. explicit bias: what research says\n"
            "  6:00 — The 'difficult patient' phenomenon and what it really means\n"
            "  10:00 — The spread effect in practice — real examples\n"
            "  14:00 — Dignity-centred care as a daily practice\n"
            "REFLECTION PROMPT (after listening): Think of a time when someone held a different assumption about you than the reality. How did it feel? How might that experience inform how you approach your patients?\n"
            "AUDIO: [Podcast file — to be generated via NotebookLM]\n"
            "VO: voiceover_2.21.mp3 (intro narration only)"
        ),
        "design_guide": (
            "Format: Podcast player with audio controls, transcript toggle, key listening points sidebar.\n"
            "Interaction: Play/pause, seekable progress bar, clickable timestamps, full transcript panel.\n"
            "Visual: Dark navy background. Lilac accent for episode label.\n"
            "Accessibility: Captions always available (CC toggle). Full transcript expandable. Downloadable MP3.\n"
            "Captions: voiceover_2.21.srt"
        ),
    },
    {
        "num":   "2.22",
        "title": "Series Progress Map",
        "step":  "Progress",
        "activities": (
            "HEADER: Accessibility First · A Three-Stage Journey\n"
            "LEDE: 18 guides organized in three stages. Guides unlock by stage.\n"
            "STAGE 1 — Foundations (Guides 01-04): Required first. Guides 01-02 complete. Guides 03-04 upcoming.\n"
            "STAGE 2 — Understanding Disability Experiences (Guides 05-09): Builds on Foundations. Locked.\n"
            "STAGE 3 — Applied Practice & Specialized Contexts (Guides 10-18): Unlock after Stages 1+2. Locked.\n"
            "STATUS: You have completed 2 of 4 Foundation guides. Complete Guides 03 and 04 to unlock Stage 2.\n"
            "VO: voiceover_2.22.mp3\n"
            "SME NOTE: Low"
        ),
        "design_guide": (
            "Format: 3-column card grid. Locked stages have overlay with lock icon.\n"
            "Visual: Guide 01 complete (filled dot). Guide 02 active (highlighted). Guides 03-04 upcoming (open dot). Stages 2+3 locked.\n"
            "Color: Stage 1 Red, Stage 2 Cobalt, Stage 3 Chartreuse.\n"
            "Accessibility: Progress map is keyboard-navigable. Screen reader announces progress state. WCAG 2.1 AA contrast.\n"
            "Captions: voiceover_2.22.srt"
        ),
    },
    {
        "num":   "2.23",
        "title": "Resources and Course Completion",
        "step":  "Completion",
        "activities": (
            "HEADER: Resources & Course Completion\n"
            "COMPLETION: You have completed Guide 02. You have earned the Accessibility First: Perceptions & Attitudes badge.\n"
            "RESOURCES:\n"
            "  AODA: ontario.ca/laws/statute/05a11\n"
            "  OHRC Human Rights Code: ohrc.on.ca\n"
            "  CAMH Anti-Stigma Resources: camh.ca/en/health-info/guides-and-publications\n"
            "  Mental Health Commission of Canada: mentalhealthcommission.ca\n"
            "  UHN IDEAA Office: [internal link]\n"
            "  UHN Accessibility Policy: [internal link]\n"
            "BADGE: Accessibility First: Perceptions & Attitudes\n"
            "NEXT: Guide 03 — Vision Disabilities: Inclusive Practices for Patients with Low or No Vision\n"
            "VO: voiceover_2.23.mp3\n"
            "NARRATION (~175 words): Congratulations — you have completed Guide 2: Perceptions, Attitudes, and Barriers. You have explored how implicit and explicit bias work, examined six types of attitudinal barriers, practised four branching scenarios, and committed to a personal action plan. That is meaningful work. The habits you practise in this guide — noticing assumptions, naming them, examining their source, and choosing a different response — these are not one-time lessons. They are ongoing practices. Bias does not disappear because we learned about it. It shifts because we develop the awareness and the commitment to catch it before it harms. Below you will find links to resources including the OHRC Human Rights Code, CAMH anti-stigma materials, and UHN's IDEAA office. You have also earned the Accessibility First: Perceptions and Attitudes badge. In Guide 3, we will shift to vision disabilities — exploring practical inclusive approaches for patients with low or no vision. Well done. See you in Guide 3.\n"
            "SME NOTE: Low"
        ),
        "design_guide": (
            "Format: 2-column layout. Left: completion status + resources + next guide + action buttons. Right: badge container.\n"
            "Completion criteria: All 23 slides visited + quiz >= 80% + at least Stop and Start MAP fields completed.\n"
            "Accessibility: All links keyboard-accessible. Badge has alt text. SCORM completion status sent.\n"
            "Captions: voiceover_2.23.srt"
        ),
    },
]


# ---------------------------------------------------------------------------
# Part 2: intro paragraph
# ---------------------------------------------------------------------------
def build_part2_intro(doc):
    p = doc.add_paragraph()
    add_run(p,
            "Per-screen 3-column tables (Step · Activities · Design Guide). "
            "The Design Guide column captures the interaction format, accessibility specifications, "
            "and image references. "
            "This document mirrors the deployed SCORM course exactly.",
            italic=True, size_pt=9, color_hex=BODY)
    p.paragraph_format.space_after = Pt(8)


# ---------------------------------------------------------------------------
# MAIN BUILD
# ---------------------------------------------------------------------------
def build():
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    doc = setup_document()

    # ---- COVER ----
    build_cover(doc)
    build_metadata_table(doc)
    build_how_to_read(doc)

    # ---- PART 1 BANNER ----
    add_part_banner(doc, "PART 1 — STORYBOARD FOR REVIEW")

    # ---- 1.1 – 1.6 ----
    build_1_1(doc)
    build_1_2(doc)
    build_1_3(doc)
    build_1_4(doc)
    build_1_5(doc)
    build_1_6(doc)

    # ---- PART 2 BANNER ----
    add_part_banner(doc, "PART 2 — DESIGNER TOOLKIT  (for Yi Jin / Developer)")
    build_part2_intro(doc)

    # ---- Per-screen tables ----
    for screen in SCREEN_DATA:
        add_screen_table(
            doc,
            screen_num=screen["num"],
            screen_title=screen["title"],
            step=screen["step"],
            activities=screen["activities"],
            design_guide=screen["design_guide"],
        )

    # ---- Confidentiality footer via core properties ----
    doc.core_properties.author   = "Yi Jin"
    doc.core_properties.comments = "Confidential — For Internal Use Only · UHN Accessibility First · Guide 02"

    doc.save(str(OUTPUT_PATH))
    print(f"Document saved to:\n  {OUTPUT_PATH}")


if __name__ == "__main__":
    build()
