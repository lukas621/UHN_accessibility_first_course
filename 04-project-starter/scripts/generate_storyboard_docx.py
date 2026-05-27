#!/usr/bin/env python3
"""
Master Storyboard DOCX Generator — UHN Accessibility First Series
Guide 01: Foundations of Disability, Inclusion and Accessible Design

Generates a fully-formatted Word document matching the v0.5.1 storyboard
reference format: Part 1 (review) + Part 2 (designer toolkit), 3-column tables.

Usage:
    python generate_storyboard_docx.py

Output:
    05-build-output/01-Foundations-of-Disability-Inclusion-and-Accessible-Design/
        02-production/master-storyboard/MASTER-STORYBOARD-GUIDE-01.docx

Requirements:
    pip install python-docx
"""

import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor, Emu
    from docx.enum.section import WD_ORIENT
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("Error: 'python-docx' not installed. Run: pip install python-docx")
    sys.exit(1)

# ---------------------------------------------------------------------------
# Paths
# ---------------------------------------------------------------------------
BASE_DIR = Path(__file__).resolve().parent.parent.parent
LOGO_PATH = BASE_DIR / "02-branding-and-style" / "logos" / "uhn_logo_2.png"
OUTPUT_DIR = (
    BASE_DIR
    / "05-build-output"
    / "01-Foundations-of-Disability-Inclusion-and-Accessible-Design"
    / "02-production"
    / "master-storyboard"
)
OUTPUT_PATH = OUTPUT_DIR / "MASTER-STORYBOARD-GUIDE-01.docx"

# ---------------------------------------------------------------------------
# Colour constants (hex strings — no leading #)
# ---------------------------------------------------------------------------
NAVY        = "192858"
COBALT      = "255CAA"
BODY        = "2B2B2B"
LIGHT_FILL  = "F5F5F5"
HEADER_FILL = "EDF0F7"
BORDER_CLR  = "CCCCCC"
WHITE       = "FFFFFF"
RED         = "C0233B"

# ---------------------------------------------------------------------------
# Helper: apply shading to a table cell
# ---------------------------------------------------------------------------
def shade_cell(cell, fill_hex: str):
    """Set the background fill of a table cell."""
    tc_pr = cell._element.get_or_add_tcPr()
    # Remove any existing shd elements
    for existing in tc_pr.findall(qn("w:shd")):
        tc_pr.remove(existing)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  fill_hex.upper())
    tc_pr.append(shd)


def set_cell_margins(cell, top_pt=6, bottom_pt=6, left_pt=8, right_pt=8):
    """Set cell internal margins in points (converted to twips)."""
    tc_pr = cell._element.get_or_add_tcPr()
    mar = OxmlElement("w:tcMar")
    for side, val in [("top", top_pt), ("bottom", bottom_pt),
                      ("left", left_pt), ("right", right_pt)]:
        el = OxmlElement(f"w:{side}")
        # 1 pt = 20 twips
        el.set(qn("w:w"), str(int(val * 20)))
        el.set(qn("w:type"), "dxa")
        mar.append(el)
    # Replace existing tcMar if any
    for existing in tc_pr.findall(qn("w:tcMar")):
        tc_pr.remove(existing)
    tc_pr.append(mar)


def set_table_borders(table, color_hex: str = BORDER_CLR, size_pt: int = 4):
    """Apply thin uniform borders to every cell in a table."""
    border_size = str(size_pt * 8)  # eighths of a point
    for row in table.rows:
        for cell in row.cells:
            tc_pr = cell._element.get_or_add_tcPr()
            tc_borders = OxmlElement("w:tcBorders")
            for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
                border = OxmlElement(f"w:{side}")
                border.set(qn("w:val"),   "single")
                border.set(qn("w:sz"),    border_size)
                border.set(qn("w:space"), "0")
                border.set(qn("w:color"), color_hex.upper())
                tc_borders.append(border)
            for existing in tc_pr.findall(qn("w:tcBorders")):
                tc_pr.remove(existing)
            tc_pr.append(tc_borders)


def set_col_widths(table, widths_inches: list):
    """Set column widths for a table (list of floats in inches)."""
    for i, row in enumerate(table.rows):
        for j, cell in enumerate(row.cells):
            if j < len(widths_inches):
                tc = cell._element
                tc_pr = tc.get_or_add_tcPr()
                tc_w = OxmlElement("w:tcW")
                tc_w.set(qn("w:w"),    str(int(widths_inches[j] * 1440)))
                tc_w.set(qn("w:type"), "dxa")
                for existing in tc_pr.findall(qn("w:tcW")):
                    tc_pr.remove(existing)
                tc_pr.append(tc_w)


def rgb(hex_str: str):
    h = hex_str.lstrip("#")
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


def add_run(para, text: str, bold=False, italic=False, size_pt=10,
            color_hex=BODY, font_name="Arial"):
    run = para.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.size       = Pt(size_pt)
    run.font.color.rgb  = rgb(color_hex)
    run.font.name       = font_name
    return run


def cell_para(cell, text: str, bold=False, italic=False, size_pt=10,
              color_hex=BODY, font_name="Arial", align=WD_ALIGN_PARAGRAPH.LEFT):
    """Clear a cell and add a single paragraph with one run."""
    # Clear existing paragraphs (keep the first one)
    for p in cell.paragraphs[1:]:
        p._element.getparent().remove(p._element)
    para = cell.paragraphs[0]
    para.clear()
    para.alignment = align
    add_run(para, text, bold=bold, italic=italic, size_pt=size_pt,
            color_hex=color_hex, font_name=font_name)
    return para


def cell_para_multiline(cell, lines: list, size_pt=10, color_hex=BODY,
                        font_name="Arial"):
    """Write multiple paragraphs into a cell. Each item in lines is either
    a plain str or a dict with keys: text, bold, italic, size_pt, color_hex."""
    for p in cell.paragraphs[1:]:
        p._element.getparent().remove(p._element)
    first = True
    for line in lines:
        if first:
            para = cell.paragraphs[0]
            para.clear()
            first = False
        else:
            para = cell.add_paragraph()
        if isinstance(line, dict):
            para.alignment = line.get("align", WD_ALIGN_PARAGRAPH.LEFT)
            add_run(para, line.get("text", ""),
                    bold=line.get("bold", False),
                    italic=line.get("italic", False),
                    size_pt=line.get("size_pt", size_pt),
                    color_hex=line.get("color_hex", color_hex),
                    font_name=line.get("font_name", font_name))
        else:
            add_run(para, str(line), size_pt=size_pt, color_hex=color_hex,
                    font_name=font_name)


# ---------------------------------------------------------------------------
# Document setup
# ---------------------------------------------------------------------------
def setup_document() -> Document:
    doc = Document()
    # Landscape 11 x 8.5 inches
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width  = Inches(11)
    section.page_height = Inches(8.5)
    section.left_margin   = Inches(0.7)
    section.right_margin  = Inches(0.7)
    section.top_margin    = Inches(0.55)
    section.bottom_margin = Inches(0.5)

    # Default paragraph style
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(10)
    style.font.color.rgb = rgb(BODY)

    return doc


# ---------------------------------------------------------------------------
# Reusable block builders
# ---------------------------------------------------------------------------
def add_section_heading(doc, number: str, title: str):
    """Add a numbered section heading (e.g. '1.1 Course Information')."""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(12)
    para.paragraph_format.space_after  = Pt(4)
    add_run(para, f"{number}  {title}", bold=True, size_pt=13,
            color_hex=NAVY, font_name="Arial Black")


def add_part_banner(doc, text: str):
    """Full-width navy banner for PART 1 / PART 2."""
    doc.add_paragraph()  # spacer
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = tbl.rows[0].cells[0]
    shade_cell(cell, NAVY)
    set_cell_margins(cell, top_pt=8, bottom_pt=8, left_pt=12, right_pt=12)
    cell_para(cell, text, bold=True, size_pt=14,
              color_hex=WHITE, font_name="Arial Black",
              align=WD_ALIGN_PARAGRAPH.CENTER)
    set_table_borders(tbl, color_hex=NAVY)
    doc.add_paragraph()  # spacer after


def add_two_col_table(doc, rows_data: list,
                      col_widths=(2.5, 6.8),
                      header_row: list = None):
    """
    Build a 2-column info table.
    rows_data: list of (label, value) tuples.
    header_row: optional [left_header, right_header] for tables with a header row.
    """
    num_rows = len(rows_data) + (1 if header_row else 0)
    tbl = doc.add_table(rows=num_rows, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    row_idx = 0
    if header_row:
        hrow = tbl.rows[row_idx].cells
        shade_cell(hrow[0], NAVY)
        shade_cell(hrow[1], NAVY)
        cell_para(hrow[0], header_row[0], bold=True, size_pt=10,
                  color_hex=WHITE, font_name="Arial Black")
        cell_para(hrow[1], header_row[1], bold=True, size_pt=10,
                  color_hex=WHITE, font_name="Arial Black")
        set_cell_margins(hrow[0])
        set_cell_margins(hrow[1])
        row_idx += 1

    for label, value in rows_data:
        cells = tbl.rows[row_idx].cells
        shade_cell(cells[0], HEADER_FILL)
        shade_cell(cells[1], WHITE)
        cell_para(cells[0], label, bold=True, size_pt=10, color_hex=NAVY)
        # Value may be multi-line if it contains newlines
        if "\n" in str(value):
            lines = []
            for line in str(value).split("\n"):
                lines.append({"text": line, "size_pt": 10, "color_hex": BODY})
            cell_para_multiline(cells[1], lines)
        else:
            cell_para(cells[1], str(value), size_pt=10, color_hex=BODY)
        set_cell_margins(cells[0])
        set_cell_margins(cells[1])
        row_idx += 1

    set_col_widths(tbl, list(col_widths))
    set_table_borders(tbl)
    return tbl


def add_four_col_table(doc, rows_data: list, col_widths=(1.8, 3.5, 1.2, 1.8),
                       headers=None):
    """Build a 4-column table (used for summative assessments)."""
    num_rows = len(rows_data) + (1 if headers else 0)
    tbl = doc.add_table(rows=num_rows, cols=4)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    row_idx = 0
    if headers:
        hcells = tbl.rows[0].cells
        for i, h in enumerate(headers):
            shade_cell(hcells[i], NAVY)
            cell_para(hcells[i], h, bold=True, size_pt=10,
                      color_hex=WHITE, font_name="Arial Black")
            set_cell_margins(hcells[i])
        row_idx += 1

    for r in rows_data:
        cells = tbl.rows[row_idx].cells
        for i, val in enumerate(r):
            shade_cell(cells[i], LIGHT_FILL if i == 0 else WHITE)
            cell_para(cells[i], str(val), bold=(i == 0), size_pt=10,
                      color_hex=NAVY if i == 0 else BODY)
            set_cell_margins(cells[i])
        row_idx += 1

    set_col_widths(tbl, list(col_widths))
    set_table_borders(tbl)
    return tbl


def add_five_col_table(doc, rows_data: list,
                       col_widths=(0.45, 0.65, 2.8, 2.1, 3.3),
                       headers=None):
    """Build a 5-column table (Master Screen Schedule)."""
    num_rows = len(rows_data) + (1 if headers else 0)
    tbl = doc.add_table(rows=num_rows, cols=5)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    row_idx = 0
    if headers:
        hcells = tbl.rows[0].cells
        for i, h in enumerate(headers):
            shade_cell(hcells[i], NAVY)
            cell_para(hcells[i], h, bold=True, size_pt=10,
                      color_hex=WHITE, font_name="Arial Black")
            set_cell_margins(hcells[i])
        row_idx += 1

    for r in rows_data:
        cells = tbl.rows[row_idx].cells
        for i, val in enumerate(r):
            is_first = (i == 0)
            shade_cell(cells[i], HEADER_FILL if is_first else WHITE)
            cell_para(cells[i], str(val), bold=is_first, size_pt=10,
                      color_hex=NAVY if is_first else BODY)
            set_cell_margins(cells[i])
        row_idx += 1

    set_col_widths(tbl, list(col_widths))
    set_table_borders(tbl)
    return tbl


# ---------------------------------------------------------------------------
# Merge cells helper (for banner rows in Part 2 tables)
# ---------------------------------------------------------------------------
def merge_row(table, row_idx: int, start_col: int, end_col: int):
    """Merge cells from start_col to end_col (inclusive) in a given row."""
    row = table.rows[row_idx]
    merged = row.cells[start_col]
    for c in range(start_col + 1, end_col + 1):
        merged = merged.merge(row.cells[c])
    return merged


# ---------------------------------------------------------------------------
# Part 2: Three-column screen table
# ---------------------------------------------------------------------------
def add_screen_table(doc, screen_num: str, screen_title: str,
                     step: str, activities: str, design_guide: str):
    """
    Build the 3-column designer-toolkit table for one screen.
    Layout:
      Row 0: merged header — "SCREEN {num}" | "Topic: {title}"
      Row 1: column headers — Step | Activities | Design Guide
      Row 2: content
    """
    tbl = doc.add_table(rows=3, cols=3)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    # --- Row 0: screen banner (2 cells merged visually — col 0+1 left, col 2 right) ---
    # We leave as 3 separate cells but span visually with color
    r0 = tbl.rows[0].cells
    # Merge col 0 and col 1 for left side
    left_cell = r0[0].merge(r0[1])
    shade_cell(left_cell, NAVY)
    set_cell_margins(left_cell, top_pt=6, bottom_pt=6, left_pt=8, right_pt=8)
    cell_para(left_cell, f"SCREEN {screen_num}  (Online · Self-Paced)",
              bold=True, size_pt=10, color_hex=WHITE, font_name="Arial Black")

    shade_cell(r0[2], NAVY)
    set_cell_margins(r0[2], top_pt=6, bottom_pt=6, left_pt=8, right_pt=8)
    cell_para(r0[2], f"Topic: {screen_title}",
              bold=True, size_pt=10, color_hex=WHITE, font_name="Arial Black")

    # --- Row 1: column headers ---
    r1 = tbl.rows[1].cells
    headers = ["Step", "Activities", "Design Guide"]
    for i, h in enumerate(headers):
        shade_cell(r1[i], "EDF0F7")
        cell_para(r1[i], h, bold=True, size_pt=10, color_hex=NAVY,
                  font_name="Arial Black")
        set_cell_margins(r1[i])

    # --- Row 2: content ---
    r2 = tbl.rows[2].cells

    # Step column
    shade_cell(r2[0], LIGHT_FILL)
    cell_para(r2[0], step, bold=True, size_pt=10, color_hex=NAVY)
    set_cell_margins(r2[0])

    # Activities column — may contain bullet-like lines separated by \n
    shade_cell(r2[1], WHITE)
    _fill_cell_lines(r2[1], activities, size_pt=10, color_hex=BODY)
    set_cell_margins(r2[1])

    # Design Guide column
    shade_cell(r2[2], LIGHT_FILL)
    _fill_cell_lines(r2[2], design_guide, size_pt=9, color_hex=BODY)
    set_cell_margins(r2[2])

    set_col_widths(tbl, [1.5, 4.5, 3.3])
    set_table_borders(tbl)
    doc.add_paragraph()  # spacer between screen tables


def _fill_cell_lines(cell, text: str, size_pt=10, color_hex=BODY):
    """Write text into a cell, splitting on newlines into separate paragraphs."""
    lines = str(text).split("\n")
    first = True
    for line in lines:
        if first:
            para = cell.paragraphs[0]
            para.clear()
            first = False
        else:
            para = cell.add_paragraph()
        add_run(para, line.strip(), size_pt=size_pt, color_hex=color_hex)


# ---------------------------------------------------------------------------
# COVER SECTION
# ---------------------------------------------------------------------------
def build_cover(doc):
    # UHN Logo
    if LOGO_PATH.exists():
        doc.add_picture(str(LOGO_PATH), width=Inches(2.0))
    else:
        p = doc.add_paragraph("[UHN LOGO]")
        p.runs[0].font.size = Pt(10)
        p.runs[0].font.color.rgb = rgb(COBALT)

    doc.add_paragraph()  # spacer

    # Series name
    series_para = doc.add_paragraph()
    add_run(series_para, "Accessibility First eLearning Series",
            bold=True, size_pt=12, color_hex=COBALT)

    # Title
    title_para = doc.add_paragraph()
    add_run(title_para,
            "Guide 01 — Foundations of Disability, Inclusion and Accessible Design",
            bold=True, size_pt=22, color_hex=NAVY, font_name="Arial Black")

    # Subtitle
    sub_para = doc.add_paragraph()
    add_run(sub_para,
            "Foundations of Disability, Inclusion, and Accessible Design",
            italic=True, size_pt=12, color_hex=NAVY)

    doc.add_paragraph()  # spacer


# ---------------------------------------------------------------------------
# METADATA TABLE
# ---------------------------------------------------------------------------
def build_metadata_table(doc):
    rows = [
        ("Audience",         "All UHN clinical & administrative staff"),
        ("Modality",         "Online · Self-paced · HTML5 SCORM 1.2 to UHN MyLearning (SumTotal)"),
        ("Duration",         "~15-20 minutes (23 slides)"),
        ("Standard",         "WCAG 2.1 AA · plain language · trauma-informed"),
        ("Author",           "Yi Jin (Instructional Designer)"),
        ("Content owner",    "Jacqueline Silvera — IDEAA, People & Culture, UHN"),
        ("Document version", "v1.0 Deployed Storyboard · 2026-05-25"),
    ]
    add_two_col_table(doc, rows, col_widths=(2.0, 7.3))
    doc.add_paragraph()


# ---------------------------------------------------------------------------
# HOW TO READ
# ---------------------------------------------------------------------------
def build_how_to_read(doc):
    p = doc.add_paragraph()
    add_run(p,
            "How to read this storyboard  ",
            bold=True, size_pt=10, color_hex=NAVY, font_name="Arial Black")
    p.paragraph_format.space_after = Pt(2)

    body = doc.add_paragraph()
    add_run(body,
            "This document has two parts. "
            "Part 1 (Storyboard for Review) is for SME and stakeholder review — "
            "it contains the course information, learning outcomes, assessments, "
            "and a complete screen schedule. "
            "Part 2 (Designer Toolkit) is for the instructional designer and developer — "
            "it contains per-screen 3-column tables with Step, Activities, and Design Guide details "
            "that mirror the deployed SCORM course exactly. "
            "Review Part 1 for content accuracy and sign-off. "
            "Part 2 is a production reference only.",
            italic=True, size_pt=9, color_hex=BODY)
    body.paragraph_format.space_after = Pt(8)


# ---------------------------------------------------------------------------
# PART 1 — STORYBOARD FOR REVIEW
# ---------------------------------------------------------------------------

def build_1_1(doc):
    add_section_heading(doc, "1.1", "Course Information")
    rows = [
        ("Program Name",        "UHN Accessibility First eLearning Series"),
        ("Course Title",        "Foundations of Disability, Inclusion and Accessible Design"),
        ("Course Code",         "AF-01"),
        ("Course Hours",        "~0.30 hours / 15-20 minutes (online, self-paced)"),
        ("Pre-requisite(s)",    "None — entry guide of the series"),
        ("Course Description",
         "Foundational guide of the Accessibility First Series. Establishes shared language, "
         "introduces the medical, social, and rights-based models of disability, and presents "
         "the Accessibility in Practice model and Accessibility Decision Path. Frames "
         "accessibility as a proactive, dignity-based practice — not a compliance exercise. "
         "Aligned with the AODA, the Ontario Human Rights Code, and UHN Accessibility Policy."),
        ("Course Rationale",
         "Hospital staff need shared vocabulary and a common understanding of accessibility "
         "before topic-specific guides can land. This intro guide creates that foundation, "
         "situates AODA and OHRC obligations within everyday practice, and orients learners "
         "to the Accessibility Decision Path used across the rest of the series."),
        ("Learning Materials",
         "All readings, audio, and podcast embedded inside the HTML5 SCORM module. "
         "Source: Accessibility First Guide Series 1-18 (Draft, April 2026), Guide 1."),
        ("Accrediting Standards",
         "AODA — Integrated Accessibility Standards Regulation; "
         "Ontario Human Rights Code (Part I, Disability) — Duty to Accommodate; "
         "UN CRPD; UHN IDEAA Strategic Framework; WCAG 2.1 AA."),
        ("Version / Last Update",
         "v1.0 Deployed Storyboard — 2026-05-25 (Yi Jin, Instructional Designer)"),
    ]
    add_two_col_table(doc, rows)
    doc.add_paragraph()


def build_1_2(doc):
    add_section_heading(doc, "1.2", "Course Development Information")
    rows = [
        ("Delivery Modality",   "Online · Self-Paced · HTML5 SCORM 1.2 to UHN MyLearning (SumTotal LMS)"),
        ("Course Developer(s)", "Yi Jin (Instructional Designer)"),
        ("SME(s)",              "Jacqueline Silvera (IDEAA, People & Culture, UHN); additional clinical SME — TBC"),
        ("ID(s)",               "Yi Jin"),
        ("Existing Course",     "None — new build."),
    ]
    add_two_col_table(doc, rows)
    doc.add_paragraph()


def build_1_3(doc):
    add_section_heading(doc, "1.3", "Course Learning Outcomes (CLOs)")
    rows = [
        ("CLO 1",
         "Define disability through multiple lenses — distinguish between the medical, "
         "social, and human-rights models of disability."),
        ("CLO 2",
         "Identify visible and invisible barriers — recognize the physical, sensory, "
         "attitudinal, communication, and systemic barriers a person may encounter at UHN."),
        ("CLO 3",
         "Apply the Accessibility in Practice model — use the four areas (Awareness, "
         "Communication, Environment, Response) to guide everyday decisions."),
        ("CLO 4",
         "Use the 5-step Decision Path — Pause and Assess, Listen, Apply, Adapt, and "
         "Seek Support when an accessibility need is identified."),
    ]
    add_two_col_table(doc, rows, header_row=["CLO", "Description"])
    doc.add_paragraph()


def build_1_4(doc):
    add_section_heading(doc, "1.4", "Summative Assessments")
    rows = [
        ("Knowledge Check 1",
         "2 MCQ questions (models of disability, AODA obligations). 2 attempts each.",
         "33%", "Slides 13-14"),
        ("Knowledge Check 2",
         "1 scenario-based MCQ (wheelchair accessibility). 2 attempts.",
         "17%", "Slide 15"),
        ("Branching Scenarios",
         "3 scenarios + 1 decision tree. Select-then-submit with feedback.",
         "30%", "Slides 10-12, 21"),
        ("My Action Plan (MAP)",
         "Stop / Start / Continue commitments. Completion-graded.",
         "20%", "Slide 18"),
    ]
    add_four_col_table(
        doc, rows,
        col_widths=(2.0, 4.0, 1.0, 2.3),
        headers=["Title of Assessment", "Brief Description", "Weight", "Position"],
    )
    doc.add_paragraph()


def build_1_5(doc):
    add_section_heading(doc, "1.5", "Master Screen Schedule")
    schedule = [
        ("1",  "1.1",  "Welcome and Course Purpose",              "Title/Cover",                "Hero image, Begin button"),
        ("2",  "1.2",  "What You Will Learn",                     "Objectives",                 "4 CLO cards"),
        ("3",  "1.3A", "Why This Matters: The Stat",              "Content + Image",            "27% statistic, key takeaway"),
        ("4",  "1.3B", "Impact: Missed Care",                     "Content + Image",            "Key takeaway"),
        ("5",  "1.3C", "Impact: Communication Gap",               "Content + Image",            "Key takeaway, cobalt accent"),
        ("6",  "1.3D", "Impact: Avoidance",                       "Content + Image",            "Indigenous context box"),
        ("7",  "1.4",  "Three Models of Disability",              "Tabbed panels (3)",          "Medical, Social, Rights-based"),
        ("8",  "1.5",  "Accessibility in Practice Model",         "4-quadrant cards",           "Awareness, Communication, Environment, Response"),
        ("9",  "1.6",  "Accessibility Decision Path",             "5-step stepper",             "Pause, Listen, Apply, Adapt, Seek Support"),
        ("10", "1.7",  "Scenario 1: Hospital Booking System",     "Branching (3-choice)",       "Mrs. Okafor scenario, graded"),
        ("11", "1.8",  "Scenario 2: Clinic Signage",              "Branching (3-choice)",       "Low vision patient, graded"),
        ("12", "1.9",  "Scenario 3: Employee Awareness",          "Branching (3-choice)",       "Checklist discussion, graded"),
        ("13", "1.10", "Knowledge Check 1 — Q1",                  "MCQ (4 options)",            "Models of disability, 2 attempts"),
        ("14", "1.11", "Knowledge Check 2 — Q2",                  "MCQ (4 options)",            "AODA obligations, 2 attempts"),
        ("15", "1.12", "Knowledge Check 3",                       "Scenario MCQ (4 options)",   "Wheelchair accessibility, 2 attempts"),
        ("16", "1.12", "Five Inclusive Practice Tips",            "5-card grid",                "Ask First, Language, Don't Assume, Multiple Options, Report Barriers"),
        ("17", "1.13", "Pause and Reflect",                       "Reflection (text entry)",    "Private — not submitted or reviewed"),
        ("18", "1.14", "My Action Plan (MAP)",                    "Action planning (3 fields)", "Stop/Start/Continue, downloadable PDF"),
        ("19", "1.15", "Key Takeaways",                           "Summary (4 cards)",          "4 key messages"),
        ("20", "1.16", "Listen and Reflect (Podcast)",            "Podcast player",             '"Five Words to Restore Patient Dignity" ~18 min'),
        ("21", "1.17", "Decision Tree: Intake Form Frustration",  "Decision tree (3-choice)",   "Loops back to Decision Path, graded"),
        ("22", "1.18", "Series Progress Map",                     "Progress map (3 stages)",    "Stage 1-3, locked stages"),
        ("23", "1.19", "Resources & Course Completion",           "Completion + certificate",   "Badge, retry quiz, exit course"),
    ]
    add_five_col_table(
        doc, schedule,
        col_widths=[0.45, 0.65, 2.8, 2.1, 3.3],
        headers=["Slide", "Screen", "Title", "Type", "Notes"],
    )
    doc.add_paragraph()


def build_1_6(doc):
    add_section_heading(doc, "1.6", "Version Control & Sign-off")

    # Version table
    version_rows = [
        ("Version", "Date",       "Author",                "Change Summary"),
        ("v0.1",    "2026-03-01", "Yi Jin",                "Initial draft"),
        ("v0.5",    "2026-04-15", "Yi Jin",                "SME review draft"),
        ("v0.5.1",  "2026-05-01", "Yi Jin / J. Silvera",  "Post-SME revisions"),
        ("v1.0",    "2026-05-25", "Yi Jin",                "Deployed storyboard — final"),
    ]
    tbl = doc.add_table(rows=len(version_rows), cols=4)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, row in enumerate(version_rows):
        cells = tbl.rows[i].cells
        for j, val in enumerate(row):
            is_header = (i == 0)
            shade_cell(cells[j], NAVY if is_header else (HEADER_FILL if j == 0 else WHITE))
            cell_para(cells[j], val, bold=is_header, size_pt=10,
                      color_hex=WHITE if is_header else (NAVY if j == 0 else BODY),
                      font_name="Arial Black" if is_header else "Arial")
            set_cell_margins(cells[j])
    set_col_widths(tbl, [0.8, 1.2, 2.5, 4.8])
    set_table_borders(tbl)

    doc.add_paragraph()

    # Sign-off box
    signoff_tbl = doc.add_table(rows=1, cols=1)
    cell = signoff_tbl.rows[0].cells[0]
    shade_cell(cell, LIGHT_FILL)
    set_cell_margins(cell, top_pt=10, bottom_pt=10, left_pt=12, right_pt=12)

    lines = [
        {"text": "SIGN-OFF", "bold": True, "size_pt": 11, "color_hex": NAVY,
         "font_name": "Arial Black"},
        {"text": ""},
        {"text": "SME Review — Jacqueline Silvera (IDEAA, People & Culture, UHN)"},
        {"text": "Reviewed by: ________________________    Date: ____________"},
        {"text": ""},
        {"text": "ID Review — Yi Jin (Instructional Designer)"},
        {"text": "Approved by: ________________________    Date: ____________"},
        {"text": ""},
        {"text": "Note: Signature on this document constitutes approval of all content, "
                 "scenarios, assessment questions, and learning outcomes as accurate, "
                 "appropriate for the intended audience, and aligned with UHN policy.",
         "italic": True, "size_pt": 9, "color_hex": BODY},
    ]
    cell_para_multiline(cell, lines)
    set_table_borders(signoff_tbl, color_hex=COBALT)
    doc.add_paragraph()


# ---------------------------------------------------------------------------
# PART 2 — DESIGNER TOOLKIT: Per-screen data
# ---------------------------------------------------------------------------

SCREEN_DATA = [
    {
        "num":   "1.1",
        "title": "Welcome and Course Purpose",
        "step":  "Hook",
        "activities": (
            "TITLE: Foundations of Disability, Inclusion & Accessible Design\n"
            "SUBTITLE: Guide 01 of 18 · Foundations\n"
            "BODY: Over the next 15 to 20 minutes, you'll build the foundation for this 18-guide series.\n"
            "METADATA: Duration 15-20 min · Audience All UHN Staff · Version v1.0 · 2026\n"
            "BUTTON: Begin\n"
            "VO: voiceover_1.1.mp3"
        ),
        "design_guide": (
            "Format: Cover slide · static layout with autoplay narration.\n"
            "Interaction: Begin button navigates to Slide 2.\n"
            "Accessibility: 1.4.3 contrast >= 4.5:1 · 2.1.1 keyboard reachable · CC ON by default.\n"
            "Visual: UHN navy background gradient, hero image right side, metadata row bottom.\n"
            "Image: g01-hero-welcome-01.png\n"
            "(A diverse group of healthcare workers walking toward a hospital entrance)"
        ),
    },
    {
        "num":   "1.2",
        "title": "What You Will Learn",
        "step":  "Outcomes",
        "activities": (
            "HEADER: What You Will Learn\n"
            "LEDE: By the end of this guide, you'll be able to do four things:\n"
            "CLO 1: Define disability through multiple lenses — Distinguish between the medical, social, and human-rights models of disability.\n"
            "CLO 2: Identify visible and invisible barriers — Recognize the physical, sensory, attitudinal, communication, and systemic barriers a person may encounter at UHN.\n"
            "CLO 3: Apply the Accessibility in Practice model — Use the four areas (Awareness, Communication, Environment, Response) to guide everyday decisions.\n"
            "CLO 4: Use the 5-step Decision Path — Pause and Assess, Listen, Apply, Adapt, and Seek Support when an accessibility need is identified.\n"
            "VO: voiceover_1.2.mp3"
        ),
        "design_guide": (
            "Format: Static objective cards (4 cards, numbered 01-04).\n"
            "Interaction: None — orientation slide. Continue always available.\n"
            "Accessibility: Tab order Header -> 1 -> 2 -> 3 -> 4. Cards have tabindex=\"0\"."
        ),
    },
    {
        "num":   "1.3A",
        "title": "Why This Matters: The Stat",
        "step":  "Context",
        "activities": (
            "HEADER: 1 in 4 Canadians live with a disability\n"
            "STAT: 27% — Statistics Canada, 2022\n"
            "BODY: According to Statistics Canada, more than 27% of Canadians aged 15 and older have at least one disability. In Ontario, that number is even higher. That means more than one in four people who come to UHN for care may experience barriers. And many disabilities are non-visible.\n"
            "KEY TAKEAWAY: Disability is not rare — it is part of everyday life at UHN. Accessible care is not optional; it is essential.\n"
            "VO: voiceover_1.3a.mp3"
        ),
        "design_guide": (
            "Format: Content + image layout. Image left, text right.\n"
            "Image: g01-context-barriers-01.png\n"
            "(A young woman in a wheelchair reaching toward directional signs in a hospital corridor)"
        ),
    },
    {
        "num":   "1.3B",
        "title": "Impact: Missed Care",
        "step":  "Impact",
        "activities": (
            "HEADER: Impact: Missed Care\n"
            "BODY: Patients avoid appointments because booking systems are unusable, entrances are inaccessible, or intake forms create barriers before care even begins. Healthcare systems were often designed without disability in mind — narrow doorways, small print on forms, phone-only booking, waiting rooms with no seating options.\n"
            "KEY TAKEAWAY: Every missed appointment is a missed diagnosis.\n"
            "VO: voiceover_1.3b.mp3"
        ),
        "design_guide": (
            "Format: Content + image layout.\n"
            "Image: g01-impact-missed-care-01.png\n"
            "(An empty hospital waiting area with a self-check-in kiosk mounted too high for wheelchair users)"
        ),
    },
    {
        "num":   "1.3C",
        "title": "Impact: Communication Gap",
        "step":  "Impact",
        "activities": (
            "HEADER: Impact: Communication Gap\n"
            "BODY: Misdiagnosis and inadequate treatment occur when communication barriers are not addressed — wrong assumptions about what a patient needs, speaking too quickly, or relying on written-only instructions. A patient who cannot access information in their preferred format cannot participate fully in their own care.\n"
            "KEY TAKEAWAY: Communication is care. If the message doesn't land, the care doesn't either.\n"
            "VO: voiceover_1.3c.mp3"
        ),
        "design_guide": (
            "Format: Content + image layout. Cobalt accent on key takeaway.\n"
            "Image: g01-impact-communication-gap-01.png\n"
            "(A healthcare worker at a computer speaking over their shoulder while a patient with a hearing aid leans forward)"
        ),
    },
    {
        "num":   "1.3D",
        "title": "Impact: Avoidance",
        "step":  "Impact",
        "activities": (
            "HEADER: Impact: Avoidance\n"
            "BODY: Patients delay or abandon care entirely because past experiences were inaccessible, dismissive, or dehumanizing. This is especially true for people who have experienced compounded barriers.\n"
            "INDIGENOUS CONTEXT: Indigenous peoples in Canada face compounded barriers — historical trauma, systemic racism, and geographic isolation intersect with disability to create unique challenges in accessing healthcare.\n"
            "VO: voiceover_1.3d.mp3"
        ),
        "design_guide": (
            "Format: Content + image layout. Indigenous context box in chartreuse/green accent.\n"
            "Image: g01-impact-avoidance-01.png\n"
            "(A man in a wheelchair sitting on the sidewalk outside a hospital entrance)\n"
            "SME: HIGH — Indigenous content requires SME sign-off."
        ),
    },
    {
        "num":   "1.4",
        "title": "Three Models of Disability",
        "step":  "Core Concept",
        "activities": (
            "HEADER: Three Models of Disability\n"
            "LEDE: How we think about disability shapes how we respond to it.\n"
            "TAB 1 — Medical Model: Focus: Individual deficit — \"What's wrong with this person?\" The medical model focuses on what's \"wrong\" with a person. It treats disability as a condition to be fixed. This model has its place in clinical care, but when it's the only lens, it puts all the responsibility on the individual.\n"
            "TAB 2 — Social Model: Focus: Environmental barriers — \"What barriers exist in our systems?\" The social model shifts the focus to the environment. It says disability happens when barriers in our systems, spaces, and attitudes prevent someone from participating. A person who uses a wheelchair isn't disabled by the wheelchair — they're disabled by stairs.\n"
            "TAB 3 — Rights-Based Model: Focus: Human rights and equal participation (UN CRPD). The rights-based model goes further. It says people with disabilities have a legal and moral right to full participation. This is the model behind the UN Convention on the Rights of Persons with Disabilities and Ontario's AODA.\n"
            "UHN APPROACH: At UHN, we use rights-based and social model thinking — focusing on removing barriers rather than expecting individuals to adapt.\n"
            "VO: voiceover_1.4.mp3"
        ),
        "design_guide": (
            "Format: Tabbed panels (3 tabs).\n"
            "Interaction: Click tab to reveal content. One panel visible at a time.\n"
            "Accessibility: role=\"tablist\", aria-selected. Tab order: tab bar -> panel content."
        ),
    },
    {
        "num":   "1.5",
        "title": "Accessibility in Practice Model",
        "step":  "Framework",
        "activities": (
            "HEADER: The Accessibility in Practice Model\n"
            "LEDE: Four interconnected areas you can use in any situation — at a patient's bedside, at the front desk, in a team meeting, or when designing a new process.\n"
            "CARD 1 — Awareness: Recognize that accessibility needs may not be visible. Be mindful of bias and systemic barriers.\n"
            "  Sub-items: Scan environments and forms. Notice assistive devices. Watch your own assumptions.\n"
            "CARD 2 — Communication: Listen, ask, adapt to individual needs and cultural context.\n"
            "  Sub-items: Ask about preferred format. Speak to the patient, not the companion. Slow down when needed.\n"
            "CARD 3 — Environment: Ensure physical, sensory, and social spaces are accessible and safe.\n"
            "  Sub-items: Clear paths and seating. Lighting and signage. Quiet space if needed.\n"
            "CARD 4 — Response: Take timely, respectful action to remove barriers.\n"
            "  Sub-items: Offer — don't impose — help. Confirm what worked. Update the chart or note.\n"
            "SERIES FRAMEWORK: This model will appear in every guide in this series.\n"
            "VO: voiceover_1.5.mp3"
        ),
        "design_guide": (
            "Format: 4-quadrant card grid. Each card has number, heading, body, and expandable detail.\n"
            "Interaction: Cards are tabbable (tabindex=\"0\"). Detail text reveals on focus/click.\n"
            "Visual: Color-coded top bars (Navy, Cobalt, Chartreuse, Red)."
        ),
    },
    {
        "num":   "1.6",
        "title": "Accessibility Decision Path",
        "step":  "Tool",
        "activities": (
            "HEADER: The Accessibility Decision Path\n"
            "LEDE: When an accessibility question comes up — and it will — you need a practical approach. Click each step to explore.\n"
            "STEP 1 — Pause & Assess: Recognize the need. Before you act, take a moment. Is there an accessibility need here? Don't assume — observe.\n"
            "STEP 2 — Listen & Ask: Engage respectfully.\n"
            "STEP 3 — Apply: Use AiP principles.\n"
            "STEP 4 — Adapt: Adjust as needed.\n"
            "STEP 5 — Seek Support: Consult resources.\n"
            "SERIES TOOL: This path isn't a one-time checklist. It's a way of thinking. You'll see it in every guide in this series, applied to different situations and different types of disability.\n"
            "VO: voiceover_1.6.mp3"
        ),
        "design_guide": (
            "Format: 5-step stepper rail with detail panel.\n"
            "Interaction: Click each step to reveal its detail. Step-through progression.\n"
            "Note: Currently only Step 1 detail is implemented in HTML. Steps 2-5 need detail content added."
        ),
    },
    {
        "num":   "1.7",
        "title": "Scenario 1: Hospital Booking System",
        "step":  "Scenario",
        "activities": (
            "HEADER: Scenario 1: Hospital Booking System\n"
            "SCENARIO: Mrs. Okafor's Booking Difficulty. You're working at the front desk when Mrs. Okafor, a 68-year-old patient, arrives looking frustrated. She explains that she tried to use the new online booking system but couldn't navigate it. The system meets accessibility standards, but she has limited digital literacy and finds it confusing. She's been trying for three days and finally came in person.\n"
            "QUESTION: What do you do?\n"
            "OPTION A (Best): You apologize for the difficulty, book her appointment right away, and let her know she can always call or come in person. You note the issue for your team lead.\n"
            "OPTION B (Acceptable): You book her appointment and suggest she ask a family member to help next time.\n"
            "OPTION C (Poor): You tell her the system is accessible and offer to walk her through it on your computer.\n"
            "DEBRIEF: Accessibility includes usability. Meeting technical standards is not enough. Offer multiple options and report barriers upstream.\n"
            "VO: voiceover_1.7.mp3"
        ),
        "design_guide": (
            "Format: Branching choice · 3 options · graded. Select-then-submit with feedback overlay.\n"
            "Image: g01-scenario-booking-01.png\n"
            "(An older woman at a hospital reception desk)"
        ),
    },
    {
        "num":   "1.8",
        "title": "Scenario 2: Clinic Signage",
        "step":  "Scenario",
        "activities": (
            "HEADER: Scenario 2: Clinic Signage\n"
            "SCENARIO: Patient with Low Vision. You notice a patient with low vision squinting at a directional sign in the outpatient clinic. The sign was recently installed and meets Ontario Building Code standards, but the font is small and it's mounted high on the wall. The patient asks you for directions to the lab.\n"
            "QUESTION: What do you do?\n"
            "OPTION A (Best): You walk the patient to the lab, then report the signage issue to facilities with a suggestion to add larger, lower signs with high-contrast text.\n"
            "OPTION B (Acceptable): You give clear verbal directions and offer to walk them partway.\n"
            "OPTION C (Poor): You point to the sign and read it aloud.\n"
            "DEBRIEF: Meeting technical standards doesn't guarantee practical accessibility. Consider accessibility from the user's perspective. Report barriers — don't just work around them.\n"
            "VO: voiceover_1.8.mp3"
        ),
        "design_guide": (
            "Format: Branching choice · 3 options · graded.\n"
            "Image: g01-scenario-signage-01.png\n"
            "(An elderly woman using a magnifying glass to read a small directional sign)"
        ),
    },
    {
        "num":   "1.9",
        "title": "Scenario 3: Employee Awareness",
        "step":  "Scenario",
        "activities": (
            "HEADER: Scenario 3: Employee Awareness\n"
            "SCENARIO: Checklist Discussion. Your team uses an accessibility checklist when supporting patients. A colleague mentions that a patient who speaks Cantonese and has a cognitive disability seemed confused during intake, even though the checklist was completed. Your colleague says, \"We followed the checklist — I'm not sure what else we can do.\"\n"
            "QUESTION: What do you do?\n"
            "OPTION A (Best): You acknowledge the checklist is a starting point but suggest the team explore what barriers the patient actually experienced — language support, cognitive load, cultural context — and adapt the process.\n"
            "OPTION B (Acceptable): You suggest requesting an interpreter for next time.\n"
            "OPTION C (Poor): You agree the checklist was completed and move on.\n"
            "DEBRIEF: Accessibility requires awareness, curiosity, and flexibility beyond any checklist. Take time to listen, ask respectful questions, and adapt to individual needs.\n"
            "VO: voiceover_1.9.mp3"
        ),
        "design_guide": (
            "Format: Branching choice · 3 options · graded.\n"
            "Image: g01-scenario-checklist-01.png\n"
            "(Three healthcare workers gathered around a tablet)"
        ),
    },
    {
        "num":   "1.10",
        "title": "Knowledge Check 1 (Q1)",
        "step":  "Assessment",
        "activities": (
            "HEADER: Knowledge Check 1\n"
            "LEDE: Two questions. Select the best answer for each.\n"
            "QUESTION: Which model of disability focuses on removing barriers in the environment rather than fixing the individual?\n"
            "A: Medical model (INCORRECT)\n"
            "B: Social model (CORRECT)\n"
            "C: Charity model (INCORRECT)\n"
            "D: Rights-based model (INCORRECT)\n"
            "FEEDBACK: Correct answer: B. The social model focuses on barriers in systems and environments, not individual deficits. The social model says people are disabled by barriers in society — stairs instead of ramps, information only in print, attitudes that exclude.\n"
            "VO: voiceover_1.10.mp3"
        ),
        "design_guide": (
            "Format: MCQ · 4 options · 2 attempts · submit-then-lock.\n"
            "Interaction: Select option -> Submit Answer button enables -> feedback panel appears. After 2 wrong attempts, correct answer revealed.\n"
            "Maps to: CLO 1"
        ),
    },
    {
        "num":   "1.11",
        "title": "Knowledge Check 2 (Q2)",
        "step":  "Assessment",
        "activities": (
            "HEADER: Knowledge Check 2\n"
            "QUESTION: Under the AODA, which of the following is an obligation for Ontario public sector organizations like UHN?\n"
            "A: Provide disability awareness training to all staff (CORRECT)\n"
            "B: Hire a minimum percentage of employees with disabilities (INCORRECT)\n"
            "C: Eliminate all physical barriers within 5 years (INCORRECT)\n"
            "D: Assign a dedicated accessibility officer (INCORRECT)\n"
            "FEEDBACK: Correct answer: A. Under the AODA, public sector organizations must provide accessibility training to all employees and volunteers. The other options are not specific AODA requirements, though they may be good practices.\n"
            "VO: voiceover_1.11.mp3"
        ),
        "design_guide": (
            "Format: MCQ · 4 options · 2 attempts · submit-then-lock.\n"
            "Maps to: CLO 2"
        ),
    },
    {
        "num":   "1.12",
        "title": "Knowledge Check 3",
        "step":  "Assessment",
        "activities": (
            "HEADER: Knowledge Check 3\n"
            "LEDE: One scenario-based question.\n"
            "QUESTION: A patient arrives at your unit using a power wheelchair. You notice the examination room door is too narrow for the chair. What is the BEST first step?\n"
            "A: Ask the patient to transfer to a standard chair (INCORRECT)\n"
            "B: Find an accessible room and apologize for the inconvenience (CORRECT)\n"
            "C: Call maintenance to widen the door (INCORRECT)\n"
            "D: Document the issue and continue with the scheduled room (INCORRECT)\n"
            "FEEDBACK: Correct answer: B. Address the immediate need first — find an accessible room. Then report the barrier for systemic fix. Consider the Accessibility Decision Path — first address the person's immediate need (Pause and Assess), then seek support for systemic improvement.\n"
            "VO: voiceover_1.14.mp3"
        ),
        "design_guide": (
            "Format: Scenario-based MCQ · 4 options · 2 attempts · submit-then-lock.\n"
            "Maps to: CLO 3, CLO 4"
        ),
    },
    {
        "num":   "1.12b",
        "title": "Five Inclusive Practice Tips",
        "step":  "Practice",
        "activities": (
            "HEADER: Five Inclusive Practice Tips\n"
            "LEDE: Five practices you can start using right away.\n"
            "TIP 1: Ask First — Ask before helping. Offer assistance but respect the person's choice.\n"
            "TIP 2: Language — Use people-first language unless the person prefers identity-first. Ask.\n"
            "TIP 3: Don't Assume — Don't assume what someone can or can't do based on visible cues.\n"
            "TIP 4: Multiple Options — Offer multiple ways to access services — phone, in-person, online.\n"
            "TIP 5: Report Barriers — Report barriers. Don't just work around them — flag them for systemic improvement.\n"
            "VO: voiceover_1.15.mp3"
        ),
        "design_guide": (
            "Format: 5-card static grid. Color-coded top bars (Red, Navy, Cobalt, Chartreuse, Lilac).\n"
            "Note: Storyboard originally specified a card carousel with prev/next. Built as static grid instead."
        ),
    },
    {
        "num":   "1.13",
        "title": "Pause and Reflect",
        "step":  "Reflection",
        "activities": (
            "HEADER: Pause and Reflect\n"
            "PROMPT: Think about a recent interaction at work. Was there a moment where accessibility could have been handled differently? What would you change?\n"
            "PRIVATE NOTICE: This reflection is private — it is not submitted or reviewed.\n"
            "BUTTON: SUBMIT REFLECTION (saves locally only)\n"
            "VO: voiceover_1.16.mp3"
        ),
        "design_guide": (
            "Format: Reflection · text entry (contenteditable div) · ungraded.\n"
            "Image: g01-reflection-01.png\n"
            "(A healthcare worker sitting thoughtfully at a desk)\n"
            "Note: Submit button saves to localStorage only. Text is never sent to server."
        ),
    },
    {
        "num":   "1.14",
        "title": "My Action Plan (MAP)",
        "step":  "Action",
        "activities": (
            "HEADER: My Action Plan (MAP)\n"
            "LEDE: Three prompts — Stop, Start, Continue. Completing at least one field is required for course completion.\n"
            "FIELD 1 — STOP: One thing I will stop doing\n"
            "FIELD 2 — START: One thing I will start doing\n"
            "FIELD 3 — CONTINUE: One thing I will continue doing (optional)\n"
            "DOWNLOAD: MAP-Template-Guide-01.pdf\n"
            "BUTTON: SAVE MY MAP\n"
            "VO: voiceover_1.17.mp3"
        ),
        "design_guide": (
            "Format: Form layout · 3 contenteditable inputs · downloadable PDF.\n"
            "Interaction: Completing at least one field required for course completion. Save stores to localStorage. Download generates branded PDF.\n"
            "Maps to: Completion requirement"
        ),
    },
    {
        "num":   "1.15",
        "title": "Key Takeaways",
        "step":  "Summary",
        "activities": (
            "HEADER: Key Takeaways\n"
            "LEDE: Four things to remember from this guide.\n"
            "TAKEAWAY 1: Beyond Physical Barriers — Accessibility goes beyond physical barriers — it includes attitudes, systems, communication, digital tools, and policies.\n"
            "TAKEAWAY 2: Remove Barriers, Not Fix People — The social and rights-based models focus on removing barriers, not fixing people.\n"
            "TAKEAWAY 3: Decision Path — The Accessibility Decision Path (Pause -> Listen -> Apply -> Adapt -> Seek Support) gives you a practical tool for any situation.\n"
            "TAKEAWAY 4: Everyone's Responsibility — Accessibility is everyone's responsibility at UHN — not just one department.\n"
            "VO: voiceover_1.18.mp3"
        ),
        "design_guide": (
            "Format: 4-card summary grid (2x2). Color-coded (Red, Cobalt, Chartreuse, Navy)."
        ),
    },
    {
        "num":   "1.16",
        "title": "Listen and Reflect (Podcast)",
        "step":  "Podcast",
        "activities": (
            "HEADER: Listen and Reflect\n"
            "EPISODE: Episode 01 of 18 · Foundations\n"
            "TITLE: Five Words to Restore Patient Dignity\n"
            "DURATION: ~17:51\n"
            "AUDIO: Five_words_to_restore_patient_dignity.m4a\n"
            "KEY LISTENING POINTS:\n"
            "  0:00 — $50M wing vs 30 seconds at the front desk\n"
            "  2:36 — The 30-second intake test\n"
            "  8:52 — Five words: \"What works best for you?\"\n"
            "  14:02 — One sentence for a busy shift\n"
            "  16:40 — Where are you defaulting to the fastest path?\n"
            "REFLECTION: Think of one patient interaction this week. Which moment would have shifted if you had paused to ask, before you acted?\n"
            "VO: voiceover_1.19.mp3 (intro narration only)"
        ),
        "design_guide": (
            "Format: Podcast player with audio controls, transcript toggle, key listening points sidebar.\n"
            "Interaction: Play/pause, seekable progress bar, clickable timestamps, full transcript panel.\n"
            "Visual: Dark navy background. Lilac accent for episode label."
        ),
    },
    {
        "num":   "1.17",
        "title": "Decision Tree: Intake Form Frustration",
        "step":  "Decision Tree",
        "activities": (
            "HEADER: Decision Tree: Intake Form Frustration\n"
            "SCENARIO: Before You Assume the Cause. The patient is frustrated trying to complete the intake form. You don't yet know whether the barrier is vision, language, digital access, cognitive load, or something else. The lobby is busy and three people are waiting behind them.\n"
            "QUESTION: What do you do next, before assuming the cause?\n"
            "OPTION A (Poor): Continue with the standard intake process. Keep working through the form so the line keeps moving.\n"
            "OPTION B (Best): Pause and ask how to support the patient. \"Is there anything I can do to make this form easier today?\"\n"
            "OPTION C (Acceptable): Ask another staff member to take over. Hand off to a colleague while you keep the queue moving.\n"
            "DEBRIEF: Pausing and asking, before assuming, is the single behaviour that distinguishes accessible care from \"trying our best.\"\n"
            "VO: voiceover_1.17.mp3"
        ),
        "design_guide": (
            "Format: Decision tree · 3 choices · loops back to Decision Path. Select-then-submit with feedback overlay.\n"
            "Image: g01-practice-tips-01.png\n"
            "(A nurse at eye level beside a patient in a power wheelchair)"
        ),
    },
    {
        "num":   "1.18",
        "title": "Series Progress Map",
        "step":  "Progress",
        "activities": (
            "HEADER: Accessibility First · A Three-Stage Journey\n"
            "LEDE: 18 guides organized in three stages. Guides unlock by stage. Foundations opens everything else.\n"
            "STAGE 1 — Foundations (Guides 01-04): Required first. Guide 01 highlighted as CURRENT.\n"
            "STAGE 2 — Understanding Disability Experiences (Guides 05-09): Builds on Foundations. Locked.\n"
            "STAGE 3 — Applied Practice (Guides 10-18): Unlock after Stages 1+2. Locked.\n"
            "VO: voiceover_1.18.mp3"
        ),
        "design_guide": (
            "Format: 3-column card grid. Locked stages have overlay with lock icon.\n"
            "Color: Stage 1 Red, Stage 2 Cobalt, Stage 3 Chartreuse."
        ),
    },
    {
        "num":   "1.19",
        "title": "Resources & Course Completion",
        "step":  "Completion",
        "activities": (
            "HEADER: Resources & Course Completion\n"
            "COMPLETION: Dynamic completion status message (JS-generated based on quiz score + slide progress + MAP completion).\n"
            "RESOURCES:\n"
            "  AODA: ontario.ca/laws/statute/05a11\n"
            "  OHRC: ohrc.on.ca\n"
            "  UHN Accessibility Policy: [internal link]\n"
            "  UHN IDEAA Office: [internal contact]\n"
            "UP NEXT: Guide 02: Perceptions, Attitudes, and Barriers\n"
            "BUTTONS: RETRY QUIZ, ACTION PLAN (download MAP), EXIT COURSE\n"
            "BADGE: \"Accessibility First: Foundations\" — Guide 01 Complete. Download badge button.\n"
            "VO: voiceover_1.19.mp3"
        ),
        "design_guide": (
            "Format: 2-column layout. Left: completion status + resources + next guide + action buttons. Right: badge container with download.\n"
            "Completion criteria: All 23 slides visited + quiz >= 80% + at least one MAP field completed."
        ),
    },
]


# ---------------------------------------------------------------------------
# Part 2: intro paragraph
# ---------------------------------------------------------------------------
def build_part2_intro(doc):
    p = doc.add_paragraph()
    add_run(p,
            "Per-screen 3-column tables (Step · Activities · Design Guide). "
            "The Design Guide column captures the interaction format, accessibility specifications, "
            "and image references. "
            "This document mirrors the deployed SCORM course exactly.",
            italic=True, size_pt=9, color_hex=BODY)
    p.paragraph_format.space_after = Pt(8)


# ---------------------------------------------------------------------------
# MAIN BUILD
# ---------------------------------------------------------------------------
def build():
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    doc = setup_document()

    # ---- COVER ----
    build_cover(doc)
    build_metadata_table(doc)
    build_how_to_read(doc)

    # ---- PART 1 BANNER ----
    add_part_banner(doc, "PART 1 — STORYBOARD FOR REVIEW")

    # ---- 1.1 – 1.6 ----
    build_1_1(doc)
    build_1_2(doc)
    build_1_3(doc)
    build_1_4(doc)
    build_1_5(doc)
    build_1_6(doc)

    # ---- PART 2 BANNER ----
    add_part_banner(doc, "PART 2 — DESIGNER TOOLKIT  (for Yi Jin / Developer)")
    build_part2_intro(doc)

    # ---- Per-screen tables ----
    for screen in SCREEN_DATA:
        add_screen_table(
            doc,
            screen_num=screen["num"],
            screen_title=screen["title"],
            step=screen["step"],
            activities=screen["activities"],
            design_guide=screen["design_guide"],
        )

    # ---- Confidentiality footer via core properties ----
    doc.core_properties.author   = "Yi Jin"
    doc.core_properties.comments = "Confidential — For Internal Use Only · UHN Accessibility First · Guide 01"

    doc.save(str(OUTPUT_PATH))
    print(f"Document saved to:\n  {OUTPUT_PATH}")


if __name__ == "__main__":
    build()
