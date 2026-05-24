#!/usr/bin/env python3
"""
Markdown to Word Converter for UHN Accessibility First Course Series
Converts all .md files in a guide folder to formatted .docx documents.

Usage:
    python convert_to_word.py --guide 1                  # Convert all MD files in Guide 1
    python convert_to_word.py --guide 1 --file master    # Only master storyboard
    python convert_to_word.py --all                      # All guides
    python convert_to_word.py --guide 1 --list           # List available MD files

Requirements:
    pip install python-docx
"""

import argparse
import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

BASE_DIR = Path(__file__).resolve().parent.parent.parent
BUILD_OUTPUT = BASE_DIR / "05-build-output"
UHN_TEMPLATE = BASE_DIR / "02-branding-and-style" / "uhn-templates" / "Primary_StandardWord (1).docx"
UHN_LOGO = BASE_DIR / "02-branding-and-style" / "logos" / "uhn_logo.png"

GUIDE_FOLDERS = {
    1: "01-Foundations-of-Disability-Inclusion-and-Accessible-Design",
    2: "02-Perceptions-Attitudes-and-Barriers",
    3: "03-Vision-Disabilities",
    4: "04-Sensory-Hearing-and-Communication-Disabilities",
    5: "05-Physical-Disabilities-and-Mobility",
    6: "06-Mental-Health-Disabilities",
    7: "07-Intellectual-Developmental-and-Learning-Disabilities",
    8: "08-Non-Visible-Disabilities",
    9: "09-Aging-Disability-and-Intersectionality",
    10: "10-Engaging-with-Confidence-and-Respect",
    11: "11-Service-Animals-Guide-Dogs-and-Non-Service-Animals",
    12: "12-Support-Persons",
    13: "13-Assistive-Devices",
    14: "14-Communication-and-Information-Accessibility",
    15: "15-Neurodiversity-and-Sensory-Regulation",
    16: "16-Trauma-Informed-Accessibility",
    17: "17-Accessibility-in-Crisis-Situations-and-De-escalation",
    18: "18-Indigenous-Peoples-and-Accessibility",
}

# UHN Brand Colours
UHN_NAVY = RGBColor(0x19, 0x28, 0x58)
UHN_COBALT = RGBColor(0x24, 0x5B, 0xAA)
UHN_TEAL = RGBColor(0x00, 0xA5, 0xA8)
UHN_RED = RGBColor(0xC0, 0x23, 0x3B)
DARK_GRAY = RGBColor(0x2E, 0x3D, 0x49)
LIGHT_GRAY = RGBColor(0xF5, 0xF7, 0xFA)


def setup_styles(doc: Document) -> None:
    """Configure document styles with UHN branding."""
    # Normal style
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    font.color.rgb = DARK_GRAY
    pf = style.paragraph_format
    pf.space_after = Pt(6)
    pf.line_spacing = 1.15

    # Heading 1 — UHN Navy
    h1 = doc.styles['Heading 1']
    h1.font.name = 'Calibri'
    h1.font.size = Pt(22)
    h1.font.color.rgb = UHN_NAVY
    h1.font.bold = True
    h1.paragraph_format.space_before = Pt(24)
    h1.paragraph_format.space_after = Pt(12)

    # Heading 2 — UHN Cobalt
    h2 = doc.styles['Heading 2']
    h2.font.name = 'Calibri'
    h2.font.size = Pt(16)
    h2.font.color.rgb = UHN_COBALT
    h2.font.bold = True
    h2.paragraph_format.space_before = Pt(18)
    h2.paragraph_format.space_after = Pt(8)

    # Heading 3 — UHN Teal
    h3 = doc.styles['Heading 3']
    h3.font.name = 'Calibri'
    h3.font.size = Pt(13)
    h3.font.color.rgb = UHN_TEAL
    h3.font.bold = True
    h3.paragraph_format.space_before = Pt(12)
    h3.paragraph_format.space_after = Pt(6)

    # Heading 4
    h4 = doc.styles['Heading 4']
    h4.font.name = 'Calibri'
    h4.font.size = Pt(11)
    h4.font.color.rgb = UHN_NAVY
    h4.font.bold = True
    h4.paragraph_format.space_before = Pt(8)
    h4.paragraph_format.space_after = Pt(4)


def add_header_footer(doc: Document, title: str, guide_num: int) -> None:
    """Add header and footer to the document."""
    section = doc.sections[0]

    # Header
    header = section.header
    header_para = header.paragraphs[0]
    header_para.text = f"Accessibility First Series | Guide {guide_num:02d}"
    header_para.style = doc.styles['Normal']
    header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header_para.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
        run.font.italic = True

    # Footer
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.text = "University Health Network | Confidential — For Internal Use Only"
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in footer_para.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    # Margins
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)


def parse_markdown_table(lines: list[str]) -> list[list[str]]:
    """Parse a markdown table into rows of cells."""
    rows = []
    for line in lines:
        line = line.strip()
        if not line.startswith('|'):
            continue
        # Skip separator row (|---|---|)
        if re.match(r'^\|[\s\-:|]+\|$', line):
            continue
        cells = [c.strip() for c in line.split('|')[1:-1]]
        if cells:
            rows.append(cells)
    return rows


def add_table(doc: Document, rows: list[list[str]]) -> None:
    """Add a formatted table to the document."""
    if not rows:
        return

    num_cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, row_data in enumerate(rows):
        row = table.rows[i]
        for j, cell_text in enumerate(row_data):
            if j < len(row.cells):
                cell = row.cells[j]
                cell.text = ''
                p = cell.paragraphs[0]

                # Handle bold markdown in cells
                parts = re.split(r'(\*\*[^*]+\*\*)', cell_text)
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        run = p.add_run(part[2:-2])
                        run.bold = True
                    else:
                        p.add_run(part)

                p.style = doc.styles['Normal']
                for run in p.runs:
                    run.font.size = Pt(9)

                # Header row styling
                if i == 0:
                    shading = cell._element.get_or_add_tcPr()
                    shading_elem = shading.makeelement(
                        qn('w:shd'),
                        {qn('w:fill'): '192858', qn('w:val'): 'clear'}
                    )
                    shading.append(shading_elem)
                    for run in p.runs:
                        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                        run.bold = True

    # Auto-fit
    table.autofit = True


def apply_inline_formatting(paragraph, text: str) -> None:
    """Apply bold and italic markdown formatting to a paragraph."""
    # Split by bold markers
    parts = re.split(r'(\*\*[^*]+\*\*)', text)

    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            inner = part[2:-2]
            # Check for italic within bold
            italic_parts = re.split(r'(\*[^*]+\*)', inner)
            for ip in italic_parts:
                if ip.startswith('*') and ip.endswith('*'):
                    run = paragraph.add_run(ip[1:-1])
                    run.bold = True
                    run.italic = True
                else:
                    run = paragraph.add_run(ip)
                    run.bold = True
        else:
            # Check for standalone italic
            italic_parts = re.split(r'(\*[^*]+\*)', part)
            for ip in italic_parts:
                if ip.startswith('*') and ip.endswith('*'):
                    run = paragraph.add_run(ip[1:-1])
                    run.italic = True
                else:
                    paragraph.add_run(ip)


def convert_md_to_docx(md_path: Path, docx_path: Path, guide_num: int) -> None:
    """Convert a markdown file to a formatted Word document."""
    with open(md_path, "r", encoding="utf-8") as f:
        content = f.read()

    doc = Document()
    setup_styles(doc)

    # Extract title from first heading
    title_match = re.search(r'^#\s+(.+)$', content, re.MULTILINE)
    title = title_match.group(1) if title_match else md_path.stem

    add_header_footer(doc, title, guide_num)

    lines = content.split('\n')
    i = 0
    in_table = False
    table_lines = []
    in_code_block = False
    code_lines = []

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Code blocks
        if stripped.startswith('```'):
            if in_code_block:
                # End code block
                code_text = '\n'.join(code_lines)
                p = doc.add_paragraph()
                run = p.add_run(code_text)
                run.font.name = 'Consolas'
                run.font.size = Pt(9)
                run.font.color.rgb = DARK_GRAY
                p.paragraph_format.left_indent = Cm(1)
                code_lines = []
                in_code_block = False
            else:
                in_code_block = True
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # Table detection
        if stripped.startswith('|') and '|' in stripped[1:]:
            if not in_table:
                in_table = True
                table_lines = []
            table_lines.append(stripped)
            i += 1
            continue
        elif in_table:
            # End of table
            rows = parse_markdown_table(table_lines)
            if rows:
                add_table(doc, rows)
            in_table = False
            table_lines = []
            # Don't increment — process current line normally

        # Empty line
        if not stripped:
            i += 1
            continue

        # Headings
        if stripped.startswith('#### '):
            doc.add_heading(stripped[5:], level=4)
        elif stripped.startswith('### '):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith('## '):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith('# '):
            doc.add_heading(stripped[2:], level=1)

        # Horizontal rule
        elif stripped == '---' or stripped == '***':
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run('─' * 50)
            run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
            run.font.size = Pt(8)

        # Bullet points
        elif stripped.startswith('- ') or stripped.startswith('* '):
            bullet_text = stripped[2:]
            p = doc.add_paragraph(style='List Bullet')
            apply_inline_formatting(p, bullet_text)

        # Numbered lists
        elif re.match(r'^\d+\.\s', stripped):
            list_text = re.sub(r'^\d+\.\s', '', stripped)
            p = doc.add_paragraph(style='List Number')
            apply_inline_formatting(p, list_text)

        # Checkbox items
        elif stripped.startswith('- [ ] ') or stripped.startswith('- [x] '):
            checked = stripped.startswith('- [x] ')
            check_text = stripped[6:]
            prefix = "☑ " if checked else "☐ "
            p = doc.add_paragraph(style='List Bullet')
            run = p.add_run(prefix)
            run.font.name = 'Segoe UI Symbol'
            apply_inline_formatting(p, check_text)

        # Regular paragraph
        else:
            p = doc.add_paragraph()
            apply_inline_formatting(p, stripped)

        i += 1

    # Handle any remaining table
    if in_table and table_lines:
        rows = parse_markdown_table(table_lines)
        if rows:
            add_table(doc, rows)

    # Save
    docx_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(docx_path))


def find_md_files(guide_folder: Path) -> list[Path]:
    """Find all markdown files in a guide folder."""
    md_files = []
    for subfolder in sorted(guide_folder.iterdir()):
        if subfolder.is_dir():
            for md_file in sorted(subfolder.glob("*.md")):
                md_files.append(md_file)
    return md_files


def main():
    parser = argparse.ArgumentParser(description="Convert guide markdown files to formatted Word documents")
    parser.add_argument("--guide", type=int, help="Guide number (1-18)")
    parser.add_argument("--all", action="store_true", help="Process all guides")
    parser.add_argument("--file", type=str, default=None,
                        help="Filter by filename keyword (e.g., 'master', 'assessment', 'narration')")
    parser.add_argument("--list", action="store_true", help="List available MD files without converting")
    parser.add_argument("--output-dir", type=str, default=None,
                        help="Custom output directory (default: same folder as .md file)")
    args = parser.parse_args()

    if not args.guide and not args.all:
        print("Error: Specify --guide N or --all")
        sys.exit(1)

    guides = range(1, 19) if args.all else [args.guide]

    total_converted = 0

    for guide_num in guides:
        if guide_num not in GUIDE_FOLDERS:
            continue

        guide_folder = BUILD_OUTPUT / GUIDE_FOLDERS[guide_num]
        if not guide_folder.exists():
            continue

        md_files = find_md_files(guide_folder)

        if args.file:
            md_files = [f for f in md_files if args.file.lower() in f.name.lower()]

        if not md_files:
            if not args.all:
                print(f"No markdown files found for Guide {guide_num:02d}")
            continue

        print(f"\nGuide {guide_num:02d}: {GUIDE_FOLDERS[guide_num]}")
        print("-" * 50)

        if args.list:
            for md_file in md_files:
                rel_path = md_file.relative_to(guide_folder)
                print(f"  {rel_path}")
            continue

        for md_file in md_files:
            docx_name = md_file.stem + ".docx"

            if args.output_dir:
                docx_path = Path(args.output_dir) / docx_name
            else:
                docx_path = md_file.parent / docx_name

            try:
                convert_md_to_docx(md_file, docx_path, guide_num)
                rel_md = md_file.relative_to(guide_folder)
                print(f"  ✓ {rel_md} → {docx_name}")
                total_converted += 1
            except Exception as e:
                print(f"  ✗ {md_file.name}: {e}")

    if not args.list:
        print(f"\nDone! Converted {total_converted} files to Word.")


if __name__ == "__main__":
    main()
