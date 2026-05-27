#!/usr/bin/env python3
"""
Master Storyboard DOCX Generator — UHN Accessibility First Series
Guide 02: Perceptions, Attitudes, and Barriers (v2 — 21 slides)

Generates a fully-formatted Word document matching the v0.5.1 storyboard
reference format: Part 1 (review) + Part 2 (designer toolkit), 3-column tables.

Changes from v1 (23 slides):
  - Removed Slide 2.6 (Intersectionality) — folded into core concept
  - Removed Slide 2.13 (Scenario 4: Mental Health Handoff) — 3 scenarios sufficient
  - Removed Slide 2.21 (Decision Tree) — 3 scenarios + worked example enough
  - Added Slide 2.6: The Spectrum of Bias (unique core concept)
  - Added Slide 2.7: AiP Applied — Where Bias Shows Up (unique AiP)
  - Added Slide 2.8: Decision Path Worked Example (unique walkthrough)
  - Revised KC slides: 2 KC slides (3 questions total)
  - Revised reflection: structured self-assessment (not open text)
  - Revised podcast: Voices of Experience

Usage:
    python generate_storyboard_g02_v2.py

Output:
    05-build-output/02-Perceptions-Attitudes-and-Barriers/
        02-production/master-storyboard/MASTER-STORYBOARD-GUIDE-02.docx

Requirements:
    pip install python-docx
"""

import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor, Emu
    from docx.enum.section import WD_ORIENT
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("Error: 'python-docx' not installed. Run: pip install python-docx")
    sys.exit(1)

# ---------------------------------------------------------------------------
# Paths
# ---------------------------------------------------------------------------
BASE_DIR = Path(__file__).resolve().parent.parent.parent
LOGO_PATH = BASE_DIR / "02-branding-and-style" / "logos" / "uhn_logo_2.png"
OUTPUT_DIR = (
    BASE_DIR
    / "05-build-output"
    / "02-Perceptions-Attitudes-and-Barriers"
    / "02-production"
    / "master-storyboard"
)
OUTPUT_PATH = OUTPUT_DIR / "MASTER-STORYBOARD-GUIDE-02.docx"

# ---------------------------------------------------------------------------
# Colour constants (hex strings — no leading #)
# ---------------------------------------------------------------------------
NAVY        = "192858"
COBALT      = "255CAA"
BODY        = "2B2B2B"
LIGHT_FILL  = "F5F5F5"
HEADER_FILL = "EDF0F7"
BORDER_CLR  = "CCCCCC"
WHITE       = "FFFFFF"
RED         = "C0233B"

# ---------------------------------------------------------------------------
# Helper: apply shading to a table cell
# ---------------------------------------------------------------------------
def shade_cell(cell, fill_hex: str):
    """Set the background fill of a table cell."""
    tc_pr = cell._element.get_or_add_tcPr()
    for existing in tc_pr.findall(qn("w:shd")):
        tc_pr.remove(existing)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  fill_hex.upper())
    tc_pr.append(shd)


def set_cell_margins(cell, top_pt=6, bottom_pt=6, left_pt=8, right_pt=8):
    """Set cell internal margins in points (converted to twips)."""
    tc_pr = cell._element.get_or_add_tcPr()
    mar = OxmlElement("w:tcMar")
    for side, val in [("top", top_pt), ("bottom", bottom_pt),
                      ("left", left_pt), ("right", right_pt)]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"), str(int(val * 20)))
        el.set(qn("w:type"), "dxa")
        mar.append(el)
    for existing in tc_pr.findall(qn("w:tcMar")):
        tc_pr.remove(existing)
    tc_pr.append(mar)


def set_table_borders(table, color_hex: str = BORDER_CLR, size_pt: int = 4):
    """Apply thin uniform borders to every cell in a table."""
    border_size = str(size_pt * 8)
    for row in table.rows:
        for cell in row.cells:
            tc_pr = cell._element.get_or_add_tcPr()
            tc_borders = OxmlElement("w:tcBorders")
            for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
                border = OxmlElement(f"w:{side}")
                border.set(qn("w:val"),   "single")
                border.set(qn("w:sz"),    border_size)
                border.set(qn("w:space"), "0")
                border.set(qn("w:color"), color_hex.upper())
                tc_borders.append(border)
            for existing in tc_pr.findall(qn("w:tcBorders")):
                tc_pr.remove(existing)
            tc_pr.append(tc_borders)


def set_col_widths(table, widths_inches: list):
    """Set column widths for a table (list of floats in inches)."""
    for i, row in enumerate(table.rows):
        for j, cell in enumerate(row.cells):
            if j < len(widths_inches):
                tc = cell._element
                tc_pr = tc.get_or_add_tcPr()
                tc_w = OxmlElement("w:tcW")
                tc_w.set(qn("w:w"),    str(int(widths_inches[j] * 1440)))
                tc_w.set(qn("w:type"), "dxa")
                for existing in tc_pr.findall(qn("w:tcW")):
                    tc_pr.remove(existing)
                tc_pr.append(tc_w)


def rgb(hex_str: str):
    h = hex_str.lstrip("#")
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


def add_run(para, text: str, bold=False, italic=False, size_pt=10,
            color_hex=BODY, font_name="Arial"):
    run = para.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.size       = Pt(size_pt)
    run.font.color.rgb  = rgb(color_hex)
    run.font.name       = font_name
    return run


def cell_para(cell, text: str, bold=False, italic=False, size_pt=10,
              color_hex=BODY, font_name="Arial", align=WD_ALIGN_PARAGRAPH.LEFT):
    """Clear a cell and add a single paragraph with one run."""
    for p in cell.paragraphs[1:]:
        p._element.getparent().remove(p._element)
    para = cell.paragraphs[0]
    para.clear()
    para.alignment = align
    add_run(para, text, bold=bold, italic=italic, size_pt=size_pt,
            color_hex=color_hex, font_name=font_name)
    return para


def cell_para_multiline(cell, lines: list, size_pt=10, color_hex=BODY,
                        font_name="Arial"):
    """Write multiple paragraphs into a cell. Each item in lines is either
    a plain str or a dict with keys: text, bold, italic, size_pt, color_hex."""
    for p in cell.paragraphs[1:]:
        p._element.getparent().remove(p._element)
    first = True
    for line in lines:
        if first:
            para = cell.paragraphs[0]
            para.clear()
            first = False
        else:
            para = cell.add_paragraph()
        if isinstance(line, dict):
            para.alignment = line.get("align", WD_ALIGN_PARAGRAPH.LEFT)
            add_run(para, line.get("text", ""),
                    bold=line.get("bold", False),
                    italic=line.get("italic", False),
                    size_pt=line.get("size_pt", size_pt),
                    color_hex=line.get("color_hex", color_hex),
                    font_name=line.get("font_name", font_name))
        else:
            add_run(para, str(line), size_pt=size_pt, color_hex=color_hex,
                    font_name=font_name)


# ---------------------------------------------------------------------------
# Document setup
# ---------------------------------------------------------------------------
def setup_document() -> Document:
    doc = Document()
    # Landscape 11 x 8.5 inches
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width  = Inches(11)
    section.page_height = Inches(8.5)
    section.left_margin   = Inches(0.7)
    section.right_margin  = Inches(0.7)
    section.top_margin    = Inches(0.55)
    section.bottom_margin = Inches(0.5)

    # Default paragraph style
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(10)
    style.font.color.rgb = rgb(BODY)

    return doc


# ---------------------------------------------------------------------------
# Reusable block builders
# ---------------------------------------------------------------------------
def add_section_heading(doc, number: str, title: str):
    """Add a numbered section heading (e.g. '1.1 Course Information')."""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(12)
    para.paragraph_format.space_after  = Pt(4)
    add_run(para, f"{number}  {title}", bold=True, size_pt=13,
            color_hex=NAVY, font_name="Arial Black")


def add_part_banner(doc, text: str):
    """Full-width navy banner for PART 1 / PART 2."""
    doc.add_paragraph()  # spacer
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = tbl.rows[0].cells[0]
    shade_cell(cell, NAVY)
    set_cell_margins(cell, top_pt=8, bottom_pt=8, left_pt=12, right_pt=12)
    cell_para(cell, text, bold=True, size_pt=14,
              color_hex=WHITE, font_name="Arial Black",
              align=WD_ALIGN_PARAGRAPH.CENTER)
    set_table_borders(tbl, color_hex=NAVY)
    doc.add_paragraph()  # spacer after


def add_two_col_table(doc, rows_data: list,
                      col_widths=(2.5, 6.8),
                      header_row: list = None):
    """
    Build a 2-column info table.
    rows_data: list of (label, value) tuples.
    header_row: optional [left_header, right_header] for tables with a header row.
    """
    num_rows = len(rows_data) + (1 if header_row else 0)
    tbl = doc.add_table(rows=num_rows, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    row_idx = 0
    if header_row:
        hrow = tbl.rows[row_idx].cells
        shade_cell(hrow[0], NAVY)
        shade_cell(hrow[1], NAVY)
        cell_para(hrow[0], header_row[0], bold=True, size_pt=10,
                  color_hex=WHITE, font_name="Arial Black")
        cell_para(hrow[1], header_row[1], bold=True, size_pt=10,
                  color_hex=WHITE, font_name="Arial Black")
        set_cell_margins(hrow[0])
        set_cell_margins(hrow[1])
        row_idx += 1

    for label, value in rows_data:
        cells = tbl.rows[row_idx].cells
        shade_cell(cells[0], HEADER_FILL)
        shade_cell(cells[1], WHITE)
        cell_para(cells[0], label, bold=True, size_pt=10, color_hex=NAVY)
        if "\n" in str(value):
            lines = []
            for line in str(value).split("\n"):
                lines.append({"text": line, "size_pt": 10, "color_hex": BODY})
            cell_para_multiline(cells[1], lines)
        else:
            cell_para(cells[1], str(value), size_pt=10, color_hex=BODY)
        set_cell_margins(cells[0])
        set_cell_margins(cells[1])
        row_idx += 1

    set_col_widths(tbl, list(col_widths))
    set_table_borders(tbl)
    return tbl


def add_four_col_table(doc, rows_data: list, col_widths=(1.8, 3.5, 1.2, 1.8),
                       headers=None):
    """Build a 4-column table (used for summative assessments)."""
    num_rows = len(rows_data) + (1 if headers else 0)
    tbl = doc.add_table(rows=num_rows, cols=4)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    row_idx = 0
    if headers:
        hcells = tbl.rows[0].cells
        for i, h in enumerate(headers):
            shade_cell(hcells[i], NAVY)
            cell_para(hcells[i], h, bold=True, size_pt=10,
                      color_hex=WHITE, font_name="Arial Black")
            set_cell_margins(hcells[i])
        row_idx += 1

    for r in rows_data:
        cells = tbl.rows[row_idx].cells
        for i, val in enumerate(r):
            shade_cell(cells[i], LIGHT_FILL if i == 0 else WHITE)
            cell_para(cells[i], str(val), bold=(i == 0), size_pt=10,
                      color_hex=NAVY if i == 0 else BODY)
            set_cell_margins(cells[i])
        row_idx += 1

    set_col_widths(tbl, list(col_widths))
    set_table_borders(tbl)
    return tbl


def add_five_col_table(doc, rows_data: list,
                       col_widths=(0.45, 0.65, 2.8, 2.1, 3.3),
                       headers=None):
    """Build a 5-column table (Master Screen Schedule)."""
    num_rows = len(rows_data) + (1 if headers else 0)
    tbl = doc.add_table(rows=num_rows, cols=5)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    row_idx = 0
    if headers:
        hcells = tbl.rows[0].cells
        for i, h in enumerate(headers):
            shade_cell(hcells[i], NAVY)
            cell_para(hcells[i], h, bold=True, size_pt=10,
                      color_hex=WHITE, font_name="Arial Black")
            set_cell_margins(hcells[i])
        row_idx += 1

    for r in rows_data:
        cells = tbl.rows[row_idx].cells
        for i, val in enumerate(r):
            is_first = (i == 0)
            shade_cell(cells[i], HEADER_FILL if is_first else WHITE)
            cell_para(cells[i], str(val), bold=is_first, size_pt=10,
                      color_hex=NAVY if is_first else BODY)
            set_cell_margins(cells[i])
        row_idx += 1

    set_col_widths(tbl, list(col_widths))
    set_table_borders(tbl)
    return tbl


# ---------------------------------------------------------------------------
# Merge cells helper (for banner rows in Part 2 tables)
# ---------------------------------------------------------------------------
def merge_row(table, row_idx: int, start_col: int, end_col: int):
    """Merge cells from start_col to end_col (inclusive) in a given row."""
    row = table.rows[row_idx]
    merged = row.cells[start_col]
    for c in range(start_col + 1, end_col + 1):
        merged = merged.merge(row.cells[c])
    return merged


# ---------------------------------------------------------------------------
# Part 2: Three-column screen table
# ---------------------------------------------------------------------------
def add_screen_table(doc, screen_num: str, screen_title: str,
                     step: str, activities: str, design_guide: str):
    """
    Build the 3-column designer-toolkit table for one screen.
    Layout:
      Row 0: merged header — "SCREEN {num}" | "Topic: {title}"
      Row 1: column headers — Step | Activities | Design Guide
      Row 2: content
    """
    tbl = doc.add_table(rows=3, cols=3)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    # --- Row 0: screen banner ---
    r0 = tbl.rows[0].cells
    left_cell = r0[0].merge(r0[1])
    shade_cell(left_cell, NAVY)
    set_cell_margins(left_cell, top_pt=6, bottom_pt=6, left_pt=8, right_pt=8)
    cell_para(left_cell, f"SCREEN {screen_num}  (Online \u00b7 Self-Paced)",
              bold=True, size_pt=10, color_hex=WHITE, font_name="Arial Black")

    shade_cell(r0[2], NAVY)
    set_cell_margins(r0[2], top_pt=6, bottom_pt=6, left_pt=8, right_pt=8)
    cell_para(r0[2], f"Topic: {screen_title}",
              bold=True, size_pt=10, color_hex=WHITE, font_name="Arial Black")

    # --- Row 1: column headers ---
    r1 = tbl.rows[1].cells
    headers = ["Step", "Activities", "Design Guide"]
    for i, h in enumerate(headers):
        shade_cell(r1[i], "EDF0F7")
        cell_para(r1[i], h, bold=True, size_pt=10, color_hex=NAVY,
                  font_name="Arial Black")
        set_cell_margins(r1[i])

    # --- Row 2: content ---
    r2 = tbl.rows[2].cells

    # Step column
    shade_cell(r2[0], LIGHT_FILL)
    cell_para(r2[0], step, bold=True, size_pt=10, color_hex=NAVY)
    set_cell_margins(r2[0])

    # Activities column
    shade_cell(r2[1], WHITE)
    _fill_cell_lines(r2[1], activities, size_pt=10, color_hex=BODY)
    set_cell_margins(r2[1])

    # Design Guide column
    shade_cell(r2[2], LIGHT_FILL)
    _fill_cell_lines(r2[2], design_guide, size_pt=9, color_hex=BODY)
    set_cell_margins(r2[2])

    set_col_widths(tbl, [1.5, 4.5, 3.3])
    set_table_borders(tbl)
    doc.add_paragraph()  # spacer between screen tables


def _fill_cell_lines(cell, text: str, size_pt=10, color_hex=BODY):
    """Write text into a cell, splitting on newlines into separate paragraphs."""
    lines = str(text).split("\n")
    first = True
    for line in lines:
        if first:
            para = cell.paragraphs[0]
            para.clear()
            first = False
        else:
            para = cell.add_paragraph()
        add_run(para, line.strip(), size_pt=size_pt, color_hex=color_hex)


# ---------------------------------------------------------------------------
# COVER SECTION
# ---------------------------------------------------------------------------
def build_cover(doc):
    if LOGO_PATH.exists():
        doc.add_picture(str(LOGO_PATH), width=Inches(2.0))
    else:
        p = doc.add_paragraph("[UHN LOGO]")
        p.runs[0].font.size = Pt(10)
        p.runs[0].font.color.rgb = rgb(COBALT)

    doc.add_paragraph()  # spacer

    # Series name
    series_para = doc.add_paragraph()
    add_run(series_para, "Accessibility First eLearning Series",
            bold=True, size_pt=12, color_hex=COBALT)

    # Title
    title_para = doc.add_paragraph()
    add_run(title_para,
            "Guide 02 \u2014 Perceptions, Attitudes, and Barriers",
            bold=True, size_pt=22, color_hex=NAVY, font_name="Arial Black")

    # Subtitle
    sub_para = doc.add_paragraph()
    add_run(sub_para,
            "Understanding How Attitudes and Assumptions Create Barriers in Healthcare",
            italic=True, size_pt=12, color_hex=NAVY)

    doc.add_paragraph()  # spacer


# ---------------------------------------------------------------------------
# METADATA TABLE
# ---------------------------------------------------------------------------
def build_metadata_table(doc):
    rows = [
        ("Audience",         "All UHN clinical & administrative staff"),
        ("Modality",         "Online \u00b7 Self-paced \u00b7 HTML5 SCORM 1.2 to UHN MyLearning (SumTotal)"),
        ("Duration",         "~15-20 minutes (20 slides)"),
        ("Standard",         "WCAG 2.1 AA \u00b7 plain language \u00b7 trauma-informed"),
        ("Author",           "Yi Jin (Instructional Designer)"),
        ("Content owner",    "Jacqueline Silvera \u2014 IDEAA, People & Culture, UHN"),
        ("Document version", "v2.0 Revised Storyboard \u2014 2026-05-25"),
    ]
    add_two_col_table(doc, rows, col_widths=(2.0, 7.3))
    doc.add_paragraph()


# ---------------------------------------------------------------------------
# HOW TO READ
# ---------------------------------------------------------------------------
def build_how_to_read(doc):
    p = doc.add_paragraph()
    add_run(p,
            "How to read this storyboard  ",
            bold=True, size_pt=10, color_hex=NAVY, font_name="Arial Black")
    p.paragraph_format.space_after = Pt(2)

    body = doc.add_paragraph()
    add_run(body,
            "This document has two parts. "
            "Part 1 (Storyboard for Review) is for SME and stakeholder review \u2014 "
            "it contains the course information, learning outcomes, assessments, "
            "and a complete screen schedule. "
            "Part 2 (Designer Toolkit) is for the instructional designer and developer \u2014 "
            "it contains per-screen 3-column tables with Step, Activities, and Design Guide details "
            "that mirror the deployed SCORM course exactly. "
            "Review Part 1 for content accuracy and sign-off. "
            "Part 2 is a production reference only.",
            italic=True, size_pt=9, color_hex=BODY)
    body.paragraph_format.space_after = Pt(8)


# ---------------------------------------------------------------------------
# PART 1 — STORYBOARD FOR REVIEW
# ---------------------------------------------------------------------------

def build_1_1(doc):
    add_section_heading(doc, "1.1", "Course Information")
    rows = [
        ("Program Name",        "UHN Accessibility First eLearning Series"),
        ("Course Title",        "Perceptions, Attitudes, and Barriers"),
        ("Course Code",         "AF-02"),
        ("Course Hours",        "~0.30 hours / 15-20 minutes (online, self-paced)"),
        ("Pre-requisite(s)",    "Guide 01 \u2014 Foundations of Disability, Inclusion and Accessible Design"),
        ("Course Description",
         "This guide examines how personal perceptions and attitudes \u2014 including implicit and explicit "
         "bias, microaggressions, and systemic assumptions \u2014 create barriers in healthcare settings. "
         "Learners explore the spectrum of bias, apply the Accessibility in Practice model to bias "
         "specifically, walk through a worked example using the Accessibility Decision Path, and "
         "practise dignity-centred responses through three branching scenarios. Includes structured "
         "self-assessment and action planning. Framed around UHN\u2019s Accessibility First principles "
         "and the OHRC duty to accommodate. Builds on the foundation established in Guide 01."),
        ("Course Rationale",
         "Technical and environmental barriers often receive the most attention in accessibility "
         "training, but research consistently shows that attitudinal barriers are among the most "
         "persistent and hardest to address. This guide equips staff to recognize bias in themselves "
         "and their workplace culture, apply practical self-reflection habits, and move from "
         "compliance-focused thinking to genuine, proactive inclusion \u2014 with a direct connection "
         "to the OHRC and UHN IDEAA framework."),
        ("Learning Materials",
         "All readings, audio, and podcast embedded inside the HTML5 SCORM module. "
         "Source: Accessibility First Guide Series 1-18 (Draft, April 2026), Guide 2."),
        ("Accrediting Standards",
         "AODA \u2014 Integrated Accessibility Standards Regulation; "
         "Ontario Human Rights Code (Part I, Disability) \u2014 Duty to Accommodate; "
         "Mental Health Commission of Canada Anti-Stigma Resources; "
         "Statistics Canada Canadian Survey on Disability (2022); "
         "UHN IDEAA Strategic Framework; WCAG 2.1 AA."),
        ("Version / Last Update",
         "v2.0 Revised Storyboard \u2014 2026-05-25 (Yi Jin, Instructional Designer)"),
    ]
    add_two_col_table(doc, rows)
    doc.add_paragraph()


def build_1_2(doc):
    add_section_heading(doc, "1.2", "Course Development Information")
    rows = [
        ("Delivery Modality",   "Online \u00b7 Self-Paced \u00b7 HTML5 SCORM 1.2 to UHN MyLearning (SumTotal LMS)"),
        ("Course Developer(s)", "Yi Jin (Instructional Designer)"),
        ("SME(s)",              "Jacqueline Silvera (IDEAA, People & Culture, UHN); CAMH language review \u2014 TBC; Indigenous content review \u2014 TBC"),
        ("ID(s)",               "Yi Jin"),
        ("Existing Course",     "None \u2014 new build."),
    ]
    add_two_col_table(doc, rows)
    doc.add_paragraph()


def build_1_3(doc):
    add_section_heading(doc, "1.3", "Course Learning Outcomes (CLOs)")
    rows = [
        ("CLO 1",
         "Define attitudinal barriers and explain how implicit and explicit bias affect "
         "accessibility and equity in healthcare."),
        ("CLO 2",
         "Recognize how assumptions about ability, culture, mental health, and identity "
         "influence patient care and employee interactions at UHN."),
        ("CLO 3",
         "Apply reflective practice strategies to identify and mitigate personal bias "
         "in everyday workplace interactions."),
        ("CLO 4",
         "Demonstrate proactive, dignity-centred responses that move accessibility practice "
         "beyond compliance."),
    ]
    add_two_col_table(doc, rows, header_row=["CLO", "Description"])
    doc.add_paragraph()


def build_1_4(doc):
    add_section_heading(doc, "1.4", "Summative Assessments")
    rows = [
        ("Knowledge Check 1\n(Q1 + Q2)",
         "2 MCQ questions on one slide: (1) definition of implicit bias, (2) identify "
         "attitudinal barrier in scenario. 2 attempts each.",
         "~22%", "Screen 2.12"),
        ("Knowledge Check 2\n(Q3)",
         "1 scenario-based MCQ \u2014 responding to bias in practice. 2 attempts.",
         "~11%", "Screen 2.13"),
        ("Branching Scenarios",
         "3 scenarios: assumption about ability, implicit bias in triage, "
         "language and cultural assumptions. Select-then-submit with 3-choice "
         "branching and debrief feedback.",
         "~47%", "Screens 2.9\u20132.11"),
        ("My Action Plan (MAP)",
         "Stop / Start / Continue commitments. Stop and Start fields required to proceed.",
         "~20%", "Screen 2.16"),
    ]
    add_four_col_table(
        doc, rows,
        col_widths=(2.0, 4.0, 0.8, 2.5),
        headers=["Title of Assessment", "Brief Description", "Weight", "Position"],
    )
    doc.add_paragraph()


def build_1_5(doc):
    add_section_heading(doc, "1.5", "Master Screen Schedule")
    schedule = [
        ("1",  "2.1",  "Welcome and Course Purpose",                      "Title / Cover",                "Hero image, Begin button"),
        ("2",  "2.2",  "Learning Objectives",                             "Objectives",                   "4 CLO cards (2x2 grid)"),
        ("3",  "2.3",  "Impact: The Stat",                                "Content + Statistic",          "3x unmet healthcare needs, navy stat block"),
        ("4",  "2.4",  "Impact: Attitudinal Barriers",                    "Content + Image",              "How attitudes create invisible barriers"),
        ("5",  "2.5",  "Impact: Stigma and Avoidance",                    "Content + Image",              "Stigma causes patients to delay/avoid care"),
        ("6",  "2.6",  "Core Concept: The Spectrum of Bias",              "Tabbed panels + definitions",  "Implicit/explicit bias, microaggressions, healthcare examples"),
        ("7",  "2.7",  "AiP Applied: Where Bias Shows Up",               "4-quadrant cards (bias)",      "Awareness, Communication, Environment, Response \u2014 bias-specific"),
        ("8",  "2.8",  "Decision Path Worked Example",                    "5-step narrative walkthrough",  "Speech disability scenario, bias-specific steps"),
        ("9",  "2.9",  "Scenario 1: Assumption About Ability",            "Branching (3-choice)",         "Patient spoken to through companion, graded"),
        ("10", "2.10", "Scenario 2: Implicit Bias in Triage",             "Branching (3-choice)",         "Mental health dx + chest pain, graded"),
        ("11", "2.11", "Scenario 3: Language and Cultural Assumptions",   "Branching (3-choice)",         "Communication device user, graded"),
        ("12", "2.12", "Knowledge Check 1 (Q1 + Q2)",                     "2 MCQ (4 options each)",       "Implicit bias definition + attitudinal barrier, 2 attempts each"),
        ("13", "2.13", "Knowledge Check 2 (Q3)",                          "Scenario MCQ (4 options)",     "Responding to bias in practice, 2 attempts"),
        ("14", "2.14", "Five Ways to Check Your Assumptions",             "5-card grid",                  "Pause, Ask, Reflect, Seek feedback, Report"),
        ("15", "2.15", "Reflection: Self-Assessment",                     "Structured rating (1-5)",      "4 dimensions, private \u2014 not submitted"),
        ("16", "2.16", "MAP Action Planning",                             "3-field form + PDF download",  "Stop/Start/Continue, MAP-Template-Guide-02.pdf"),
        ("17", "2.17", "Key Takeaways",                                   "Summary (4 cards)",            "4 takeaway cards, navy/cobalt"),
        ("18", "2.18", "Podcast: Voices of Experience",                   "Podcast player + transcript",  "Episode 02 \u00b7 Voices of Experience \u00b7 ~15-18 min"),
        ("19", "2.19", "Series Progress Map",                             "Progress map (3 stages)",      "2 of 4 Foundations complete; Stage 2+3 locked"),
        ("20", "2.20", "Resources & Completion",                          "Completion + badge",           "Badge, resource links, Guide 03 next"),
    ]
    add_five_col_table(
        doc, schedule,
        col_widths=[0.45, 0.65, 2.8, 2.1, 3.3],
        headers=["Slide", "Screen", "Title", "Type", "Notes"],
    )
    doc.add_paragraph()


def build_1_6(doc):
    add_section_heading(doc, "1.6", "Version Control & Sign-off")

    version_rows = [
        ("Version", "Date",       "Author",   "Change Summary"),
        ("v0.1",    "2026-05-25", "Yi Jin",   "Initial draft storyboard (23 slides)"),
        ("v2.0",    "2026-05-25", "Yi Jin",   "Revised storyboard \u2014 21 slides, unique core concept/AiP/DP slides, 3 scenarios"),
        ("v0.5",    "TBD",        "Yi Jin",   "SME review draft"),
        ("v0.5.1",  "TBD",        "Yi Jin / J. Silvera", "Post-SME revisions"),
        ("v1.0",    "TBD",        "Yi Jin",   "Deployed storyboard \u2014 final"),
    ]
    tbl = doc.add_table(rows=len(version_rows), cols=4)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, row in enumerate(version_rows):
        cells = tbl.rows[i].cells
        for j, val in enumerate(row):
            is_header = (i == 0)
            shade_cell(cells[j], NAVY if is_header else (HEADER_FILL if j == 0 else WHITE))
            cell_para(cells[j], val, bold=is_header, size_pt=10,
                      color_hex=WHITE if is_header else (NAVY if j == 0 else BODY),
                      font_name="Arial Black" if is_header else "Arial")
            set_cell_margins(cells[j])
    set_col_widths(tbl, [0.8, 1.2, 2.5, 4.8])
    set_table_borders(tbl)

    doc.add_paragraph()

    # Sign-off box
    signoff_tbl = doc.add_table(rows=1, cols=1)
    cell = signoff_tbl.rows[0].cells[0]
    shade_cell(cell, LIGHT_FILL)
    set_cell_margins(cell, top_pt=10, bottom_pt=10, left_pt=12, right_pt=12)

    lines = [
        {"text": "SIGN-OFF", "bold": True, "size_pt": 11, "color_hex": NAVY,
         "font_name": "Arial Black"},
        {"text": ""},
        {"text": "SME Review \u2014 Jacqueline Silvera (IDEAA, People & Culture, UHN)"},
        {"text": "Reviewed by: ________________________    Date: ____________"},
        {"text": ""},
        {"text": "CAMH Language Review \u2014 [Reviewer TBC]"},
        {"text": "Reviewed by: ________________________    Date: ____________"},
        {"text": ""},
        {"text": "Indigenous Content Review \u2014 [Reviewer TBC]"},
        {"text": "Reviewed by: ________________________    Date: ____________"},
        {"text": ""},
        {"text": "ID Review \u2014 Yi Jin (Instructional Designer)"},
        {"text": "Approved by: ________________________    Date: ____________"},
        {"text": ""},
        {"text": "Note: Signature on this document constitutes approval of all content, "
                 "scenarios, assessment questions, and learning outcomes as accurate, "
                 "appropriate for the intended audience, and aligned with UHN policy.",
         "italic": True, "size_pt": 9, "color_hex": BODY},
    ]
    cell_para_multiline(cell, lines)
    set_table_borders(signoff_tbl, color_hex=COBALT)
    doc.add_paragraph()


# ---------------------------------------------------------------------------
# PART 2 — DESIGNER TOOLKIT: Per-screen data (21 slides)
# ---------------------------------------------------------------------------

SCREEN_DATA = [
    # -----------------------------------------------------------------------
    # SLIDE 1 — Screen 2.1: Welcome and Course Purpose
    # -----------------------------------------------------------------------
    {
        "num":   "2.1",
        "title": "Welcome and Course Purpose",
        "step":  "Hook",
        "activities": (
            "TITLE: Perceptions, Attitudes, and Barriers\n"
            "SUBTITLE: Guide 02 of 18 \u00b7 Foundations Stage\n"
            "BODY: Over the next 15 to 20 minutes, you will examine how our own perceptions and attitudes shape the care and environments we create.\n"
            "METADATA: Duration 15-20 min \u00b7 Pre-requisite: Guide 01 \u00b7 Version v2.0 \u00b7 2026\n"
            "BUTTON: Begin\n"
            "VO: voiceover_2.1.mp3\n"
            "NARRATION (~148 words): Welcome back to the Accessibility First series. This is Guide 2: Perceptions, Attitudes, and Barriers. In Guide 1, you explored what disability means, the models that shape how we think about it, and a practical framework for everyday decisions. This guide goes deeper. It asks a more personal question: how do our own perceptions and attitudes shape the care and environments we create? We all carry assumptions. Some we are aware of. Many we are not. And in a healthcare setting, those assumptions \u2014 even well-intentioned ones \u2014 can create barriers as real as any locked door or missing ramp. Over the next 15 to 20 minutes, you will examine those assumptions, explore where they come from, and build practical strategies to move from compliance toward genuine inclusion. This guide is relevant for everyone at UHN. Let us get started."
        ),
        "design_guide": (
            "Format: Cover slide \u00b7 static layout with autoplay narration.\n"
            "Interaction: Begin button navigates to Screen 2.2.\n"
            "Accessibility: 1.4.3 contrast >= 4.5:1 \u00b7 2.1.1 keyboard reachable \u00b7 CC ON by default \u00b7 Begin button Tab + Enter.\n"
            "Visual: Warm, inclusive hero photo \u2014 diverse UHN staff and patients in a bright hospital corridor.\n"
            "Image: g02-hero-welcome-01.png\n"
            "Alt: Diverse UHN healthcare workers and patients in conversation in a bright hospital corridor\n"
            "Captions: voiceover_2.1.srt"
        ),
    },
    # -----------------------------------------------------------------------
    # SLIDE 2 — Screen 2.2: Learning Objectives
    # -----------------------------------------------------------------------
    {
        "num":   "2.2",
        "title": "Learning Objectives",
        "step":  "Outcomes",
        "activities": (
            "HEADER: What You Will Learn\n"
            "LEDE: By the end of this guide, you will be able to do four things:\n"
            "CLO 1: Understand attitudinal barriers \u2014 define attitudinal barriers and explain how implicit and explicit bias affect accessibility and equity in healthcare.\n"
            "CLO 2: Recognize assumptions in practice \u2014 identify how assumptions about ability, culture, mental health, and identity influence patient care and employee interactions at UHN.\n"
            "CLO 3: Apply reflective practice \u2014 use self-reflection strategies to recognize and mitigate personal bias in everyday interactions.\n"
            "CLO 4: Act with dignity beyond compliance \u2014 demonstrate proactive, respectful responses that move accessibility practice beyond meeting minimum standards.\n"
            "VO: voiceover_2.2.mp3\n"
            "NARRATION (~122 words): By the end of this guide, you will be able to do four things. First, you will be able to define attitudinal barriers and explain the difference between implicit and explicit bias \u2014 and why both matter in healthcare. Second, you will be able to recognize how assumptions about ability, culture, mental health, and identity show up in patient care and workplace interactions. Third, you will be able to apply reflective practice strategies to identify your own assumptions and adjust your behaviour. And fourth, you will be able to demonstrate responses that go beyond compliance \u2014 responses grounded in dignity, curiosity, and genuine inclusion. These four objectives build on Guide 1. Each one connects directly to a scenario you will encounter later in this guide."
        ),
        "design_guide": (
            "Format: Static objective cards (4 cards, 2x2 grid, navy top border, large number + title + description).\n"
            "Interaction: None \u2014 orientation slide. Continue always available.\n"
            "Accessibility: Ordered list markup (ol/li). Each icon has alt text. WCAG 2.1 AA contrast.\n"
            "Captions: voiceover_2.2.srt"
        ),
    },
    # -----------------------------------------------------------------------
    # SLIDE 3 — Screen 2.3: Impact: The Stat
    # -----------------------------------------------------------------------
    {
        "num":   "2.3",
        "title": "Impact: The Stat",
        "step":  "Context",
        "activities": (
            "HEADER: The Invisible Barrier\n"
            "STAT: People with disabilities are 3x more likely to report unmet healthcare needs \u2014 Statistics Canada, 2022\n"
            "PULL QUOTE: Often, the barrier is not a ramp or a doorway. It is an attitude.\n"
            "BODY: Attitudinal barriers are invisible to those who hold them \u2014 and deeply felt by those who experience them.\n"
            "VO: voiceover_2.3.mp3\n"
            "NARRATION (~90 words): Here is a number worth sitting with. According to Statistics Canada, people with disabilities are nearly three times more likely than people without disabilities to report unmet healthcare needs. Not because services do not exist. But because the experience of seeking care \u2014 the assumptions made, the language used, the way people are treated \u2014 creates barriers that make people give up before they even get the help they need. The most persistent barriers in healthcare are not physical. They are attitudinal. And the good news is that attitudes can change. That is what this guide is about.\n"
            "REFS: Statistics Canada (2023). Canadian Survey on Disability, 2022.\n"
            "SME NOTE: HIGH \u2014 verify stat accuracy with SME"
        ),
        "design_guide": (
            "Format: Large statistic prominent, centred. Animated counter. Clean white background with navy stat block.\n"
            "Accessibility: Stat coded as text, not image. Animation respects prefers-reduced-motion.\n"
            "Image: g02-stats-attitudinal-01.png\n"
            "Alt: Statistic: people with disabilities are 3x more likely to report unmet healthcare needs\n"
            "Captions: voiceover_2.3.srt"
        ),
    },
    # -----------------------------------------------------------------------
    # SLIDE 4 — Screen 2.4: Impact: Attitudinal Barriers
    # -----------------------------------------------------------------------
    {
        "num":   "2.4",
        "title": "Impact: Attitudinal Barriers",
        "step":  "Impact",
        "activities": (
            "HEADER: When Attitudes Become Barriers\n"
            "BODY: Attitudinal barriers occur when assumptions, stereotypes, or lack of awareness lead to exclusion \u2014 even without intent.\n"
            "BODY 2: In healthcare, these barriers affect diagnosis, treatment decisions, and whether patients feel safe returning. An employee who speaks over a patient with a cognitive disability. A clinician who assumes a patient with a mental health history cannot make sound decisions. A colleague who speaks to a patient\u2019s support person instead of the patient directly.\n"
            "KEY TAKEAWAY: None of these require bad intent. They require awareness. And awareness is something we can build.\n"
            "VO: voiceover_2.4.mp3\n"
            "NARRATION (~85 words): Attitudinal barriers are among the most common \u2014 and the hardest to see \u2014 in any workplace. In healthcare, they can affect who receives equitable treatment, who gets taken seriously, and who feels safe enough to come back. An employee who speaks over a patient with a cognitive disability. A clinician who assumes a patient with a mental health history cannot make sound decisions. A colleague who speaks to a patient\u2019s support person instead of the patient directly. None of these require bad intent. They require awareness. And awareness is something we can build.\n"
            "SME NOTE: Medium"
        ),
        "design_guide": (
            "Format: Two-column layout. Left: photo placeholder (540px height). Right: red accent bar + title + body text + key takeaway box.\n"
            "Accessibility: Red accent bar for visual emphasis. WCAG 2.1 AA contrast.\n"
            "Image: g02-impact-attitudinal-01.png\n"
            "Alt: Healthcare provider and patient in a tense or disconnected interaction suggesting a communication gap\n"
            "Captions: voiceover_2.4.srt"
        ),
    },
    # -----------------------------------------------------------------------
    # SLIDE 5 — Screen 2.5: Impact: Stigma and Avoidance
    # -----------------------------------------------------------------------
    {
        "num":   "2.5",
        "title": "Impact: Stigma and Avoidance",
        "step":  "Impact",
        "activities": (
            "HEADER: The Weight of Stigma\n"
            "BODY: Stigma is one of the most powerful attitudinal barriers in healthcare. It causes patients to delay or avoid care entirely \u2014 sometimes for years.\n"
            "STAT: Mental health stigma in healthcare settings delays help-seeking by an average of 11 years. \u2014 Mental Health Commission of Canada\n"
            "BODY 2: For employees with disabilities, fear of disclosure can limit career growth and access to supports. Patients may not disclose conditions relevant to their care. The result is the same: people suffer in silence because the environment feels unsafe.\n"
            "INTERSECTIONALITY NOTE: A person\u2019s experience of stigma is shaped by race, gender, language, income, age, and culture \u2014 not disability alone. Indigenous Peoples face compounded barriers rooted in colonial history, systemic racism, and cultural disconnection from care.\n"
            "KEY TAKEAWAY: When we examine our own assumptions about mental health, ability, and identity \u2014 our language, our reactions, our expectations \u2014 we begin to break down stigma from the inside out.\n"
            "VO: voiceover_2.5.mp3\n"
            "NARRATION (~100 words): Stigma is one of the most powerful attitudinal barriers we know of. The Mental Health Commission of Canada reports that fear of stigma causes many people to delay seeking help for mental health conditions for over a decade. At UHN, this affects both patients and employees. And the experience of stigma does not exist in isolation. A person\u2019s experience is shaped by race, gender, language, income, and culture. Indigenous Peoples face compounded barriers rooted in colonial history, systemic racism, and geographic isolation. When we examine our own assumptions \u2014 our language, our reactions, our expectations \u2014 we begin to break down stigma from the inside out.\n"
            "SME NOTE: HIGH \u2014 CAMH language review required; Indigenous content requires SME sign-off"
        ),
        "design_guide": (
            "Format: Two-column layout. Left: photo placeholder (540px height). Right: cobalt accent bar + title + body text + intersectionality callout box (earth tones) + key takeaway.\n"
            "Note: This slide folds in the intersectionality content from the removed Screen 2.6 of the 23-slide version.\n"
            "Accessibility: Cobalt accent bar for visual emphasis. WCAG 2.1 AA contrast.\n"
            "Image: g02-impact-stigma-01.png\n"
            "Alt: Person sitting alone in a waiting area, looking anxious or withdrawn\n"
            "Captions: voiceover_2.5.srt"
        ),
    },
    # -----------------------------------------------------------------------
    # SLIDE 6 — Screen 2.6: Core Concept: The Spectrum of Bias (UNIQUE)
    # -----------------------------------------------------------------------
    {
        "num":   "2.6",
        "title": "Core Concept: The Spectrum of Bias",
        "step":  "Core Concept",
        "activities": (
            "HEADER: The Spectrum of Bias\n"
            "LEDE: Bias exists on a spectrum \u2014 from unconscious assumptions to deliberate discrimination. Understanding where bias lives helps you know where to intervene.\n"
            "\n"
            "SECTION 1 \u2014 IMPLICIT BIAS:\n"
            "Definition: An unconscious attitude or stereotype that affects judgement and behaviour without awareness.\n"
            "Healthcare example: A nurse automatically speaks louder when addressing a patient using a wheelchair, assuming hearing loss. A triage nurse unconsciously attributes a patient\u2019s chest pain to anxiety because the chart shows a mental health diagnosis.\n"
            "Impact: Patients feel patronized, trust erodes, clinical decisions are compromised.\n"
            "\n"
            "SECTION 2 \u2014 MICROAGGRESSIONS:\n"
            "Definition: Brief, everyday exchanges that communicate negative or demeaning messages to people based on their membership in a marginalized group \u2014 often unintentionally.\n"
            "Healthcare example: \u201cYou speak so well for someone with...\u201d \u201cYou don\u2019t look disabled.\u201d Asking a patient\u2019s companion about the patient\u2019s needs instead of asking the patient directly.\n"
            "Impact: Cumulative harm. Each instance is small; the pattern is exhausting and dehumanizing.\n"
            "\n"
            "SECTION 3 \u2014 EXPLICIT BIAS:\n"
            "Definition: A conscious attitude, belief, or stereotype that is deliberately held and acted upon.\n"
            "Healthcare example: A provider believes that patients with mental health histories are \u2018difficult\u2019 and delays their intake. A manager withholds a leadership opportunity from an employee because of a disclosed disability.\n"
            "Impact: Direct discrimination; creates unsafe care and work environments.\n"
            "\n"
            "KEY INSIGHT: Most bias in healthcare is implicit \u2014 well-meaning people acting on unconscious patterns. That is why awareness, not punishment, is the first step. The question is not whether you have bias. The question is whether you can recognize it and choose a different response.\n"
            "VO: voiceover_2.6.mp3\n"
            "NARRATION (~200 words): Bias exists on a spectrum \u2014 and understanding where on that spectrum a behaviour falls helps you know how to respond. At one end is implicit bias: the unconscious attitudes and stereotypes that shape our behaviour without our awareness. A triage nurse who unconsciously attributes a patient\u2019s chest pain to anxiety because the chart shows a mental health diagnosis. A receptionist who speaks to the companion instead of the patient. These are not acts of malice. They are patterns we have absorbed from culture, media, and experience. In the middle are microaggressions: brief, everyday exchanges that communicate demeaning messages, often unintentionally. \u2018You don\u2019t look disabled.\u2019 \u2018You speak so well for someone with...\u2019 Each instance seems small. The cumulative effect is exhausting. At the far end is explicit bias: conscious, deliberate attitudes acted upon. This is rarer in healthcare settings, but it exists. The key insight is this: most bias in healthcare is implicit. It comes from well-meaning people acting on unconscious patterns. That is why awareness \u2014 not punishment \u2014 is the first step. You cannot change what you cannot see. This guide will help you see it.\n"
            "SME NOTE: Medium"
        ),
        "design_guide": (
            "Format: Three-panel layout (spectrum visualization). Left: Implicit (cobalt). Centre: Microaggressions (chartreuse). Right: Explicit (red). Each panel expands on click to reveal definition, example, and impact.\n"
            "Interaction: Click each panel to expand. One panel open at a time. Keyboard-navigable.\n"
            "Accessibility: Panels keyboard-navigable. role=\"tablist\", aria-selected. Screen reader announces active panel.\n"
            "Visual: Horizontal spectrum bar at top connecting the three panels. Gradient from cobalt through chartreuse to red.\n"
            "Image: g02-spectrum-bias-01.png (flat vector infographic)\n"
            "Alt: Three-panel spectrum showing implicit bias, microaggressions, and explicit bias with definitions and healthcare examples\n"
            "Captions: voiceover_2.6.srt"
        ),
    },
    # -----------------------------------------------------------------------
    # SLIDE 7 — Screen 2.7: AiP Applied: Where Bias Shows Up (UNIQUE)
    # -----------------------------------------------------------------------
    {
        "num":   "2.7",
        "title": "AiP Applied: Where Bias Shows Up",
        "step":  "Framework",
        "activities": (
            "HEADER: Where Bias Shows Up: Accessibility in Practice\n"
            "LEDE: You learned the four AiP areas in Guide 01. Now let\u2019s apply them specifically to bias \u2014 because bias shows up differently in each area.\n"
            "\n"
            "CARD 1 \u2014 AWARENESS (Bias lens): What assumptions do I carry?\n"
            "  Am I making assumptions based on how this person looks, moves, or communicates?\n"
            "  Am I reacting differently to this patient than I did to the last one? Why?\n"
            "  Have I checked my own comfort level \u2014 or am I avoiding a situation because it feels unfamiliar?\n"
            "\n"
            "CARD 2 \u2014 COMMUNICATION (Bias lens): Am I speaking differently based on visible cues?\n"
            "  Am I speaking to the patient or to their companion?\n"
            "  Am I simplifying my language based on assumptions about cognitive ability?\n"
            "  Am I using labels in handoffs that transfer bias rather than clinical information? (e.g. \u2018difficult patient\u2019)\n"
            "\n"
            "CARD 3 \u2014 ENVIRONMENT (Bias lens): Who feels welcome here?\n"
            "  Does our space signal that all bodies, communication styles, and identities are expected and valued?\n"
            "  Are our intake processes designed for the \u2018standard\u2019 patient \u2014 or for the full range of people who actually come through our doors?\n"
            "  Would a person with an invisible disability feel safe disclosing their needs here?\n"
            "\n"
            "CARD 4 \u2014 RESPONSE (Bias lens): Whose needs get prioritized?\n"
            "  When I respond to an accessibility need, am I treating it with the same urgency as other needs?\n"
            "  Am I offering help \u2014 or imposing it?\n"
            "  Am I reporting attitudinal barriers upstream, or just working around them?\n"
            "\n"
            "SERIES FRAMEWORK: These questions are specific to bias and perception. In other guides, the AiP model is applied to different accessibility topics.\n"
            "VO: voiceover_2.7.mp3\n"
            "NARRATION (~180 words): In Guide 1, you learned the Accessibility in Practice model \u2014 four interconnected areas: Awareness, Communication, Environment, and Response. Now let\u2019s apply those areas specifically to bias. Awareness: what assumptions do I carry? Am I reacting differently to this patient than the last? Am I avoiding a situation because it feels unfamiliar? These are the questions that make implicit bias visible. Communication: am I speaking differently based on visible cues? Am I talking to the patient or to their companion? Am I using labels like \u2018difficult patient\u2019 that transfer bias rather than clinical information? Environment: who feels welcome here? Are our processes designed for the \u2018standard\u2019 patient or for the full range of people who actually come through our doors? Would someone with an invisible disability feel safe disclosing their needs? And Response: whose needs get prioritized? Am I offering help or imposing it? Am I reporting attitudinal barriers upstream or just working around them? These are not generic questions. They are specific to bias. And they will help you catch the patterns that matter most.\n"
            "SME NOTE: Low"
        ),
        "design_guide": (
            "Format: 4-quadrant card grid. Each card: number, heading, bias-specific question, 3 sub-questions. Color-coded top bars (Navy, Cobalt, Chartreuse, Red).\n"
            "Interaction: Cards are tabbable (tabindex=\"0\"). Sub-questions reveal on focus/click.\n"
            "Accessibility: Tab order: header \u2192 card 1 \u2192 card 2 \u2192 card 3 \u2192 card 4. Cards have tabindex=\"0\".\n"
            "Visual: Same layout as Guide 01 Screen 1.5 but with bias-specific content. NOT the generic AiP text.\n"
            "Captions: voiceover_2.7.srt"
        ),
    },
    # -----------------------------------------------------------------------
    # SLIDE 8 — Screen 2.8: Decision Path Worked Example (UNIQUE)
    # -----------------------------------------------------------------------
    {
        "num":   "2.8",
        "title": "Decision Path Worked Example",
        "step":  "Tool",
        "activities": (
            "HEADER: The Decision Path in Action: A Bias Scenario\n"
            "LEDE: In Guide 01, you learned the 5-step Accessibility Decision Path. Here is what it looks like when the barrier is an attitude \u2014 not a ramp.\n"
            "\n"
            "SCENARIO: You are explaining a treatment plan to a patient who has a speech disability. Partway through, you realize you have been directing your explanation to the patient\u2019s support worker \u2014 not to the patient. You assumed, without thinking about it, that because the patient has difficulty speaking, they might not understand complex medical information.\n"
            "\n"
            "STEP 1 \u2014 PAUSE & ASSESS: Notice the assumption.\n"
            "You catch yourself. You realize you shifted your eye contact and your language to the support worker. The patient has not said they do not understand. You assumed it. This is the moment of recognition \u2014 the most important step. Without it, the bias continues unexamined.\n"
            "\n"
            "STEP 2 \u2014 LISTEN & ASK: Ask the patient how they prefer to receive information.\n"
            "You turn to the patient directly: \u201cI want to make sure I\u2019m explaining this in the way that works best for you. Would you like me to continue, or would you prefer written information, or is there another format that works better?\u201d You listen to what they say \u2014 not what you expected them to say.\n"
            "\n"
            "STEP 3 \u2014 APPLY: Use the AiP Communication area.\n"
            "The patient indicates they understand spoken information but would also like a written summary. You adjust your approach: you continue speaking directly to the patient, with the support worker present but not as the primary audience. You prepare a written summary to leave with them.\n"
            "\n"
            "STEP 4 \u2014 ADAPT: Adjust based on the patient\u2019s response.\n"
            "Midway through, the patient uses a communication device to ask a clarifying question. You pause, give them time, and respond to their question directly. You do not rush. You do not look to the support worker for interpretation. You adapt your pace to the person in front of you.\n"
            "\n"
            "STEP 5 \u2014 SEEK SUPPORT: Consult accessibility resources if needed.\n"
            "After the interaction, you note in the patient\u2019s chart their communication preferences for future visits. You also flag to your team lead that the unit could benefit from a brief refresher on communication with patients who use AAC (augmentative and alternative communication) devices.\n"
            "\n"
            "DEBRIEF: This is what the Decision Path looks like when the barrier is bias. The steps are the same \u2014 but the application is personal. It starts with catching yourself, and it ends with making the system better for the next person.\n"
            "VO: voiceover_2.8.mp3\n"
            "NARRATION (~220 words): In Guide 1, you learned the five-step Accessibility Decision Path: Pause and Assess, Listen and Ask, Apply, Adapt, and Seek Support. That tool works for physical barriers, communication barriers, and environmental barriers. But what does it look like when the barrier is an attitude \u2014 not a ramp? Here is a worked example. You are explaining a treatment plan to a patient with a speech disability. Partway through, you realize you have been directing your explanation to the support worker, not to the patient. You assumed \u2014 without thinking about it \u2014 that because the patient has difficulty speaking, they might not understand complex medical information. Step one: Pause and Assess. You catch yourself. You notice the assumption. This is the most important step. Step two: Listen and Ask. You turn to the patient directly and ask how they prefer to receive information. Step three: Apply. You use the Communication area of the AiP model \u2014 you continue speaking to the patient and prepare a written summary. Step four: Adapt. The patient uses a communication device to ask a question. You give them time. You do not rush. Step five: Seek Support. You note their preferences in the chart and suggest a team refresher on AAC devices. The steps are the same. The application is personal.\n"
            "SME NOTE: Medium"
        ),
        "design_guide": (
            "Format: Narrative walkthrough with 5-step stepper rail. Each step reveals the scenario detail. NOT the generic stepper from Guide 01 \u2014 this is a bias-specific narrative.\n"
            "Interaction: Step-through progression. Click each step to reveal narrative content. Cannot skip ahead.\n"
            "Accessibility: Stepper keyboard-navigable. Each step announced by screen reader. Focus indicator on active step.\n"
            "Visual: Left rail with 5 numbered steps (same styling as Guide 01 Screen 1.6). Right panel shows narrative text for each step. Scenario setup shown above stepper.\n"
            "Image: g02-decision-path-worked-01.png\n"
            "Alt: Healthcare worker redirecting attention from support worker to patient with communication device\n"
            "Captions: voiceover_2.8.srt"
        ),
    },
    # -----------------------------------------------------------------------
    # SLIDE 9 — Screen 2.9: Scenario 1: Assumption About Ability
    # -----------------------------------------------------------------------
    {
        "num":   "2.9",
        "title": "Scenario 1: Assumption About Ability",
        "step":  "Scenario",
        "activities": (
            "HEADER: Scenario 1: Assumption About Ability\n"
            "SCENARIO: A patient with a visible physical disability \u2014 Ms. Okeke uses a power wheelchair \u2014 arrives at the outpatient clinic for a follow-up appointment. The intake clerk addresses all questions to Ms. Okeke\u2019s companion, a friend who drove her to the appointment. The clerk makes eye contact with the companion, directs questions like \u201cDoes she need any accommodations?\u201d and \u201cWhat medications is she on?\u201d to the friend. Ms. Okeke is sitting right there.\n"
            "QUESTION: What should the intake clerk do?\n"
            "\n"
            "CHOICE A (Best): The clerk catches themselves, turns to Ms. Okeke directly, and says: \u201cI\u2019m sorry \u2014 I should be asking you directly. How would you like me to proceed? Is there anything you need to make this visit work well for you?\u201d The clerk redirects all subsequent questions to Ms. Okeke.\n"
            "OUTCOME A: Ms. Okeke feels respected and seen. Her autonomy is affirmed. The companion\u2019s role is clarified as a support person, not a spokesperson. The clerk models inclusive behaviour for other staff who may observe the interaction.\n"
            "\n"
            "CHOICE B (Acceptable): The clerk continues to include both Ms. Okeke and the companion in conversation, but does not explicitly redirect. Questions go to \u201cboth of you.\u201d\n"
            "OUTCOME B: Better than excluding the patient, but the implicit assumption \u2014 that the companion is needed for communication \u2014 is not challenged. Ms. Okeke may feel her autonomy is partially respected but not fully affirmed.\n"
            "\n"
            "CHOICE C (Poor): The clerk continues directing all questions to the companion, reasoning that it is more efficient and the companion seems to know the answers.\n"
            "OUTCOME C: Ms. Okeke feels invisible and patronized. Her dignity and autonomy are disregarded. This reinforces the attitudinal barrier: the assumption that physical disability means cognitive or communicative limitation. This is the spread effect in action.\n"
            "\n"
            "DEBRIEF: Always address the patient directly. A visible physical disability does not indicate cognitive or communication limitations. When you catch yourself making an assumption, the most powerful thing you can do is pause, correct course, and ask the patient how they prefer to be supported.\n"
            "PRINCIPLE: People First & Dignity \u00b7 Independence & Autonomy\n"
            "VO: voiceover_2.9.mp3 \u00b7 SME NOTE: HIGH"
        ),
        "design_guide": (
            "Format: Branching choice \u00b7 3 options \u00b7 graded. Select-then-submit with feedback overlay.\n"
            "Interaction: All choices keyboard-selectable. Consequence screens have back navigation. No auto-advance.\n"
            "Image: g02-scenario-ability-01.png\n"
            "Alt: Patient in power wheelchair at clinic intake desk with companion standing beside them\n"
            "Captions: voiceover_2.9.srt\n"
            "Maps to: CLO 2, CLO 4"
        ),
    },
    # -----------------------------------------------------------------------
    # SLIDE 10 — Screen 2.10: Scenario 2: Implicit Bias in Triage
    # -----------------------------------------------------------------------
    {
        "num":   "2.10",
        "title": "Scenario 2: Implicit Bias in Triage",
        "step":  "Scenario",
        "activities": (
            "HEADER: Scenario 2: Implicit Bias in Triage\n"
            "SCENARIO: Mr. Davies, a 42-year-old patient, presents at the emergency department with chest pain and shortness of breath. His chart shows a history of generalized anxiety disorder and panic attacks. The triage nurse reviews the chart, notes the mental health history, and begins to consider whether the symptoms might be anxiety-related rather than cardiac. The nurse is experienced and well-intentioned \u2014 but the mental health diagnosis has shifted their clinical lens.\n"
            "QUESTION: What should the triage nurse do?\n"
            "\n"
            "CHOICE A (Best): The nurse assesses the presenting symptoms \u2014 chest pain and shortness of breath \u2014 on their own clinical merits, using standard cardiac triage protocols. The mental health history is noted but does not change the urgency of the assessment. The nurse says: \u201cI can see you\u2019re in distress. Let\u2019s get you assessed right away.\u201d\n"
            "OUTCOME A: Mr. Davies receives timely, appropriate care. His symptoms are not filtered through his mental health diagnosis. The nurse demonstrates that clinical assessment should be based on presenting symptoms, not diagnostic labels. If the symptoms do turn out to be anxiety-related, that conclusion comes from evidence \u2014 not assumption.\n"
            "\n"
            "CHOICE B (Acceptable): The nurse proceeds with standard triage but mentions to a colleague: \u201cHe has anxiety on his chart, so it could be a panic attack.\u201d The assessment continues, but the framing has shifted.\n"
            "OUTCOME B: The patient is assessed, but the handoff language introduces bias. A colleague who hears \u201cit could be a panic attack\u201d may unconsciously deprioritize the patient. This is how implicit bias spreads through teams \u2014 not through bad intent, but through casual framing.\n"
            "\n"
            "CHOICE C (Poor): The nurse attributes the symptoms to anxiety and places Mr. Davies in a lower-priority queue, reasoning that panic attacks, while distressing, are not life-threatening.\n"
            "OUTCOME C: Mr. Davies waits longer for assessment. If the symptoms are cardiac, this delay could be dangerous. This is diagnostic overshadowing \u2014 a well-documented form of implicit bias where physical symptoms are attributed to a pre-existing mental health condition. It is one of the most serious patient safety risks associated with bias in healthcare.\n"
            "\n"
            "DEBRIEF: Diagnostic overshadowing is a form of implicit bias where physical symptoms are dismissed or deprioritized because of a mental health diagnosis. It is a patient safety issue, not just an equity issue. Assess presenting symptoms on their clinical merits. The chart informs \u2014 it does not diagnose.\n"
            "PRINCIPLE: Clinical Safety \u00b7 Equity \u00b7 People First\n"
            "VO: voiceover_2.10.mp3 \u00b7 SME NOTE: HIGH \u2014 CAMH review + clinical accuracy"
        ),
        "design_guide": (
            "Format: Branching choice \u00b7 3 options \u00b7 graded. Same interaction pattern as Screen 2.9.\n"
            "Image: g02-scenario-triage-01.png\n"
            "Alt: Triage nurse reviewing patient chart in emergency department, patient visible in background holding their chest\n"
            "Captions: voiceover_2.10.srt\n"
            "Maps to: CLO 1, CLO 2, CLO 3"
        ),
    },
    # -----------------------------------------------------------------------
    # SLIDE 11 — Screen 2.11: Scenario 3: Language and Cultural Assumptions
    # -----------------------------------------------------------------------
    {
        "num":   "2.11",
        "title": "Scenario 3: Language and Cultural Assumptions",
        "step":  "Scenario",
        "activities": (
            "HEADER: Scenario 3: Language and Cultural Assumptions\n"
            "SCENARIO: Mr. Park, a 35-year-old patient, uses an augmentative and alternative communication (AAC) device \u2014 a tablet with a speech-generating app \u2014 to communicate. He arrives for a routine lab appointment. The phlebotomist notices the AAC device and begins to simplify their language significantly: speaking very slowly, using single-word instructions (\u201cSit.\u201d \u201cArm.\u201d \u201cStay.\u201d), and avoiding eye contact. The phlebotomist assumes that because Mr. Park uses a communication device, he has a cognitive impairment.\n"
            "QUESTION: What should the phlebotomist do?\n"
            "\n"
            "CHOICE A (Best): The phlebotomist speaks to Mr. Park using normal adult language, at a natural pace. They ask: \u201cIs there anything you need from me to make this appointment go smoothly?\u201d They give Mr. Park time to respond using his device and maintain eye contact while waiting.\n"
            "OUTCOME A: Mr. Park is treated as an adult with full cognitive capacity. The phlebotomist\u2019s approach communicates respect and patience. The appointment proceeds smoothly because the patient feels safe and engaged.\n"
            "\n"
            "CHOICE B (Acceptable): The phlebotomist uses simplified language but maintains eye contact and a respectful tone. They do not ask about Mr. Park\u2019s communication preferences but complete the appointment without incident.\n"
            "OUTCOME B: The tone is respectful, but the assumption remains unchallenged. Mr. Park may not feel harmed in this interaction, but the underlying pattern \u2014 equating AAC use with cognitive impairment \u2014 persists and will affect other patients.\n"
            "\n"
            "CHOICE C (Poor): The phlebotomist calls a colleague to \u201chelp\u201d with the patient, explaining that \u201che can\u2019t really communicate.\u201d Both staff members proceed to speak about Mr. Park in the third person while he is present.\n"
            "OUTCOME C: Mr. Park is dehumanized. His communication ability is dismissed. Being spoken about in the third person while present is a common and deeply harmful experience for people who use AAC devices. The phlebotomist\u2019s assumption has become the barrier.\n"
            "\n"
            "DEBRIEF: Using a communication device does not indicate cognitive impairment. AAC users have the same range of cognitive abilities as anyone else. The appropriate response is to ask how the person prefers to communicate, give them time, and treat them as the expert on their own needs.\n"
            "PRINCIPLE: Communication \u00b7 People First & Dignity \u00b7 Inclusion\n"
            "VO: voiceover_2.11.mp3 \u00b7 SME NOTE: HIGH \u2014 AAC clinical review"
        ),
        "design_guide": (
            "Format: Branching choice \u00b7 3 options \u00b7 graded. Same interaction pattern as Screen 2.9.\n"
            "Image: g02-scenario-language-01.png\n"
            "Alt: Patient using tablet-based communication device during a healthcare appointment\n"
            "Captions: voiceover_2.11.srt\n"
            "Maps to: CLO 2, CLO 3, CLO 4"
        ),
    },
    # -----------------------------------------------------------------------
    # SLIDE 12 — Screen 2.12: Knowledge Check 1 (Q1 + Q2)
    # -----------------------------------------------------------------------
    {
        "num":   "2.12",
        "title": "Knowledge Check 1 (Q1 + Q2)",
        "step":  "Assessment",
        "activities": (
            "HEADER: Knowledge Check 1\n"
            "LEDE: Two questions. Select the best answer for each.\n"
            "\n"
            "QUESTION 1: Which of the following BEST describes implicit bias?\n"
            "A: A deliberate belief that a group of people is inferior (INCORRECT)\n"
            "B: An unconscious attitude that shapes behaviour without awareness (CORRECT)\n"
            "C: A policy that intentionally excludes people with disabilities (INCORRECT)\n"
            "D: A formal complaint about discriminatory treatment (INCORRECT)\n"
            "FEEDBACK CORRECT (B): That is right. Implicit bias is unconscious \u2014 we may not even be aware it is influencing our actions. That is what makes reflective practice so important.\n"
            "FEEDBACK INCORRECT (A): This describes explicit bias \u2014 a deliberate, conscious attitude. Implicit bias is different: it operates below the level of awareness.\n"
            "FEEDBACK INCORRECT (C): This describes systemic or structural discrimination in policy. Implicit bias is an individual-level unconscious attitude.\n"
            "FEEDBACK INCORRECT (D): A formal complaint is a response to discrimination, not a type of bias.\n"
            "\n"
            "QUESTION 2: A triage nurse unconsciously deprioritizes a patient\u2019s chest pain because their chart shows a history of anxiety disorder. This is an example of:\n"
            "A: Paternalism \u2014 deciding what the patient needs without asking (INCORRECT)\n"
            "B: Explicit bias \u2014 deliberately discriminating against the patient (INCORRECT)\n"
            "C: Diagnostic overshadowing \u2014 physical symptoms attributed to a mental health condition (CORRECT)\n"
            "D: Microaggression \u2014 a brief demeaning exchange (INCORRECT)\n"
            "FEEDBACK CORRECT (C): Exactly. Diagnostic overshadowing occurs when physical symptoms are dismissed or deprioritized because of a pre-existing mental health diagnosis. It is a patient safety issue rooted in implicit bias.\n"
            "FEEDBACK INCORRECT (A): Paternalism involves making decisions for someone without their input. This scenario is about a clinical assessment being distorted by bias, not a care decision made on the patient\u2019s behalf.\n"
            "FEEDBACK INCORRECT (B): The nurse is not acting deliberately \u2014 this is an unconscious pattern, which is the hallmark of implicit rather than explicit bias.\n"
            "FEEDBACK INCORRECT (D): A microaggression is a verbal or behavioural exchange. This scenario describes a clinical decision-making pattern, not an interpersonal exchange.\n"
            "\n"
            "2 attempts allowed per question. After 2 wrong attempts, correct answer revealed. Submit button locks after final attempt.\n"
            "VO: voiceover_2.12.mp3 (brief intro only)"
        ),
        "design_guide": (
            "Format: 2 MCQ on one slide \u00b7 4 options each \u00b7 2 attempts per question \u00b7 submit-then-lock.\n"
            "Interaction: Each question has its own submit button. Questions are independent (answering Q1 does not affect Q2). Feedback panel appears below each question.\n"
            "Accessibility: Tab order: Q1 options \u2192 Q1 submit \u2192 Q2 options \u2192 Q2 submit.\n"
            "Captions: voiceover_2.12.srt\n"
            "Maps to: CLO 1, CLO 2"
        ),
    },
    # -----------------------------------------------------------------------
    # SLIDE 13 — Screen 2.13: Knowledge Check 2 (Q3)
    # -----------------------------------------------------------------------
    {
        "num":   "2.13",
        "title": "Knowledge Check 2 (Q3)",
        "step":  "Assessment",
        "activities": (
            "HEADER: Knowledge Check 2\n"
            "LEDE: One scenario-based question.\n"
            "\n"
            "QUESTION: During a shift handoff, a colleague describes a patient who uses an AAC device as \u201cdifficult to communicate with\u201d and suggests \u201cjust getting the info from the family.\u201d You have not yet met the patient. What is the BEST response?\n"
            "A: Accept the colleague\u2019s assessment and direct your questions to the family when you enter the room. (INCORRECT)\n"
            "B: Thank the colleague for the handoff, then meet the patient yourself. Introduce yourself, ask how they prefer to communicate, and give them time to respond using their device. (CORRECT)\n"
            "C: Report the colleague to a supervisor for using biased language. (INCORRECT)\n"
            "D: Enter the room and try to communicate with the patient, but if it takes too long, switch to the family member. (INCORRECT)\n"
            "FEEDBACK CORRECT (B): That is right. Each patient deserves a fresh interaction based on your own assessment, not on a label from a colleague. Asking the patient how they prefer to communicate is the most respectful and effective approach.\n"
            "FEEDBACK INCORRECT (A): Accepting a biased framing without question transfers that bias to your own care. The patient has not had the opportunity to communicate with you directly.\n"
            "FEEDBACK INCORRECT (C): While biased handoff language is a concern, the immediate priority is the patient\u2019s care. You can address the language concern separately. Reporting without first providing good care does not help the patient right now.\n"
            "FEEDBACK INCORRECT (D): Setting a time limit on communication undermines the patient\u2019s right to participate in their own care. AAC communication takes time \u2014 that time is part of equitable care.\n"
            "\n"
            "2 attempts allowed. After 2 wrong attempts, correct answer revealed.\n"
            "VO: voiceover_2.13.mp3"
        ),
        "design_guide": (
            "Format: Scenario-based MCQ \u00b7 4 options \u00b7 2 attempts \u00b7 submit-then-lock.\n"
            "Captions: voiceover_2.13.srt\n"
            "Maps to: CLO 3, CLO 4"
        ),
    },
    # -----------------------------------------------------------------------
    # SLIDE 14 — Screen 2.14: Five Ways to Check Your Assumptions
    # -----------------------------------------------------------------------
    {
        "num":   "2.14",
        "title": "Five Ways to Check Your Assumptions",
        "step":  "Practice",
        "activities": (
            "HEADER: Five Ways to Check Your Assumptions\n"
            "LEDE: Five practices you can start using in your next shift.\n"
            "\n"
            "TIP 1 \u2014 PAUSE BEFORE ACTING: Before you assist, redirect, or simplify \u2014 pause. Ask yourself: am I responding to what this person has told me, or to what I assumed about them? One second of awareness can change an entire interaction.\n"
            "\n"
            "TIP 2 \u2014 ASK, DON\u2019T ASSUME: Replace assumptions with questions. \u201cHow would you like me to support you?\u201d \u201cWhat format works best for you?\u201d \u201cIs there anything I can do differently?\u201d These questions communicate respect and give the person control.\n"
            "\n"
            "TIP 3 \u2014 REFLECT ON PATTERNS: At the end of a shift, take 30 seconds to review: Did I treat all patients the same way? Did I change my approach based on visible cues rather than expressed needs? Pattern recognition is a skill \u2014 build it intentionally.\n"
            "\n"
            "TIP 4 \u2014 SEEK FEEDBACK: Ask trusted colleagues to observe and give you honest feedback. Bias is easier to see from the outside. Create a culture where feedback is a gift, not a threat.\n"
            "\n"
            "TIP 5 \u2014 REPORT WHEN YOU SEE IT: When you notice attitudinal barriers \u2014 in yourself, in a process, in a colleague\u2019s language \u2014 flag it. Report it. Don\u2019t just work around it. Systemic change requires upstream action.\n"
            "\n"
            "VO: voiceover_2.14.mp3\n"
            "NARRATION (~180 words): Before we move to the self-assessment, here are five ways to check your assumptions \u2014 practices you can start using in your next shift. First: pause before acting. Before you assist, redirect, or simplify, ask yourself: am I responding to what this person has told me, or to what I assumed? One second of awareness can change an entire interaction. Second: ask, do not assume. Replace assumptions with questions. How would you like me to support you? What format works best? These questions communicate respect. Third: reflect on patterns. At the end of a shift, take 30 seconds to review. Did I treat all patients the same way? Did I change my approach based on visible cues rather than expressed needs? Fourth: seek feedback. Ask trusted colleagues to observe and give you honest feedback. Bias is easier to see from the outside. And fifth: report when you see it. When you notice attitudinal barriers, flag them. Systemic change requires upstream action.\n"
            "SME NOTE: Low"
        ),
        "design_guide": (
            "Format: 5-card static grid. Color-coded top bars (Navy, Cobalt, Chartreuse, Red, Lilac).\n"
            "Interaction: Cards are tabbable. All content visible without interaction.\n"
            "Accessibility: Cards are semantic list items. WCAG 2.1 AA contrast.\n"
            "Captions: voiceover_2.14.srt"
        ),
    },
    # -----------------------------------------------------------------------
    # SLIDE 15 — Screen 2.15: Reflection: Self-Assessment
    # -----------------------------------------------------------------------
    {
        "num":   "2.15",
        "title": "Reflection: Self-Assessment",
        "step":  "Reflection",
        "activities": (
            "HEADER: Reflection: How Confident Are You?\n"
            "LEDE: Rate your confidence from 1 (not at all confident) to 5 (very confident) on each of the following dimensions. This is private \u2014 it is not submitted or reviewed.\n"
            "\n"
            "DIMENSION 1 \u2014 RECOGNIZING IMPLICIT BIAS:\n"
            "I can recognize when my own implicit bias may be influencing my behaviour or clinical decisions.\n"
            "SCALE: 1 (Not at all confident) \u2014 2 \u2014 3 \u2014 4 \u2014 5 (Very confident)\n"
            "\n"
            "DIMENSION 2 \u2014 USING INCLUSIVE LANGUAGE:\n"
            "I consistently use language that respects each person\u2019s identity, agency, and communication preferences.\n"
            "SCALE: 1 (Not at all confident) \u2014 2 \u2014 3 \u2014 4 \u2014 5 (Very confident)\n"
            "\n"
            "DIMENSION 3 \u2014 CHALLENGING ASSUMPTIONS:\n"
            "When I notice an assumption forming \u2014 about a patient, a colleague, or a situation \u2014 I can pause, name it, and choose a different response.\n"
            "SCALE: 1 (Not at all confident) \u2014 2 \u2014 3 \u2014 4 \u2014 5 (Very confident)\n"
            "\n"
            "DIMENSION 4 \u2014 RESPONDING WITH DIGNITY:\n"
            "I consistently respond to accessibility needs with the same urgency, respect, and attention I would give any other patient need.\n"
            "SCALE: 1 (Not at all confident) \u2014 2 \u2014 3 \u2014 4 \u2014 5 (Very confident)\n"
            "\n"
            "PRIVATE NOTICE: This self-assessment is private. It is not submitted, scored, or reviewed. It is for your own awareness.\n"
            "BUTTON: SAVE MY REFLECTION (saves locally only)\n"
            "VO: voiceover_2.15.mp3\n"
            "NARRATION (~120 words): Take a moment to reflect on what you have learned. This is a private self-assessment. No one will see your answers. For each of the four dimensions, rate your current confidence from 1 to 5. How confident are you in recognizing when implicit bias may be influencing your decisions? How confident are you in using language that respects each person\u2019s identity and preferences? How confident are you in pausing when you notice an assumption forming and choosing a different response? And how confident are you in responding to accessibility needs with the same urgency you would give any other patient need? Be honest. There are no wrong answers. The purpose is awareness \u2014 and awareness is the first step toward change.\n"
            "SME NOTE: Low"
        ),
        "design_guide": (
            "Format: Structured self-assessment \u00b7 4 dimensions \u00b7 1-5 radio button scale per dimension \u00b7 ungraded.\n"
            "Interaction: Radio button selection (1-5) for each dimension. Save button stores to localStorage only. NOT an open text field.\n"
            "Visual: Clean layout. Each dimension in its own card with slider or radio buttons. Calm, reflective mood.\n"
            "Accessibility: Radio buttons keyboard-accessible. Labels associated with inputs. Tab order: D1 \u2192 D2 \u2192 D3 \u2192 D4 \u2192 Save.\n"
            "Image: g02-reflection-self-assessment-01.png\n"
            "Alt: Quiet, reflective image \u2014 healthcare worker in a thoughtful moment\n"
            "Captions: voiceover_2.15.srt"
        ),
    },
    # -----------------------------------------------------------------------
    # SLIDE 16 — Screen 2.16: MAP Action Planning (FIXED)
    # -----------------------------------------------------------------------
    {
        "num":   "2.16",
        "title": "MAP Action Planning",
        "step":  "Action",
        "activities": (
            "HEADER: My Action Plan (MAP)\n"
            "LEDE: Three prompts \u2014 Stop, Start, Continue. Completing at least the Stop and Start fields is required to proceed.\n"
            "FIELD 1 \u2014 STOP: One thing I will stop doing\n"
            "  Placeholder: \u2018I will stop making assumptions about what patients can do before asking.\u2019\n"
            "FIELD 2 \u2014 START: One thing I will start doing\n"
            "  Placeholder: \u2018I will start pausing before I speak to check whether I\u2019m responding to what the patient said or what I assumed.\u2019\n"
            "FIELD 3 \u2014 CONTINUE: One thing I will continue doing (optional)\n"
            "  Placeholder: \u2018I will continue asking patients directly how they prefer to be addressed and supported.\u2019\n"
            "DOWNLOAD: MAP-Template-Guide-02.pdf\n"
            "BUTTON: SAVE MY MAP\n"
            "VO: voiceover_2.16.mp3\n"
            "NARRATION (~168 words): You are almost done. Before we reach the final summary, it is time to make a personal commitment. This is your My Action Plan \u2014 the MAP. The MAP appears in every guide in this series, and it is always about translating what you have learned into one concrete behaviour change. In the fields below, complete three prompts. Stop: one thing you will stop doing \u2014 an assumption, a habit, a phrase. Start: one thing you will start doing \u2014 a question you will ask, a pause you will take, a practice you will adopt. And Continue: one thing you are already doing well that you want to keep doing. Take two to three minutes. Be specific and honest. The goal is not a grand declaration. It is a small, concrete action that you can take in your next shift. Your MAP is private. It will not be reviewed or assessed. It is yours.\n"
            "CLO: 4 \u00b7 SME NOTE: Low"
        ),
        "design_guide": (
            "Format: Form layout \u00b7 3 contenteditable inputs \u00b7 downloadable PDF. Lilac accent (MAP activity signature colour).\n"
            "Interaction: Completing at least Stop and Start fields required to proceed. Save stores to localStorage. Download generates branded PDF.\n"
            "Accessibility: All fields keyboard-accessible. Download button clearly labelled. PDF is accessible (tagged PDF).\n"
            "Download: MAP-Template-Guide-02.pdf\n"
            "Captions: voiceover_2.16.srt\n"
            "Maps to: Completion requirement \u00b7 CLO 4"
        ),
    },
    # -----------------------------------------------------------------------
    # SLIDE 17 — Screen 2.17: Key Takeaways (FIXED)
    # -----------------------------------------------------------------------
    {
        "num":   "2.17",
        "title": "Key Takeaways",
        "step":  "Summary",
        "activities": (
            "HEADER: Key Takeaways\n"
            "LEDE: Four things to remember from this guide.\n"
            "\n"
            "TAKEAWAY 1: Most barriers are attitudinal, not physical \u2014 The most persistent barriers in healthcare are not ramps or doorways. They are assumptions, stereotypes, and unconscious patterns that shape how we treat people.\n"
            "\n"
            "TAKEAWAY 2: Bias is usually implicit \u2014 awareness is the first step \u2014 Most bias in healthcare comes from well-meaning people acting on unconscious patterns. You cannot change what you cannot see. Awareness is not weakness \u2014 it is the beginning of change.\n"
            "\n"
            "TAKEAWAY 3: Reflective practice is a skill you can build \u2014 Catching your own assumptions, naming them, and choosing a different response is not a personality trait. It is a practice. And like any practice, it gets stronger with repetition.\n"
            "\n"
            "TAKEAWAY 4: Dignity means asking, not assuming \u2014 The single most powerful behaviour in accessible care is asking the person in front of you what they need \u2014 before you decide for them. Ask first. Listen. Adapt.\n"
            "\n"
            "VO: voiceover_2.17.mp3\n"
            "NARRATION (~160 words): Let us recap what you have learned. First: the most persistent barriers in healthcare are not physical. They are attitudinal \u2014 assumptions, stereotypes, and unconscious patterns that shape how we treat people. Second: most bias is implicit. It comes from well-meaning people acting on patterns they have not examined. Awareness is not a weakness. It is the beginning of change. Third: reflective practice is a skill, not a personality trait. Catching your own assumptions, naming them, and choosing a different response \u2014 this gets stronger with repetition. You do not need to be perfect. You need to be willing to notice. And fourth: dignity means asking, not assuming. The single most powerful behaviour in accessible care is asking the person in front of you what they need before you decide for them. These are the habits that build an accessible culture \u2014 one interaction at a time.\n"
            "SME NOTE: Low"
        ),
        "design_guide": (
            "Format: 4-card summary grid (2x2). Color-coded (Navy, Cobalt, Chartreuse, Red).\n"
            "Accessibility: Summary points are semantic list. WCAG 2.1 AA contrast.\n"
            "Captions: voiceover_2.17.srt"
        ),
    },
    # -----------------------------------------------------------------------
    # SLIDE 18 — Screen 2.18: Podcast: Voices of Experience
    # -----------------------------------------------------------------------
    {
        "num":   "2.18",
        "title": "Podcast: Voices of Experience",
        "step":  "Podcast",
        "activities": (
            "HEADER: Listen and Reflect\n"
            "EPISODE: Episode 02 of 18 \u00b7 Perceptions, Attitudes, and Barriers\n"
            "TITLE: Voices of Experience: Bias, Dignity, and the Patient Journey\n"
            "DURATION: ~15-18 min (full podcast)\n"
            "DESCRIPTION: Personal stories about bias in healthcare from patient advisors and staff who have experienced attitudinal barriers first-hand. Hear what it feels like to be spoken about instead of spoken to, to have your abilities assumed away, and to encounter dignity in unexpected moments.\n"
            "\n"
            "KEY LISTENING POINTS:\n"
            "  0:00 \u2014 Introduction: Why personal stories matter in accessibility training\n"
            "  3:15 \u2014 \u201cThey talked to my mother, not to me\u201d \u2014 a patient advisor describes being spoken through rather than spoken to, and the moment a staff member changed everything by simply asking\n"
            "  7:30 \u2014 The weight of \u201cYou don\u2019t look disabled\u201d \u2014 microaggressions and the invisible tax of having to prove your needs over and over\n"
            "  11:00 \u2014 \u201cHe has anxiety on his chart\u201d \u2014 a patient\u2019s experience of diagnostic overshadowing and how a bias in the chart followed them from visit to visit\n"
            "  14:30 \u2014 What dignity looks like in practice \u2014 three small moments that restored trust: a pause, a question, and eye contact\n"
            "\n"
            "REFLECTION PROMPT (after listening): Think of a moment when someone assumed something about you that was not true. How did it feel? How might that feeling inform how you approach your next patient interaction?\n"
            "AUDIO: [Podcast file \u2014 to be generated via NotebookLM]\n"
            "VO: voiceover_2.18.mp3 (intro narration only)"
        ),
        "design_guide": (
            "Format: Podcast player with audio controls, transcript toggle, key listening points sidebar.\n"
            "Interaction: Play/pause, seekable progress bar, clickable timestamps, full transcript panel.\n"
            "Visual: Dark navy background. Lilac accent for episode label.\n"
            "Accessibility: Captions always available (CC toggle). Full transcript expandable. Downloadable MP3.\n"
            "Captions: voiceover_2.18.srt"
        ),
    },
    # -----------------------------------------------------------------------
    # SLIDE 19 — Screen 2.19: Series Progress Map (FIXED)
    # -----------------------------------------------------------------------
    {
        "num":   "2.19",
        "title": "Series Progress Map",
        "step":  "Progress",
        "activities": (
            "HEADER: Accessibility First \u00b7 A Three-Stage Journey\n"
            "LEDE: 18 guides organized in three stages. Guides unlock by stage.\n"
            "STAGE 1 \u2014 Foundations (Guides 01-04): Required first. Guides 01-02 complete. Guides 03-04 upcoming.\n"
            "STAGE 2 \u2014 Understanding Disability Experiences (Guides 05-09): Builds on Foundations. Locked.\n"
            "STAGE 3 \u2014 Applied Practice & Specialized Contexts (Guides 10-18): Unlock after Stages 1+2. Locked.\n"
            "STATUS: You have completed 2 of 4 Foundation guides. Complete Guides 03 and 04 to unlock Stage 2.\n"
            "VO: voiceover_2.19.mp3\n"
            "SME NOTE: Low"
        ),
        "design_guide": (
            "Format: 3-column card grid. Locked stages have overlay with lock icon.\n"
            "Visual: Guide 01 complete (filled dot). Guide 02 active/highlighted. Guides 03-04 upcoming (open dot). Stages 2+3 locked.\n"
            "Color: Stage 1 Red, Stage 2 Cobalt, Stage 3 Chartreuse.\n"
            "Accessibility: Progress map keyboard-navigable. Screen reader announces progress state.\n"
            "Captions: voiceover_2.19.srt"
        ),
    },
    # -----------------------------------------------------------------------
    # SLIDE 20 — Screen 2.20: Resources & Completion (FIXED)
    # -----------------------------------------------------------------------
    {
        "num":   "2.20",
        "title": "Resources & Completion",
        "step":  "Completion",
        "activities": (
            "HEADER: Resources & Course Completion\n"
            "COMPLETION: You have completed Guide 02. You have earned the Accessibility First: Perceptions & Attitudes badge.\n"
            "RESOURCES:\n"
            "  AODA: ontario.ca/laws/statute/05a11\n"
            "  OHRC Human Rights Code: ohrc.on.ca\n"
            "  CAMH Anti-Stigma Resources: camh.ca/en/health-info/guides-and-publications\n"
            "  Mental Health Commission of Canada: mentalhealthcommission.ca\n"
            "  UHN IDEAA Office: [internal link]\n"
            "  UHN Accessibility Policy: [internal link]\n"
            "BADGE: Accessibility First: Perceptions & Attitudes\n"
            "NEXT: Guide 03 \u2014 Vision Disabilities: Inclusive Practices for Patients with Low or No Vision\n"
            "BUTTONS: RETRY QUIZ, ACTION PLAN (download MAP), EXIT COURSE\n"
            "VO: voiceover_2.20.mp3\n"
            "NARRATION (~150 words): Congratulations \u2014 you have completed Guide 2: Perceptions, Attitudes, and Barriers. You have explored the spectrum of bias, applied the Accessibility in Practice model to your own assumptions, walked through a worked example of the Decision Path, and practised three branching scenarios grounded in real healthcare situations. You have also completed a self-assessment and committed to a personal action plan. The habits you practise here \u2014 pausing before assuming, asking instead of deciding, reflecting on patterns, and reporting barriers upstream \u2014 these are not one-time lessons. They are ongoing practices. Below you will find links to resources including the OHRC Human Rights Code, CAMH anti-stigma materials, and UHN\u2019s IDEAA office. You have also earned the Accessibility First: Perceptions and Attitudes badge. In Guide 3, we will shift to vision disabilities. Well done. See you in Guide 3.\n"
            "SME NOTE: Low"
        ),
        "design_guide": (
            "Format: 2-column layout. Left: completion status + resources + next guide + action buttons. Right: badge container.\n"
            "Completion criteria: All 20 slides visited + quiz >= 80% + at least Stop and Start MAP fields completed.\n"
            "Accessibility: All links keyboard-accessible. Badge has alt text. SCORM completion status sent.\n"
            "Captions: voiceover_2.20.srt"
        ),
    },
]


# ---------------------------------------------------------------------------
# Part 2: intro paragraph
# ---------------------------------------------------------------------------
def build_part2_intro(doc):
    p = doc.add_paragraph()
    add_run(p,
            "Per-screen 3-column tables (Step \u00b7 Activities \u00b7 Design Guide). "
            "The Design Guide column captures the interaction format, accessibility specifications, "
            "and image references. "
            "This document mirrors the deployed SCORM course exactly.",
            italic=True, size_pt=9, color_hex=BODY)
    p.paragraph_format.space_after = Pt(8)


# ---------------------------------------------------------------------------
# MAIN BUILD
# ---------------------------------------------------------------------------
def build():
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    doc = setup_document()

    # ---- COVER ----
    build_cover(doc)
    build_metadata_table(doc)
    build_how_to_read(doc)

    # ---- PART 1 BANNER ----
    add_part_banner(doc, "PART 1 \u2014 STORYBOARD FOR REVIEW")

    # ---- 1.1 – 1.6 ----
    build_1_1(doc)
    build_1_2(doc)
    build_1_3(doc)
    build_1_4(doc)
    build_1_5(doc)
    build_1_6(doc)

    # ---- PART 2 BANNER ----
    add_part_banner(doc, "PART 2 \u2014 DESIGNER TOOLKIT  (for Yi Jin / Developer)")
    build_part2_intro(doc)

    # ---- Per-screen tables ----
    for screen in SCREEN_DATA:
        add_screen_table(
            doc,
            screen_num=screen["num"],
            screen_title=screen["title"],
            step=screen["step"],
            activities=screen["activities"],
            design_guide=screen["design_guide"],
        )

    # ---- Confidentiality footer via core properties ----
    doc.core_properties.author   = "Yi Jin"
    doc.core_properties.comments = "Confidential \u2014 For Internal Use Only \u00b7 UHN Accessibility First \u00b7 Guide 02"

    doc.save(str(OUTPUT_PATH))
    print(f"Document saved to:\n  {OUTPUT_PATH}")


if __name__ == "__main__":
    build()
